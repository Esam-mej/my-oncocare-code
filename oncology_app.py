from docx.shared import Inches
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import json
import sys
import os
import tempfile
import pandas as pd
from datetime import datetime
from tkcalendar import Calendar
import bcrypt
import matplotlib
matplotlib.use('TkAgg')
import matplotlib.pyplot as plt
from collections import defaultdict
import webbrowser
from docx import Document
from docx.shared import Pt
import shutil
import socket
import threading
import time
from PIL import Image, ImageTk
import firebase_admin
from firebase_admin import credentials, firestore
from firebase_admin.exceptions import FirebaseError
from concurrent.futures import ThreadPoolExecutor
from google.cloud.exceptions import NotFound
from google.auth.exceptions import RefreshError
import schedule
import io
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from googleapiclient.errors import HttpError
from configparser import ConfigParser
import configparser
from tkinter import simpledialog
import subprocess
    
# Path to the JSON file
DROPDOWN_FILE = 'dropdown_lists.json'

def load_users_config():
    """Load user configuration from users_data.json"""
    default = {'last_modified': '1970-01-01T00:00:00', 'data': {}}
    try:
        with open('users_data.json', 'r') as f:
            users_data = json.load(f)
            return {
                'last_modified': datetime.now().isoformat(),
                'data': users_data  # Directly use the existing format
            }
    except (FileNotFoundError, json.JSONDecodeError):
        return default

def save_users_config(config):
    """Save user configuration to users_data.json"""
    with open('users_data.json', 'w') as f:
        json.dump(config['data'], f, indent=4)  # Save only the data part(config, f, indent=4)

def load_dropdown_config():
    """Load dropdown configuration from local file"""
    default = {'last_modified': '1970-01-01T00:00:00', 'data': {}}
    try:
        with open('dropdown_lists.json', 'r') as f:
            config = json.load(f)
            return config if 'data' in config else default
    except:
        return default

def save_dropdown_config(config):
    """Save dropdown configuration to local file"""
    with open('dropdown_lists.json', 'w') as f:
        json.dump(config, f, indent=4)

def validate_dropdown_structure(data):
    if not isinstance(data, dict):
        return False
    for key, value in data.items():
        if not isinstance(value, list):
            return False
        for item in value:
            if not isinstance(item, str):
                return False
    return True

def get_resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    path = os.path.join(base_path, relative_path)
    
    # Special handling for Windows EXE case
    if not os.path.exists(path) and hasattr(sys, '_MEIPASS'):
        # Try parent directory for nested EXE structures
        base_path = os.path.join(base_path, '..')
        path = os.path.join(base_path, relative_path)
    
    return path

# Constants
LAB_RANGES = {
# Hematology
"HB": {"normal_range": "11.5-15.5", "critical_low": "<7", "critical_high": ">20"},
"PLT": {"normal_range": "150-450", "critical_low": "<20", "critical_high": ">1000"},
"WBC": {"normal_range": "4.5-13.5", "critical_low": "<1", "critical_high": ">30"},
"NEUTROPHILS": {"normal_range": "1.5-8.5", "critical_low": "<0.5", "critical_high": ">15"},
"LYMPHOCYTES": {"normal_range": "1.5-7.0", "critical_low": "<0.5", "critical_high": ">10"},
"HCT": {"normal_range": "35-45", "critical_low": "<20", "critical_high": ">60"},

# Chemistry
"UREA": {"normal_range": "10-40", "critical_low": "<10", "critical_high": ">100"},
"CREATININE": {"normal_range": "0.3-1.0", "critical_low": "<0.3", "critical_high": ">2.0"},
"NA": {"normal_range": "135-145", "critical_low": "<120", "critical_high": ">160"},
"K": {"normal_range": "3.5-5.0", "critical_low": "<2.5", "critical_high": ">6.5"},
"CL": {"normal_range": "98-107", "critical_low": "<80", "critical_high": ">115"},
"CA": {"normal_range": "8.8-10.8", "critical_low": "<7", "critical_high": ">13"},
"MG": {"normal_range": "1.7-2.4", "critical_low": "<1", "critical_high": ">4"},
"PHOSPHORUS": {"normal_range": "3.0-6.0", "critical_low": "<1.5", "critical_high": ">8"},
"URIC ACID": {"normal_range": "2.5-7.0", "critical_low": "<1", "critical_high": ">10"},
"LDH": {"normal_range": "100-300", "critical_low": "<50", "critical_high": ">1000"},
"FERRITIN": {"normal_range": "10-300", "critical_low": "<10", "critical_high": ">1000"},

# Liver
"GPT (ALT)": {"normal_range": "5-45", "critical_low": "<5", "critical_high": ">200"},
"GOT (AST)": {"normal_range": "10-40", "critical_low": "<10", "critical_high": ">200"},
"ALK PHOS": {"normal_range": "50-400", "critical_low": "<50", "critical_high": ">1000"},
"T.BILIRUBIN": {"normal_range": "0.2-1.2", "critical_low": "<0.2", "critical_high": ">5"},
"D.BILIRUBIN": {"normal_range": "0-0.4", "critical_low": "<0", "critical_high": ">2"},

# Coagulation
"PT": {"normal_range": "11-14", "critical_low": "<10", "critical_high": ">30"},
"APTT": {"normal_range": "25-35", "critical_low": "<20", "critical_high": ">60"},
"INR": {"normal_range": "0.9-1.2", "critical_low": "<0.8", "critical_high": ">5"},
"FIBRINOGEN": {"normal_range": "200-400", "critical_low": "<100", "critical_high": ">700"},
"D-DIMER": {"normal_range": "<0.5", "critical_low": "<0.1", "critical_high": ">5"},

# ABG
"PH": {"normal_range": "7.35-7.45", "critical_low": "<7.2", "critical_high": ">7.6"},
"PCO2": {"normal_range": "35-45", "critical_low": "<25", "critical_high": ">60"},
"PO2": {"normal_range": "80-100", "critical_low": "<60", "critical_high": ">120"},
"HCO3": {"normal_range": "22-26", "critical_low": "<15", "critical_high": ">35"},
"BASE EXCESS": {"normal_range": "-2 to +2", "critical_low": "<-5", "critical_high": ">5"},

# Viral Screen
"HBSAG": {"normal_range": "Negative", "critical_low": "", "critical_high": "Positive"},
"HCV": {"normal_range": "Negative", "critical_low": "", "critical_high": "Positive"},
"HIV": {"normal_range": "Negative", "critical_low": "", "critical_high": "Positive"},
"CMV": {"normal_range": "Negative", "critical_low": "", "critical_high": "Positive"},
"EBV": {"normal_range": "Negative", "critical_low": "", "critical_high": "Positive"}
}

EF_NOTES_TEMPLATES = [
    "EF < 50%",
    ">10% drop from baseline",
    "Clinical symptoms",
    "Consider cardiology consult"
]

MALIGNANCIES = [
    "ALL", "AML", "LYMPHOMA", "EWING", "OSTEO", "NEUROBLASTOMA",
    "BRAIN T", "RHABDO", "RETINO", "HEPATO", "GERM CELL"
]

DROPDOWN_OPTIONS = {
    "GENDER": ["MALE", "FEMALE"],
    "B_SYMPTOMS": ["FEVER", "NIGHT SWEAT", "WT LOSS", "NO B-SYMPTOMS"],
    "EXAMINATION": [
        "FEBRILE", "HEPATOMEGALLY", "SPLENOMEGALLY", "LYMPHADENOPATHY",
        "TESTICULAR ENLARGMENT", "GUM HYPATROPHY", "CRANIAL PULSY",
        "RACCON EYES", "BONY SWELLING", "JOINT SWELLING", "SUBCONJUNCTIVAL HE",
        "HIGH BP", "HYPOTENSION", "OTHERS"
    ],
    "SYMPTOMS": [
        "ASYMPTOMATIC", "FEVER", "BONE PAIN", "BRUSIS", "WT LOSS",
        "NIGHT SWEATTING", "NECK SWELLING", "JOINT SWELLING",
        "ABDOMINAL DISTENSION", "HEMATURIA", "FRACTURE", "PALLOR",
        "VOMITTING", "DIARRHEA", "CONSTIPATION", "EPISTAXIS",
        "HEMATEMESIS", "GUM BLEEDING", "MELENA", "HEMATOCHEZIA", "OTHERS"
    ],
    "SURGERY": ["INDICATED", "NOT INDICATED", "DONE"],
    "SR_TYPE": [
        "COMPLETE RESECTION", "DEBULKING", "PALLIATIVE",
        "INCOMPLETE RESECTION", "DIAGNOSTIC &/OR STAGING"
    ],
    "STATE": ["ALIVE", "DECEASED", "MISSED FOLLOW UP"],
    "STEROID_R": ["GOOD", "POOR"],
    "TREATMENT_STAGE": [
        "PRE PHASE", "INDUCTION", "CONSOLIDATION",
        "MAINTENANCE", "OFF THERAPY"
    ],
    "RISK_GROUP": ["LRG", "SRG", "IRG", "HRG"],
    "BMT": ["INDICATED", "NO INDICATION", "DONE", "UNKNOWN"],
    "RADIOTHERAPY": ["INDICATED", "NOT INDICATED", "DONE"],
    "NECROSIS_GRADE": [
        "GRADE 1 <50%", "GRADE 2 >50%", "GRADE 3 >90%", "GRADE 4 >100%"
    ],
    "CSF": [
        "INITIAL 0", "INITIAL 1", "INITIAL 2", "INITIAL 3",
        "RELPASE 1", "RELAPSE 2", "RELAPSE 3"
    ],
    "STAGE": [
        "1", "2", "3", "4", "4S", "5", "REFRACTORY 1",
        "REFRACTORY 2", "REFRACTORY 3", "LOCAL", "METS"
    ],
    "THERAPY_SIDE_EFFECTS": [
        "MTX TOXICITY", "LEUKOENCEPHALOPATHY", "HGE CYSTITIS",
        "PERIPHERAL NEUROPATHY", "PULMONARY FIBROSIS",
        "L-ASP. SENSITIVITY", "OTHERS"
    ]
}

MALIGNANCY_FIELDS = {
    "ALL": [
        "GROUP", "CYTOGENETIC", "CSF", "MRD", "STEROID_R", 
        "TREATMENT_STAGE", "CYCLE", "RADIOTHERAPY", "BMT", 
        "STATE", "THERAPY_SIDE_EFFECTS", "NOTES"
    ],
    "AML": [
        "CYTOGENETIC", "CSF", "MRD", "TREATMENT_STAGE", 
        "CYCLE", "RADIOTHERAPY", "BMT", "STATE", 
        "THERAPY_SIDE_EFFECTS", "NOTES"
    ],
    "LYMPHOMA": [
        "GROUP", "STAGE", "CYTOGENETIC", "CSF", "MRD", 
        "STEROID_R", "TREATMENT_STAGE", "CYCLE", 
        "RADIOTHERAPY", "BMT", "STATE", "THERAPY_SIDE_EFFECTS", 
        "NOTES"
    ],
    "EWING": [
        "STAGE", "HISTOPATHOLOGY", "CYTOGENETIC", "SURGERY", 
        "SR_TYPE", "SR_DATE", "TREATMENT_STAGE", "CYCLE", 
        "RADIOTHERAPY", "BMT", "STATE", "THERAPY_SIDE_EFFECTS", 
        "NOTES"
    ],
    "OSTEO": [
        "STAGE", "HISTOPATHOLOGY", "CYTOGENETIC", "SURGERY", 
        "SR_TYPE", "SR_DATE", "TREATMENT_STAGE", "CYCLE", 
        "NECROSIS_GRADE", "RADIOTHERAPY", "BMT", "STATE", 
        "THERAPY_SIDE_EFFECTS", "NOTES"
    ],
    "NEUROBLASTOMA": [
        "STAGE", "HISTOPATHOLOGY", "CYTOGENETIC", "SURGERY", 
        "SR_TYPE", "SR_DATE", "TREATMENT_STAGE", "CYCLE", 
        "RADIOTHERAPY", "BMT", "STATE", "THERAPY_SIDE_EFFECTS", 
        "NOTES"
    ],
    "BRAIN T": [
        "GROUP", "STAGE", "GRADE", "HISTOPATHOLOGY", 
        "CYTOGENETIC", "SURGERY", "SR_TYPE", "SR_DATE", 
        "TREATMENT_STAGE", "CYCLE", "RADIOTHERAPY", "BMT", 
        "STATE", "THERAPY_SIDE_EFFECTS", "NOTES"
    ],
    "RHABDO": [
        "GROUP", "STAGE", "HISTOPATHOLOGY", "CYTOGENETIC", 
        "SURGERY", "SR_TYPE", "SR_DATE", "TREATMENT_STAGE", 
        "CYCLE", "RADIOTHERAPY", "BMT", "STATE", 
        "THERAPY_SIDE_EFFECTS", "NOTES"
    ],
    "RETINO": [
        "GROUP", "STAGE", "GRADE", "HISTOPATHOLOGY", 
        "CYTOGENETIC", "SURGERY", "SR_TYPE", "SR_DATE", 
        "TREATMENT_STAGE", "CYCLE", "RADIOTHERAPY", "BMT", 
        "STATE", "THERAPY_SIDE_EFFECTS", "NOTES"
    ],
    "HEPATO": [
        "GROUP", "STAGE", "GRADE", "HISTOPATHOLOGY", 
        "CYTOGENETIC", "SURGERY", "SR_TYPE", "SR_DATE", 
        "TREATMENT_STAGE", "CYCLE", "RADIOTHERAPY", "BMT", 
        "STATE", "THERAPY_SIDE_EFFECTS", "NOTES"
    ],
    "GERM CELL": [
        "GROUP", "STAGE", "GRADE", "HISTOPATHOLOGY", 
        "CYTOGENETIC", "SURGERY", "SR_TYPE", "SR_DATE", 
        "TREATMENT_STAGE", "CYCLE", "RADIOTHERAPY", "BMT", 
        "STATE", "THERAPY_SIDE_EFFECTS", "NOTES"
    ]
}

COMMON_FIELDS = [
    "FILE NUMBER", "NAME", "GENDER", "DATE OF BIRTH", "NATIONALITY",
    "NATIONAL NUMBER", "ADDRESS", "PHONE NUMBER", "FAMILY HISTORY OF MALIGNANCY",
    "ASSOCIATED DISEASE", "SYMPTOMS", "OTHER HISTORY NOTES", "EXAMINATION", "DIAGNOSIS",
    "AGE ON DIAGNOSIS", "INITIAL WBC"
]

MALIGNANCY_COLORS = {
    "ALL": "#FFCCCC", "AML": "#CCE5FF", "LYMPHOMA": "#FFFFCC",
    "EWING": "#E5FFCC", "OSTEO": "#FFCCE5", "NEUROBLASTOMA": "#CCFFFF",
    "BRAIN T": "#E5CCFF", "RHABDO": "#FFE5CC", "RETINO": "#CCFFCC",
    "HEPATO": "#FFCCFF", "GERM CELL": "#CCE5E5"
}

# Google Drive API Scopes
SCOPES = ['https://www.googleapis.com/auth/drive']
# Google Drive folder ID for the app (will be created if not exists)
DRIVE_FOLDER_NAME = "OncoCare"
DRIVE_PATIENTS_FOLDER_NAME = "OncoCare_Patients"

# Chemo Stocks Google Sheet URL
CHEMO_STOCKS_URL = "https://docs.google.com/spreadsheets/d/1QxcuCM4JLxPyePbaCvB1fdQOmEMgHwAL-bqRCaMvMfM/edit?usp=sharing"

# Supported file extensions for sync
SUPPORTED_EXTENSIONS = {
    'documents': ['.doc', '.docx', '.txt', '.pdf', '.rtf', '.odt', '.wps', '.md'],
    'spreadsheets': ['.xls', '.xlsx', '.xlsm', '.csv', '.ods', '.tsv'],
    'images': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.svg', '.dcm', '.webp', '.heif'],
    'data': ['.json', '.xml', '.yaml', '.csv', '.parquet'],
    'presentations': ['.ppt', '.pptx', '.odp', '.key'],
    'archives': ['.zip', '.rar', '.tar', '.gz', '.7z', '.bz2'],
    'audio': ['.mp3', '.wav', '.aac', '.ogg', '.flac'],
    'video': ['.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv'],
    'markup': ['.html', '.htm', '.xml', '.markdown']
}

class GoogleDriveManager:
    def __init__(self):
        self.service = None
        self.initialized = False
        self.app_folder_id = None
        self.patients_folder_id = None
        self.creds = None  # Add this line
        self.initialize_drive()

    def initialize_drive(self):
        """Initialize Google Drive connection with automatic token refresh"""
        try:
            token_path = 'token.json'
            
            # First try to load credentials
            if os.path.exists(token_path):
                try:
                    self.creds = Credentials.from_authorized_user_file(token_path, SCOPES)
                    # Force credential validation
                    if self.creds and self.creds.expired and self.creds.refresh_token:
                        self.creds.refresh(Request())
                        # Save immediately after refresh
                        with open(token_path, 'w') as token:
                            token.write(self.creds.to_json())
                except (RefreshError, ValueError) as e:
                    print(f"Token error: {e}. Starting fresh authentication...")
                    os.remove(token_path)
                    self.creds = None

            # If no valid credentials, start auth flow
            if not self.creds or not self.creds.valid:
                flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
                self.creds = flow.run_local_server(
                    port=0,
                    authorization_prompt_message='Please visit this URL to authorize OncoCare: {url}',
                    success_message='Authentication complete! You may close this window.',
                    open_browser=True
                )
                # Save new credentials
                with open(token_path, 'w') as token:
                    token.write(self.creds.to_json())

            self.service = build('drive', 'v3', credentials=self.creds)
            self.initialized = True
            self.setup_app_folders()

        except Exception as e:
            print(f"Google Drive initialization failed: {e}")
            self.initialized = False

    def setup_app_folders(self):
        """Set up application folder structure in Drive"""
        try:
            query = f"name='{DRIVE_FOLDER_NAME}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
            results = self.service.files().list(q=query).execute()
            items = results.get('files', [])
            
            if not items:
                folder_metadata = {
                    'name': DRIVE_FOLDER_NAME,
                    'mimeType': 'application/vnd.google-apps.folder'
                }
                folder = self.service.files().create(body=folder_metadata).execute()
                self.app_folder_id = folder['id']
            else:
                self.app_folder_id = items[0]['id']
            
            query = f"name='{DRIVE_PATIENTS_FOLDER_NAME}' and '{self.app_folder_id}' in parents"
            results = self.service.files().list(q=query).execute()
            items = results.get('files', [])
            
            if not self.creds or not self.creds.valid:
                self.initialize_drive()  # Re-initialize if needed
            
            if not items:
                folder_metadata = {
                    'name': DRIVE_PATIENTS_FOLDER_NAME,
                    'parents': [self.app_folder_id],
                    'mimeType': 'application/vnd.google-apps.folder'
                }
                folder = self.service.files().create(body=folder_metadata).execute()
                self.patients_folder_id = folder['id']
            else:
                self.patients_folder_id = items[0]['id']
                
        except HttpError as e:
            print(f"Drive folder setup error: {e}")
            if e.resp.status == 404:
                self.app_folder_id = None
                self.setup_app_folders()
    
    def upload_file(self, file_path, file_name, folder_id=None):
        """Upload a file to Google Drive with improved error handling"""
        if not self.initialized:
            return False
        
        try:
            # Check if file already exists
            query = f"name='{file_name}' and trashed=false"
            if folder_id:
                query += f" and '{folder_id}' in parents"
            
            existing_files = self.service.files().list(q=query, fields="files(id, modifiedTime)").execute().get('files', [])
            
            file_metadata = {
                'name': file_name
            }
            
            if folder_id:
                file_metadata['parents'] = [folder_id]
            
            media = MediaFileUpload(file_path)
            
            if existing_files:
                # Get the existing file's modification time
                drive_mtime = datetime.strptime(existing_files[0].get('modifiedTime'), '%Y-%m-%dT%H:%M:%S.%fZ').timestamp()
                local_mtime = os.path.getmtime(file_path)
                
                # Only update if local file is newer
                if local_mtime > drive_mtime:
                    file_id = existing_files[0]['id']
                    file = self.service.files().update(
                        fileId=file_id,
                        body=file_metadata,
                        media_body=media,
                        fields='id'
                    ).execute()
                    print(f"Updated file: {file_name}")
                else:
                    print(f"File {file_name} is up to date in Drive")
                    return True
            else:
                # Create new file
                file = self.service.files().create(
                    body=file_metadata,
                    media_body=media,
                    fields='id'
                ).execute()
                print(f"Uploaded new file: {file_name}")
            
            return True
        except HttpError as error:
            print(f"An error occurred while uploading file: {error}")
            return False
        except Exception as e:
            print(f"Error uploading file {file_name}: {e}")
            return False
    
    def download_file(self, file_id, save_path):
        """Download a file from Google Drive by file ID"""
        if not self.initialized:
            return False
        
        try:
            request = self.service.files().get_media(fileId=file_id)
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            
            fh = io.FileIO(save_path, 'wb')
            downloader = MediaIoBaseDownload(fh, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()
                print(f"Download {int(status.progress() * 100)}%")
            
            return True
        except HttpError as error:
            print(f"An error occurred while downloading file: {error}")
            return False
        except Exception as e:
            print(f"Error downloading file: {e}")
            return False
    
    def create_patient_folder(self, file_no):
        """Create a folder for a patient in Google Drive"""
        if not self.initialized or not self.patients_folder_id:
            return None
        
        folder_name = f"Patient_{file_no}"
        try:
            # Check if folder already exists
            query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder' and '{self.patients_folder_id}' in parents and trashed=false"
            results = self.service.files().list(q=query, fields="files(id)").execute()
            items = results.get('files', [])
            
            if items:
                return items[0]['id']
            
            # Create new folder
            folder_metadata = {
                'name': folder_name,
                'mimeType': 'application/vnd.google-apps.folder',
                'parents': [self.patients_folder_id],
                'description': f'Patient folder for file number {file_no}'
            }
            folder = self.service.files().create(body=folder_metadata, fields='id').execute()
            return folder.get('id')
        except Exception as e:
            print(f"Error creating patient folder: {e}")
            return None
    
    def upload_patient_data(self, patient_data):
        """Upload patient data to Google Drive"""
        if not self.initialized:
            return False
        
        try:
            # Save patient data to a temporary file
            temp_file = "temp_patient_data.json"
            with open(temp_file, 'w') as f:
                json.dump(patient_data, f)
            
            # Upload to the app folder
            file_name = f"patient_{patient_data['FILE NUMBER']}.json"
            success = self.upload_file(temp_file, file_name, self.app_folder_id)
            
            # Remove temporary file
            os.remove(temp_file)
            
            return success
        except Exception as e:
            print(f"Error uploading patient data: {e}")
            return False
    
    def download_all_patient_data(self):
        """Download all patient data from Google Drive"""
        if not self.initialized or not self.app_folder_id:
            return None
        
        try:
            # List all patient data files
            query = f"'{self.app_folder_id}' in parents and name contains 'patient_' and name contains '.json' and trashed=false"
            results = self.service.files().list(q=query, fields="files(id, name)").execute()
            items = results.get('files', [])
            
            all_patients = []
            
            for item in items:
                # Download each file
                request = self.service.files().get_media(fileId=item['id'])
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                
                done = False
                while done is False:
                    status, done = downloader.next_chunk()
                
                # Parse the JSON data
                fh.seek(0)
                patient_data = json.load(fh)
                all_patients.append(patient_data)
            
            return all_patients
        except Exception as e:
            print(f"Error downloading patient data: {e}")
            return None
    
    def sync_patient_files(self, file_no):
        """Sync all files for a specific patient between local and Google Drive with two-way sync"""
        if not self.initialized:
            return False, "Google Drive not initialized"
        
        local_folder = f"Patient_{file_no}"
        if not os.path.exists(local_folder):
            os.makedirs(local_folder, exist_ok=True)
        
        # Get or create the patient folder in Drive
        drive_folder_id = self.create_patient_folder(file_no)
        if not drive_folder_id:
            return False, "Failed to create/get patient folder in Drive"
        
        try:
            # List files in local folder
            local_files = {}
            for root, _, files in os.walk(local_folder):
                for file in files:
                    file_path = os.path.join(root, file)
                    rel_path = os.path.relpath(file_path, local_folder)
                    local_files[rel_path] = {
                        'path': file_path,
                        'mtime': os.path.getmtime(file_path),
                        'size': os.path.getsize(file_path)
                    }
            
            # List files in Drive folder
            query = f"'{drive_folder_id}' in parents and trashed=false"
            results = self.service.files().list(
                q=query, 
                fields="files(id, name, modifiedTime, size, mimeType)"
            ).execute()
            drive_files = {
                item['name']: {
                    'id': item['id'],
                    'mtime': datetime.strptime(item['modifiedTime'], '%Y-%m-%dT%H:%M:%S.%fZ').timestamp(),
                    'size': int(item.get('size', 0)),
                    'mimeType': item['mimeType']
                } 
                for item in results.get('files', [])
            }
            
            # Process files for two-way sync
            upload_count = 0
            download_count = 0
            
            # Upload new or modified files to Drive
            for rel_path, local_info in local_files.items():
                file_name = os.path.basename(rel_path)
                
                # Skip unsupported file types
                if not self.is_supported_file(file_name):
                    continue
                
                if file_name not in drive_files:
                    # New file - upload to Drive
                    if self.upload_file(local_info['path'], file_name, drive_folder_id):
                        upload_count += 1
                else:
                    # File exists in both - compare modification times
                    drive_info = drive_files[file_name]
                    if local_info['mtime'] > drive_info['mtime']:
                        # Local file is newer - upload to Drive
                        if self.upload_file(local_info['path'], file_name, drive_folder_id):
                            upload_count += 1
            
            # Download files from Drive that don't exist locally or are newer
            for file_name, drive_info in drive_files.items():
                local_path = os.path.join(local_folder, file_name)
                
                if not os.path.exists(local_path):
                    # File doesn't exist locally - download
                    if self.download_file(drive_info['id'], local_path):
                        download_count += 1
                else:
                    # File exists - compare modification times
                    local_mtime = os.path.getmtime(local_path)
                    if drive_info['mtime'] > local_mtime:
                        # Drive file is newer - download
                        if self.download_file(drive_info['id'], local_path):
                            download_count += 1
            
            return True, f"Sync complete. Uploaded: {upload_count}, Downloaded: {download_count}"
        except Exception as e:
            print(f"Error syncing patient files: {e}")
            return False, f"Sync failed: {str(e)}"
    
    def is_supported_file(self, filename):
        """Check if file extension is supported for sync"""
        ext = os.path.splitext(filename)[1].lower()
        for category, extensions in SUPPORTED_EXTENSIONS.items():
            if ext in extensions:
                return True
        return False
    
    def get_drive_file_mtime(self, file_id):
        """Get the modification time of a file in Google Drive by file ID"""
        try:
            file = self.service.files().get(fileId=file_id, fields='modifiedTime').execute()
            modified_time = file.get('modifiedTime')
            if modified_time:
                return datetime.strptime(modified_time, '%Y-%m-%dT%H:%M:%S.%fZ').timestamp()
        except Exception as e:
            print(f"Error getting file modification time: {e}")
        
        return 0

class FirebaseManager:
    def __init__(self):
        self.db = None
        self.initialized = False
        self.initialize_firebase()

    def initialize_firebase(self):
        """Initialize Firebase connection with error handling"""
        try:
            if not firebase_admin._apps:
                cred_path = self.get_resource_path('serviceAccountKey.json')
                if os.path.exists(cred_path):
                    cred = credentials.Certificate(cred_path)
                    firebase_admin.initialize_app(cred)
                    self.db = firestore.client()
                    self.initialized = True
                else:
                    print("Firebase credentials not found. Offline mode only.")
        except Exception as e:
            print(f"Firebase initialization failed: {e}")
            self.initialized = False

    def get_resource_path(self, relative_path):
        """Get absolute path to resource"""
        try:
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")
        return os.path.join(base_path, relative_path)
    
    def sync_patients(self, local_data):
        """Synchronize patient data with Firebase with improved conflict resolution"""
        if not self.initialized:
            return False, "Firebase not initialized"
        
        try:
            firebase_data = []
            docs = self.db.collection('patients').stream()
            for doc in docs:
                patient = doc.to_dict()
                patient['_firestore_id'] = doc.id
                firebase_data.append(patient)
            
            merged_data = self.merge_data(local_data, firebase_data)
            
            batch = self.db.batch()
            patients_ref = self.db.collection('patients')
            
            for patient in merged_data:
                doc_id = patient.pop('_firestore_id', None) or patient['FILE NUMBER']
                doc_ref = patients_ref.document(doc_id)
                batch.set(doc_ref, patient)
            
            batch.commit()
            return True, "Sync completed successfully"
        
        except Exception as e:
            return False, f"Unexpected error: {e}"
    
    def sync_users(self, local_users):
        """Synchronize user data with Firebase"""
        if not self.initialized:
            return False, "Firebase not initialized"
        
        try:
            firebase_users = []
            docs = self.db.collection('users').stream()
            for doc in docs:
                user = doc.to_dict()
                user['_firestore_id'] = doc.id
                firebase_users.append(user)
            
            merged_users = self.merge_users(local_users, firebase_users)
            
            batch = self.db.batch()
            users_ref = self.db.collection('users')
            
            for user in merged_users:
                doc_id = user.pop('_firestore_id', None) or user['username']
                doc_ref = users_ref.document(doc_id)
                batch.set(doc_ref, user)
            
            batch.commit()
            return True, "User sync completed"
        
        except Exception as e:
            return False, f"User sync error: {str(e)}"
    
    def sync_dropdowns(self, local_dropdowns):
        """Enhanced sync with default values preservation"""
        try:
            firebase_dropdowns = self.get_all_dropdowns()
            merged = {'data': {}}

            # Always include default values first
            default_data = self.get_all_dropdowns()['data']
            for name, data in default_data.items():
                merged['data'][name] = data

            # Merge Firebase data
            for name, fire_data in firebase_dropdowns['data'].items():
                local_data = local_dropdowns.get(name, {})
                fire_time = datetime.fromisoformat(fire_data['LAST_MODIFIED'])
                local_time = datetime.fromisoformat(local_data.get('LAST_MODIFIED', '1970-01-01T00:00:00'))
                
                if local_time > fire_time:
                    merged['data'][name] = local_data
                else:
                    merged['data'][name] = fire_data

            # Push merged data back to Firebase
            batch = self.db.batch()
            dropdown_ref = self.db.collection('dropdowns')
            
            for name, data in merged['data'].items():
                doc_ref = dropdown_ref.document(name)
                batch.set(doc_ref, {
                    'values': data['values'],
                    'LAST_MODIFIED': firestore.SERVER_TIMESTAMP
                })
            
            batch.commit()
            
            # Save merged data locally with atomic write
            temp_path = f"{DROPDOWN_FILE}.tmp"
            with open(temp_path, 'w', encoding='utf-8') as f:
                json.dump(merged, f, indent=4, ensure_ascii=False)
            
            if os.path.exists(DROPDOWN_FILE):
                os.replace(temp_path, DROPDOWN_FILE)
            else:
                os.rename(temp_path, DROPDOWN_FILE)
            
            return True, "Dropdowns synced successfully"
            
        except Exception as e:
            print(f"Dropdown sync error: {e}")
            return False, f"Dropdown sync failed: {str(e)}"
    
    def merge_data(self, local_data, firebase_data):
        """Merge local and Firebase data with improved conflict resolution"""
        merged = []
        local_by_id = {p['FILE NUMBER']: p for p in local_data}
        firebase_by_id = {p['FILE NUMBER']: p for p in firebase_data}
        
        all_ids = set(local_by_id.keys()).union(set(firebase_by_id.keys()))
        
        for file_no in all_ids:
            local_patient = local_by_id.get(file_no)
            firebase_patient = firebase_by_id.get(file_no)
            
            if local_patient and not firebase_patient:
                merged.append(local_patient)
            elif firebase_patient and not local_patient:
                merged.append(firebase_patient)
            else:
                try:
                    local_date = datetime.strptime(local_patient['LAST_MODIFIED_DATE'], '%d/%m/%Y %H:%M:%S')
                    firebase_date = datetime.strptime(firebase_patient['LAST_MODIFIED_DATE'], '%d/%m/%Y %H:%M:%S')
                    
                    if local_date > firebase_date:
                        local_patient['_firestore_id'] = firebase_patient.get('_firestore_id', file_no)
                        merged.append(local_patient)
                    else:
                        merged.append(firebase_patient)
                except:
                    local_patient['_conflict'] = True
                    merged.append(local_patient)
                    merged.append(firebase_patient)
        
        return merged
    
    def merge_users(self, local_users, firebase_users):
        """Merge user data with conflict resolution"""
        merged = []
        local_by_id = {u['username']: u for u in local_users}
        firebase_by_id = {u['username']: u for u in firebase_users}
        
        all_usernames = set(local_by_id.keys()).union(set(firebase_by_id.keys()))
        
        for username in all_usernames:
            local = local_by_id.get(username)
            remote = firebase_by_id.get(username)
            
            if not remote:
                merged.append(local)
                continue
            if not local:
                merged.append(remote)
                continue
                
            local_date = datetime.strptime(local.get('LAST_MODIFIED', '1970-01-01'), '%Y-%m-%d %H:%M:%S')
            remote_date = datetime.strptime(remote.get('LAST_MODIFIED', '1970-01-01'), '%Y-%m-%d %H:%M:%S')
            
            if local_date > remote_date:
                merged.append(local)
            else:
                merged.append(remote)
        
        return merged
    
    def merge_dropdowns(self, local_dropdowns, firebase_dropdowns):
        """Merge dropdowns with timestamp comparison"""
        merged = []
        
        # Convert both to comparable format
        local_map = {d['name']: d for d in local_dropdowns}
        firebase_map = {d['name']: d for d in firebase_dropdowns}
        
        all_names = set(local_map.keys()).union(set(firebase_map.keys()))
        
        for name in all_names:
            local = local_map.get(name)
            remote = firebase_map.get(name)
            
            if not remote:
                merged.append(local)
                continue
            if not local:
                merged.append(remote)
                continue
                
            # Convert timestamps to datetime objects
            try:
                # Handle local timestamp (string)
                local_time = datetime.strptime(local['LAST_MODIFIED'], '%Y-%m-%d %H:%M:%S')
                
                # Handle Firebase timestamp (could be Timestamp object or string)
                if isinstance(remote['LAST_MODIFIED'], firestore.firestore.Timestamp):
                    remote_time = remote['LAST_MODIFIED'].to_datetime()
                else:
                    remote_time = datetime.fromisoformat(remote['LAST_MODIFIED'])
                
                if local_time > remote_time:
                    merged.append(local)
                else:
                    merged.append(remote)
            except Exception as e:
                print(f"Timestamp conversion error: {e}")
                merged.append(remote if 'LAST_MODIFIED' in remote else local)
        
        return merged
    
    def get_all_patients(self):
        """Retrieve all patients from Firebase with error handling"""
        if not self.initialized:
            return None
        
        try:
            patients = []
            docs = self.db.collection('patients').stream()
            for doc in docs:
                patient = doc.to_dict()
                patient['_firestore_id'] = doc.id
                patients.append(patient)
            return patients
        except Exception as e:
            print(f"Error fetching patients from Firebase: {e}")
            return None
    
    def get_all_users(self):
        """Retrieve all users from Firebase"""
        try:
            users = []
            docs = self.db.collection('users').stream()
            for doc in docs:
                user = doc.to_dict()
                users.append(user)
            return users
        except:
            return None
    
    def get_all_dropdowns(self):
        """Retrieve dropdowns with enhanced validation and default fallbacks"""
        try:
            default_dropdowns = {
                'data': {
                    'GENDER': {'values': ['MALE', 'FEMALE'], 'LAST_MODIFIED': '1970-01-01T00:00:00'},
                    'B_SYMPTOMS': {'values': ['FEVER', 'NIGHT SWEAT', 'WT LOSS', 'NO B-SYMPTOMS'], 'LAST_MODIFIED': '1970-01-01T00:00:00'},
                    # Add other default dropdowns here
                }
            }
            
            if not self.initialized:
                return default_dropdowns

            dropdowns = {'data': {}}
            docs = self.db.collection('dropdowns').stream()
            
            for doc in docs:
                try:
                    data = doc.to_dict()
                    valid_values = data.get('values', [])
                    if not isinstance(valid_values, list):
                        valid_values = []
                    
                    dropdowns['data'][doc.id] = {
                        'values': [str(v) for v in valid_values],
                        'LAST_MODIFIED': self.parse_firebase_timestamp(data.get('LAST_MODIFIED')) or datetime.now().isoformat()
                    }
                except Exception as doc_error:
                    print(f"Skipping invalid dropdown {doc.id}: {doc_error}")

            # Merge with defaults to ensure completeness
            for name, default_data in default_dropdowns['data'].items():
                if name not in dropdowns['data']:
                    dropdowns['data'][name] = default_data

            return dropdowns
            
        except Exception as e:
            print(f"Firebase dropdown fetch error: {e}")
            return default_dropdowns

    def parse_firebase_timestamp(self, timestamp):
        """Handle various timestamp formats with fallbacks"""
        if timestamp is None:
            return datetime.now().isoformat()
            
        try:
            if isinstance(timestamp, firestore.firestore.Timestamp):
                return timestamp.to_datetime().isoformat()
            elif isinstance(timestamp, datetime):
                return timestamp.isoformat()
            else:
                return datetime.fromisoformat(timestamp).isoformat()
        except:
            return datetime.now().isoformat()

class OncologyApp:
    def __init__(self, root):
        self.root = root
        self.root.title("OncoCare - Pediatric Oncology Patient Management System")
        self.root.geometry("1200x800")
        self.root.state('zoomed')
        self.setup_keyboard_shortcuts()
        self.setup_login_screen()

        self.setup_protocols_config()
        self.setup_chemo_sheets_config()
        self.setup_drug_documentation_config()
        self.setup_fn_documentation_config()

        # Initialize connection status
        self.internet_connected = False
        self.sync_in_progress = False
        self.last_sync_time = None
        
        # Initialize services
        self.firebase = FirebaseManager()
        self.drive = GoogleDriveManager()
        
        # Check and restore essential files
        self.ensure_essential_files()
        
        # Load data
        self.load_users()
        self.load_patient_data()
        
        # Setup UI components
        self.setup_styles()
        self.setup_status_bar()
        self.check_internet_connection()
        
        # Initialize F&N Documentation configuration
        self.setup_fn_documentation_config()
        self.current_user = os.getenv("USERNAME").lower()  # Get Windows username
        
        # Final initialization
        self.executor = ThreadPoolExecutor(max_workers=4)
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    def setup_fn_documentation_config(self):
        """Setup configuration for F&N Documentation software"""
        self.fn_config = ConfigParser()
        self.fn_config_file = os.path.join(os.environ['APPDATA'], 'fn_config.ini')
        
        if not os.path.exists(self.fn_config_file):
            self.fn_config['FN_DOCUMENTATION'] = {'path': 'C:\\Path\\To\\Your\\Software.exe'}
            with open(self.fn_config_file, 'w') as configfile:
                self.fn_config.write(configfile)
        else:
            self.fn_config.read(self.fn_config_file)

    def setup_drug_documentation_config(self):
        """Setup configuration for Drug Documentation with encoding fixes"""
        self.drug_config = ConfigParser()
        self.drug_config_file = os.path.join(os.environ['APPDATA'], 'drug_config.ini')
        
        # Default path with safe characters
        default_path = os.path.join(os.path.expanduser('~'), 'Documents', 'Drug_Docs')
        
        try:
            # Read existing config with UTF-8 encoding
            if os.path.exists(self.drug_config_file):
                with open(self.drug_config_file, 'r', encoding='utf-8') as configfile:
                    self.drug_config.read_file(configfile)
                
                # Validate config structure
                if not self.drug_config.has_section('DRUG_DOCUMENTATION'):
                    raise ValueError("Invalid config structure")
                    
                current_path = self.drug_config.get('DRUG_DOCUMENTATION', 'path', fallback=default_path)
                
                # Validate path format
                if not isinstance(current_path, str) or len(current_path) < 2:
                    raise ValueError("Invalid path format")
                    
            else:
                # Create new config with safe defaults
                self.drug_config['DRUG_DOCUMENTATION'] = {'path': default_path}
                with open(self.drug_config_file, 'w', encoding='utf-8') as configfile:
                    self.drug_config.write(configfile)
                
            # Ensure directory exists
            current_path = self.drug_config.get('DRUG_DOCUMENTATION', 'path', fallback=default_path)
            os.makedirs(current_path, exist_ok=True)
            
        except Exception as e:
            # Reset to default config on error
            self.drug_config = ConfigParser()
            self.drug_config['DRUG_DOCUMENTATION'] = {'path': default_path}
            with open(self.drug_config_file, 'w', encoding='utf-8') as configfile:
                self.drug_config.write(configfile)
            os.makedirs(default_path, exist_ok=True)

    def get_drug_documentation_path(self):
        """Safely get path with encoding fixes"""
        try:
            with open(self.drug_config_file, 'r', encoding='utf-8') as configfile:
                self.drug_config.read_file(configfile)
                
            path = self.drug_config.get('DRUG_DOCUMENTATION', 'path', 
                                      fallback=os.path.join(os.path.expanduser('~'), 'Documents', 'Drug_Docs'))
            
            # Validate path format
            if not isinstance(path, str) or len(path) < 2:
                raise ValueError("Invalid path format")
                
            return os.path.normpath(path)
            
        except Exception as e:
            # Return safe default path
            default_path = os.path.join(os.path.expanduser('~'), 'Documents', 'Drug_Docs')
            os.makedirs(default_path, exist_ok=True)
            return default_path

    def set_drug_documentation_path(self, new_path):
        """Set new path for Drug Documentation folder"""
        self.drug_config['DRUG_DOCUMENTATION']['path'] = new_path
        with open(self.drug_config_file, 'w') as configfile:
            self.drug_config.write(configfile)

    def handle_drug_documentation(self):
        """Handle Drug Documentation button click (same as Protocols)"""
        try:
            folder_path = self.get_drug_documentation_path()
            
            dialog = tk.Toplevel(self.root)
            dialog.title("Drug Documentation")
            dialog.geometry("400x200")
            
            msg = f"Current Drug Documentation folder:\n{folder_path}"
            ttk.Label(dialog, text=msg, wraplength=380).pack(pady=20)
            
            btn_frame = ttk.Frame(dialog)
            btn_frame.pack(pady=10)
            
            ttk.Button(btn_frame, text="Open Folder", 
                     command=lambda: self.open_folder(folder_path, dialog),
                     style='Green.TButton').pack(side=tk.LEFT, padx=10)
            
            if self.current_user.lower() == "mej.esam":
                ttk.Button(btn_frame, text="Change Path", 
                          command=lambda: self.change_drug_path(dialog),
                          style='Blue.TButton').pack(side=tk.LEFT, padx=10)
            
            ttk.Button(btn_frame, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT, padx=10)

        except Exception as e:
            messagebox.showerror("Error", f"Drug Docs Error: {str(e)}")

    def open_folder(self, path, dialog=None):
        """Generic folder opening function"""
        try:
            if sys.platform == 'win32':
                os.startfile(path)
            else:
                subprocess.run(['open', path] if sys.platform == 'darwin' 
                             else ['xdg-open', path], check=True)
            if dialog:
                dialog.destroy()
        except Exception as e:
            messagebox.showerror("Error", f"Could not open folder: {str(e)}")

    def change_drug_path(self, dialog):
        """Change path with comprehensive error handling"""
        new_path = filedialog.askdirectory(title="Select Drug Documentation Folder")
        if new_path:
            try:
                # Normalize and validate path
                new_path = os.path.normpath(new_path)
                if len(new_path) < 2 or not isinstance(new_path, str):
                    raise ValueError("Invalid path format")
                
                # Update config with proper encoding
                self.drug_config.set('DRUG_DOCUMENTATION', 'path', new_path)
                with open(self.drug_config_file, 'w', encoding='utf-8') as configfile:
                    self.drug_config.write(configfile)
                
                # Create directory with safe permissions
                os.makedirs(new_path, exist_ok=True, mode=0o755)
                
                messagebox.showinfo("Success", f"Path updated to:\n{new_path}")
                dialog.destroy()
                
            except Exception as e:
                messagebox.showerror("Error", 
                    f"Failed to update path: {str(e)}\n"
                    "Please choose a different location with:\n"
                    "- Standard Latin characters\n"
                    "- No special symbols\n"
                    "- Short path length")

    def open_drug_documentation(self, dialog=None):
        """Open the Drug Documentation folder"""
        try:
            folder_path = self.get_drug_documentation_path()
            if not os.path.exists(folder_path):
                messagebox.showerror("Error", f"Folder not found at:\\n{folder_path}")
                return
            
            os.startfile(folder_path)
            if dialog:
                dialog.destroy()
            messagebox.showinfo("Success", "Drug Documentation folder opened successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Could not open folder: {str(e)}")

    def change_drug_documentation_path(self, dialog):
        """Change the path to Drug Documentation folder"""
        new_path = filedialog.askdirectory(
            title="Select Drug Documentation Folder"
        )
        
        if new_path:
            self.set_drug_documentation_path(new_path)
            messagebox.showinfo("Success", f"Path updated to:\\n{new_path}")
            dialog.destroy()

    def get_fn_documentation_path(self):
        """Get the path to F&N Documentation software"""
        self.fn_config.read(self.fn_config_file)
        return self.fn_config['FN_DOCUMENTATION']['path']

    def set_fn_documentation_path(self, new_path):
        """Set new path for F&N Documentation software"""
        self.fn_config['FN_DOCUMENTATION']['path'] = new_path
        with open(self.fn_config_file, 'w') as configfile:
            self.fn_config.write(configfile)

    def handle_fn_documentation(self):
        """Handle F&N Documentation button click"""
        software_path = self.get_fn_documentation_path()
        
        dialog = tk.Toplevel(self.root)
        dialog.title("F&N Documentation")
        dialog.geometry("400x200")
        
        msg = f"Do you want to start F&N Documentation?\n\nCurrent path: {software_path}"
        ttk.Label(dialog, text=msg, wraplength=380).pack(pady=20)
        
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="Launch", 
                  command=lambda: self.launch_fn_documentation(dialog),
                  style='Green.TButton').pack(side=tk.LEFT, padx=10)
        
        if self.current_user.lower() == "mej.esam":
            ttk.Button(btn_frame, text="Change Path", 
                      command=lambda: self.change_fn_documentation_path(dialog),
                      style='Blue.TButton').pack(side=tk.LEFT, padx=10)
        
        ttk.Button(btn_frame, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT, padx=10)

    def launch_fn_documentation(self, dialog=None):
        """Launch the F&N Documentation software"""
        try:
            software_path = self.get_fn_documentation_path()
            if not os.path.exists(software_path):
                messagebox.showerror("Error", f"Software not found at:\n{software_path}")
                return
            
            subprocess.Popen(software_path)
            if dialog:
                dialog.destroy()
            messagebox.showinfo("Success", "F&N Documentation started successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Could not start software: {str(e)}")

    def change_fn_documentation_path(self, dialog):
        """Change the path to F&N Documentation software"""
        new_path = filedialog.askopenfilename(
            title="Select F&N Documentation executable",
            filetypes=[("Executable files", "*.exe"), ("All files", "*.*")]
        )
        
        if new_path:
            self.set_fn_documentation_path(new_path)
            messagebox.showinfo("Success", f"Path updated to:\n{new_path}")
            dialog.destroy()

    def ensure_essential_files(self):
        """Guaranteed file creation with multiple fallbacks"""
        max_attempts = 3
        for attempt in range(max_attempts):
            try:
                if not os.path.exists(DROPDOWN_FILE):
                    if self.firebase.initialized:
                        if self.fetch_dropdowns_from_firebase(force=True):
                            return
                    # Fallback to embedded defaults
                    self._create_default_dropdown_file()
                    return
            except Exception as e:
                print(f"File creation attempt {attempt+1} failed: {e}")
                if attempt == max_attempts - 1:
                    self._create_default_dropdown_file()

    def _create_default_dropdown_file(self):
        """Nuclear option: create default file from scratch"""
        default_data = {
            'data': {
                'GENDER': {'values': ['MALE', 'FEMALE'], 'LAST_MODIFIED': datetime.now().isoformat()},
                'B_SYMPTOMS': {'values': ['FEVER', 'NIGHT SWEAT', 'WT LOSS', 'NO B-SYMPTOMS'], 'LAST_MODIFIED': datetime.now().isoformat()},
                # Add all other default dropdowns here
            }
        }
        with open(DROPDOWN_FILE, 'w', encoding='utf-8') as f:
            json.dump(default_data, f, indent=4, ensure_ascii=False)
        print("Created new dropdown file with default values")

    def fetch_dropdowns_from_firebase(self, force=False):
        """Force fetch dropdowns with emergency defaults"""
        try:
            success = False
            firebase_dropdowns = self.firebase.get_all_dropdowns()
            
            if firebase_dropdowns.get('data'):
                temp_path = f"{DROPDOWN_FILE}.tmp"
                with open(temp_path, 'w', encoding='utf-8') as f:
                    json.dump(firebase_dropdowns, f, indent=4, ensure_ascii=False)
                
                if os.path.exists(DROPDOWN_FILE):
                    os.replace(temp_path, DROPDOWN_FILE)
                else:
                    os.rename(temp_path, DROPDOWN_FILE)
                print("Dropdowns fully restored from Firebase")
                success = True
            else:
                print("No valid dropdown data in Firebase, using defaults")
                
            if not success and force:
                messagebox.showerror("Sync Error", 
                    "Could not retrieve dropdowns from Firebase. Using default values.")
                
            return success
            
        except Exception as e:
            print(f"Final dropdown restoration error: {e}")
            if force:
                messagebox.showerror("Critical Error", 
                    "Could not restore dropdowns. Please check internet connection and Firebase configuration.")
            return False
    
    def setup_keyboard_shortcuts(self):
        """Configure keyboard shortcuts for quick navigation."""
        # General shortcuts
        self.root.bind('<Control-q>', lambda e: self.on_close())  # Ctrl+Q to quit
        self.root.bind('<Control-l>', lambda e: self.setup_login_screen())  # Ctrl+L to logout
        self.root.bind('<F1>', lambda e: self.show_help())  # F1 for help
        
        # Patient management
        self.root.bind('<Control-n>', lambda e: self.add_patient())  # Ctrl+N for new patient
        self.root.bind('<Control-f>', lambda e: self.search_patient())  # Ctrl+F to search
        
        # Navigation
        self.root.bind('<Escape>', lambda e: self.main_menu())  # ESC to return to main menu

    def show_help(self):
        """Display keyboard shortcuts help dialog."""
        shortcuts = {
            "Ctrl+Q": "Quit the application",
            "Ctrl+L": "Log out",
            "F1": "Show this help dialog",
            "Ctrl+N": "Add new patient",
            "Ctrl+F": "Search patients",
            "Esc": "Return to main menu"
        }
        
        help_text = "Keyboard Shortcuts:\n\n" + "\n".join(
            f"{key}: {desc}" for key, desc in shortcuts.items()
        )
        
        messagebox.showinfo("Keyboard Shortcuts", help_text)

    def setup_styles(self):
        """Configure the visual styles for the application"""
        self.style = ttk.Style()
        self.style.theme_use('clam')
        
        # Configure colors
        self.primary_color = '#2c3e50'  # Dark blue
        self.secondary_color = '#3498db'  # Blue
        self.accent_color = '#e74c3c'  # Red
        self.light_color = '#ecf0f1'  # Light gray
        self.dark_color = '#2c3e50'  # Dark blue
        self.success_color = '#27ae60'  # Green
        self.warning_color = '#f39c12'  # Orange
        self.danger_color = '#e74c3c'  # Red
        
        # Configure styles
        self.style.configure('TFrame', background=self.light_color)
        self.style.configure('TButton', font=('Helvetica', 10, 'bold'), padding=8)
        self.style.map('TButton',
            background=[('active', self.secondary_color)],
            foreground=[('active', 'white')]
        )
        
        # Custom button styles
        self.style.configure('Blue.TButton', background=self.secondary_color, foreground='white')
        self.style.map('Blue.TButton',
            background=[('active', '#2980b9')],
            foreground=[('active', 'white')]
        )
        
        self.style.configure('Green.TButton', background=self.success_color, foreground='white')
        self.style.map('Green.TButton',
            background=[('active', '#219653')],
            foreground=[('active', 'white')]
        )
        
        self.style.configure('Yellow.TButton', background=self.warning_color, foreground='white')
        self.style.map('Yellow.TButton',
            background=[('active', '#e67e22')],
            foreground=[('active', 'white')]
        )
        
        self.style.configure('Brown.TButton', background='#8B4513', foreground='white')
        self.style.map('Brown.TButton',
            background=[('active', '#A0522D')],
            foreground=[('active', 'white')]
        )
        
        self.style.configure('Purple.TButton', background='#9b59b6', foreground='white')
        self.style.map('Purple.TButton',
            background=[('active', '#8e44ad')],
            foreground=[('active', 'white')]
        )

        self.style.configure('Orange.TButton', background='#e67e22', foreground='white')
        self.style.map('Orange.TButton',
            background=[('active', '#d35400')],
            foreground=[('active', 'white')]
        )

        self.style.configure('Red.TButton', background=self.danger_color, foreground='white')
        self.style.map('Red.TButton',
            background=[('active', '#c0392b')],
            foreground=[('active', 'white')]
        )
        
        self.style.configure('TLabel', font=('Helvetica', 10), background=self.light_color)
        self.style.configure('Header.TLabel', font=('Helvetica', 24, 'bold'), foreground=self.primary_color)
        self.style.configure('Subheader.TLabel', font=('Helvetica', 18), foreground=self.secondary_color)
        self.style.configure('TEntry', font=('Helvetica', 10), padding=5)
        self.style.configure('TCombobox', font=('Helvetica', 10))
        self.style.configure('Status.TFrame', background=self.primary_color)
        self.style.configure('Status.TLabel', background=self.primary_color, foreground='white')
        self.style.configure('Malignancy.TCombobox', font=('Helvetica', 12, 'bold'))
    
    def setup_status_bar(self):
        """Setup the status bar with hidden sync button"""
        self.status_frame = ttk.Frame(self.root, style='Status.TFrame', height=30)
        self.status_frame.pack(side=tk.BOTTOM, fill=tk.X)
        
        # User info
        self.user_label = ttk.Label(self.status_frame, text="Not logged in", style='Status.TLabel')
        self.user_label.pack(side=tk.LEFT, padx=10)
        
        # Internet status indicator
        self.internet_indicator = tk.Canvas(self.status_frame, width=20, height=20, highlightthickness=0)
        self.internet_indicator.pack(side=tk.LEFT, padx=5)
        self.update_internet_indicator()
        
        # Sync status
        self.sync_status = ttk.Label(self.status_frame, text="", style='Status.TLabel')
        self.sync_status.pack(side=tk.LEFT, padx=10)
        
        # Date and time
        self.datetime_label = ttk.Label(self.status_frame, text="", style='Status.TLabel')
        self.datetime_label.pack(side=tk.RIGHT, padx=10)
        
        # Create but hide sync button initially
        self.sync_btn = ttk.Button(self.status_frame, text="Synchronize Data", 
                                 command=self.sync_data, style='Blue.TButton')
        
        # Don't pack it yet
        self.update_datetime()
    
    def load_patient_data(self):
        """Load patient data from file or create empty list"""
        if not os.path.exists('patients_data.json'):
            self.patient_data = []
            return
            
        try:
            with open('patients_data.json', 'r') as f:
                self.patient_data = json.load(f)
        except json.JSONDecodeError:
            self.patient_data = []
    
    def save_patient_data(self):
        """Save patient data to file"""
        with open('patients_data.json', 'w') as f:
            json.dump(self.patient_data, f, indent=4)
    
    def update_internet_indicator(self):
        """Update the internet connection indicator"""
        color = "green" if self.internet_connected else "red"
        self.internet_indicator.delete("all")
        self.internet_indicator.create_oval(5, 5, 15, 15, fill=color, outline="")
    
    def update_datetime(self):
        """Update the date and time display"""
        if hasattr(self, 'datetime_label') and self.datetime_label.winfo_exists():
            now = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
            self.datetime_label.config(text=now)
            # Use after_idle to prevent "invalid command name" errors
            self.root.after(1000, self.update_datetime)
    
    def check_internet_connection(self):
        """Check internet connection status periodically"""
        def internet_check():
            try:
                socket.create_connection(("www.google.com", 80), timeout=5)
                self.internet_connected = True
            except:
                self.internet_connected = False
            
            self.update_internet_indicator()
            # Use after_idle to prevent "invalid command name" errors
            self.root.after(30000, internet_check)  # Check every 30 seconds
        
        internet_check()
    
    def load_users(self):
        """Load user data from file or Firebase"""
        try:
            if self.firebase.initialized:
                firebase_users = self.firebase.get_all_users()
                if firebase_users:
                    self.users = {u['username']: u for u in firebase_users}
                    # Add timestamps to legacy entries
                    for username, data in self.users.items():
                        if 'LAST_MODIFIED' not in data:
                            data['LAST_MODIFIED'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    return
        except Exception as e:
            print(f"Error loading users from Firebase: {e}")
        
        # Fallback to local file
        if os.path.exists("users_data.json"):
            try:
                with open("users_data.json", "r") as f:
                    self.users = json.load(f)
            except json.JSONDecodeError:
                self.users = self.create_default_users()
        else:
            self.users = self.create_default_users()
        
        # Ensure timestamps exist
        for username, data in self.users.items():
            if 'LAST_MODIFIED' not in data:
                data['LAST_MODIFIED'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    def create_default_users(self):
        """Create default user data with hashed passwords and timestamps"""
        now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        return {
            "mej.esam": {
                "password": bcrypt.hashpw("wjap19527".encode('utf-8'), bcrypt.gensalt()).decode('utf-8'),
                "role": "admin",
                "LAST_MODIFIED": now
            },
            "doctor1": {
                "password": bcrypt.hashpw("doc123".encode('utf-8'), bcrypt.gensalt()).decode('utf-8'),
                "role": "editor",
                "LAST_MODIFIED": now
            },
            "nurse1": {
                "password": bcrypt.hashpw("nur123".encode('utf-8'), bcrypt.gensalt()).decode('utf-8'),
                "role": "viewer",
                "LAST_MODIFIED": now
            },
            "pharmacist1": {
                "password": bcrypt.hashpw("pharm123".encode('utf-8'), bcrypt.gensalt()).decode('utf-8'),
                "role": "pharmacist",
                "LAST_MODIFIED": now
            },
            "seraj": {
                "password": bcrypt.hashpw("steve8288".encode('utf-8'), bcrypt.gensalt()).decode('utf-8'),
                "role": "admin",
                "LAST_MODIFIED": now
            }
        }
    
    def save_users_to_file(self):
        """Save user data to file with timestamp updates"""
        for username, data in self.users.items():
            if 'LAST_MODIFIED' not in data:
                data['LAST_MODIFIED'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        with open("users_data.json", "w") as f:
            json.dump(self.users, f, indent=4)
    
    def on_close(self):
        """Handle window close event"""
        if messagebox.askyesno("Exit Confirmation", "Are you sure you want to exit?"):
            # Cancel all pending after events
            for after_id in self.root.tk.eval('after info').split():
                self.root.after_cancel(after_id)
            
            self.executor.shutdown(wait=False)
            self.root.destroy()
    
    def clear_frame(self):
        """Clear all widgets from the main frame except status bar"""
        for widget in self.root.winfo_children():
            if widget not in [self.status_frame]:
                widget.destroy()
    
    def setup_login_screen(self):
        """Setup the login screen with modern design"""
        self.clear_frame()
        
        # Main container with gradient background
        main_frame = tk.Frame(self.root, bg='white')
        main_frame.pack(expand=True, fill=tk.BOTH)
        
        # Left side with logo and app name
        left_frame = tk.Frame(main_frame, bg='#3498db')
        left_frame.pack(side=tk.LEFT, fill=tk.Y, expand=False)
        
        # App logo and name
        logo_frame = tk.Frame(left_frame, bg='#3498db')
        logo_frame.pack(expand=True, fill=tk.BOTH, padx=40, pady=40)
        
        # App name with modern font
        tk.Label(logo_frame, text="OncoCare", font=('Helvetica', 36, 'bold'), 
                bg='#3498db', fg='white').pack(pady=(40, 10))
        
        tk.Label(logo_frame, text="Pediatric Oncology Patient Management System", 
                font=('Helvetica', 14), bg='#3498db', fg='white').pack(pady=(0, 40))
        
        # Right side with login form
        right_frame = tk.Frame(main_frame, bg='white')
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        # Login form container
        form_container = tk.Frame(right_frame, bg='white')
        form_container.place(relx=0.5, rely=0.5, anchor='center')
        
        # Login header
        tk.Label(form_container, text="Login to OncoCare", font=('Helvetica', 24, 'bold'), 
                bg='white').pack(pady=(0, 30))
        
        # Username field
        tk.Label(form_container, text="Username", font=('Helvetica', 12), 
                bg='white', anchor='w').pack(fill=tk.X, pady=(10, 0))
        self.entry_username = ttk.Entry(form_container, font=('Helvetica', 12))
        self.entry_username.pack(fill=tk.X, ipady=5, pady=(0, 20))
        
        # Password field
        tk.Label(form_container, text="Password", font=('Helvetica', 12), 
                bg='white', anchor='w').pack(fill=tk.X)
        self.entry_password = ttk.Entry(form_container, show="*", font=('Helvetica', 12))
        self.entry_password.pack(fill=tk.X, ipady=5, pady=(0, 30))
        
        # Login button
        login_btn = ttk.Button(form_container, text="Login", command=self.login, 
                             style='Blue.TButton')
        login_btn.pack(fill=tk.X, ipady=10, pady=(0, 20))
        
        # Exit button
        exit_btn = ttk.Button(form_container, text="Exit", command=self.root.quit)
        exit_btn.pack(fill=tk.X, ipady=10)
        
        # Focus on username field
        self.entry_username.focus()
        
        # Bind Enter key to login
        self.entry_password.bind('<Return>', lambda event: self.login())
    
    def login(self):
        """Handle user login and show sync button"""
        username = self.entry_username.get()
        password = self.entry_password.get()

        if username in self.users:
            stored_hash = self.users[username]["password"].encode('utf-8')
            if bcrypt.checkpw(password.encode('utf-8'), stored_hash):
                self.current_user = username
                self.user_label.config(text=f"User: {username} ({self.users[username]['role']})")
                
                # Show sync button after successful login
                self.sync_btn.pack(side=tk.RIGHT, padx=10)
                
                messagebox.showinfo("Login Successful", f"Welcome, {username}!")
                self.main_menu()
                return

        messagebox.showerror("Login Failed", "Invalid username or password.")
        self.entry_password.delete(0, tk.END)
    
    def main_menu(self):
        """Display the modern main menu with improved visual design"""
        self.clear_frame()

        # Main container using grid for better layout control
        main_container = ttk.Frame(self.root, style='TFrame')
        main_container.pack(fill=tk.BOTH, expand=True)

        # Left sidebar with improved design
        sidebar = ttk.Frame(main_container, style='Dark.TFrame', width=280)
        sidebar.pack(side=tk.LEFT, fill=tk.Y)
        sidebar.pack_propagate(False)

        # Logo and user info section
        logo_frame = ttk.Frame(sidebar, style='Dark.TFrame')
        logo_frame.pack(pady=40, padx=20, fill=tk.X)

        # Modern logo layout
        ttk.Label(logo_frame, text="🩺 OncoCare", style='Logo.TLabel').pack(pady=(0, 10))
        ttk.Label(logo_frame, text="Pediatric Oncology Management", 
                 style='SubLogo.TLabel').pack(pady=(0, 30))
        
        # User info panel
        user_frame = ttk.Frame(sidebar, style='Dark.TFrame')
        user_frame.pack(padx=20, fill=tk.X)
        ttk.Label(user_frame, text=f"Logged in as:", style='User Info.TLabel').pack(anchor='w')
        ttk.Label(user_frame, text=self.current_user.upper(), 
                 style='Username.TLabel').pack(anchor='w')
        ttk.Label(user_frame, text=f"Role: {self.users[self.current_user]['role'].title()}",
                 style='Role.TLabel').pack(anchor='w', pady=(0, 20))

        # Quick action buttons in sidebar
        quick_actions = ttk.Frame(sidebar, style='Dark.TFrame')
        quick_actions.pack(padx=20, fill=tk.X)
        ttk.Button(quick_actions, text="🔄 Sync Data", command=self.sync_data,
                  style='Sidebar.TButton').pack(fill=tk.X, pady=3)
        ttk.Button(quick_actions, text="❓ Quick Help", command=self.show_help,
                  style='Sidebar.TButton').pack(fill=tk.X, pady=3)

        # Main content area
        content = ttk.Frame(main_container, style='TFrame')
        content.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        # Main button grid using notebook-style navigation
        button_grid = ttk.Frame(content, style='TFrame')
        button_grid.pack(padx=40, pady=40, fill=tk.BOTH, expand=True)

        # Section 1: Patient Management
        patient_frame = ttk.LabelFrame(button_grid, text="PATIENT MANAGEMENT", 
                                      style='Section.TLabelframe')
        patient_frame.grid(row=0, column=0, sticky='nsew', padx=5, pady=5)

        if self.users[self.current_user]["role"] in ["admin", "editor"]:
            btn1 = ttk.Button(patient_frame, text="➕ New Patient", command=self.add_patient,
                             style='Accent.TButton')
            btn1.pack(fill=tk.X, pady=2)

        btn2 = ttk.Button(patient_frame, text="🔍 Search Patients", command=self.search_patient,
                         style='Accent.TButton')
        btn2.pack(fill=tk.X, pady=2)

        # Statistics button with role check
        if self.users[self.current_user]["role"] in ["admin"]:
            ttk.Button(patient_frame, text="📊 Statistics", command=self.show_statistics,
                     style='Accent.TButton').pack(fill=tk.X, pady=2)

        if self.users[self.current_user]["role"] in ["admin"]:
            btn3 = ttk.Button(patient_frame, text="📋 All Patients", command=self.view_all_patients,
                             style='Accent.TButton')
            btn3.pack(fill=tk.X, pady=2)

        # Section 2: Clinical Tools
        clinical_frame = ttk.LabelFrame(button_grid, text="CLINICAL TOOLS", 
                                       style='Section.TLabelframe')
        clinical_frame.grid(row=0, column=1, sticky='nsew', padx=5, pady=5)

        ttk.Button(clinical_frame, text="📈 Lab & EF Docs", command=self.show_lab_ef_window,
                  style='Clinical.TButton').pack(fill=tk.X, pady=2)
        ttk.Button(clinical_frame, text="🧮 Calculators", command=self.show_calculators,
                  style='Clinical.TButton').pack(fill=tk.X, pady=2)
        ttk.Button(clinical_frame, text="⚠️ Extravasation", command=self.show_extravasation_management,
                  style='Clinical.TButton').pack(fill=tk.X, pady=2)
        ttk.Button(clinical_frame, text="📄 F&N Documentation", command=self.handle_fn_documentation,
                  style='Clinical.TButton').pack(fill=tk.X, pady=2)
    
        # Section 3: Chemotherapy
        chemo_frame = ttk.LabelFrame(button_grid, text="CHEMOTHERAPY", 
                                    style='Section.TLabelframe')
        chemo_frame.grid(row=1, column=0, sticky='nsew', padx=5, pady=5)
        
        ttk.Button(chemo_frame, text="📚 Protocols", command=self.handle_protocols,
                  style='Chemo.TButton').pack(fill=tk.X, pady=2)
        ttk.Button(chemo_frame, text="📝 Sheets", command=self.handle_chemo_sheet,
                  style='Chemo.TButton').pack(fill=tk.X, pady=2)
        
        # Add Drug Documentation button
        if self.users[self.current_user]["role"] in ["pharmacist"] or self.current_user == "mej.esam":
            ttk.Button(chemo_frame, text="💊 Drug Docs", command=self.handle_drug_documentation,
                      style='Chemo.TButton').pack(fill=tk.X, pady=2)
        
        if self.users[self.current_user]["role"] in ["admin", "pharmacist"]:
            ttk.Button(chemo_frame, text="💊 Stocks", command=self.show_chemo_stocks,
                      style='Chemo.TButton').pack(fill=tk.X, pady=2)

        # Section 4: Data Management
        data_frame = ttk.LabelFrame(button_grid, text="DATA MANAGEMENT", 
                                    style='Section.TLabelframe')
        data_frame.grid(row=1, column=1, sticky='nsew', padx=5, pady=5)

        if self.users[self.current_user]["role"] in ["admin", "editor"]:
            ttk.Button(data_frame, text="💾 Backup", command=self.backup_data,
                       style='Data.TButton').pack(fill=tk.X, pady=2)
        if self.users[self.current_user]["role"] == "admin":
            ttk.Button(data_frame, text="📤 Export Data", command=self.export_all_data,
                       style='Data.TButton').pack(fill=tk.X, pady=2)
        if self.current_user == "mej.esam":
            ttk.Button(data_frame, text="🔄 Restore", command=self.restore_data,
                       style='Data.TButton').pack(fill=tk.X, pady=2)
            ttk.Button(data_frame, text="📋 Manage Drop-downs", command=self.manage_dropdowns,
                       style='Data.TButton').pack(fill=tk.X, pady=2)

        # Section 5: Administration
        admin_frame = ttk.LabelFrame(button_grid, text="ADMINISTRATION", 
                                    style='Section.TLabelframe')
        admin_frame.grid(row=2, column=0, columnspan=2, sticky='nsew', padx=5, pady=5)
        
        if self.users[self.current_user]["role"] != "viewer":
            ttk.Button(admin_frame, text="🔑 Change Password", command=self.change_password,
                      style='Admin.TButton').pack(side=tk.LEFT, expand=True, fill=tk.X, padx=2)
        if self.users[self.current_user]["role"] in ["admin", "editor"] or self.current_user == "mej.esam":
            ttk.Button(admin_frame, text="👥 Manage Users", command=self.manage_users,
                      style='Admin.TButton').pack(side=tk.LEFT, expand=True, fill=tk.X, padx=2)
        ttk.Button(admin_frame, text="🚪 Logout", command=self.setup_login_screen,
                  style='Logout.TButton').pack(side=tk.RIGHT, expand=True, fill=tk.X, padx=2)
        
        # Configure grid weights
        button_grid.columnconfigure(0, weight=1)
        button_grid.columnconfigure(1, weight=1)
        button_grid.rowconfigure(0, weight=1)
        button_grid.rowconfigure(1, weight=1)
        button_grid.rowconfigure(2, weight=0)

        # Add version info
        version_frame = ttk.Frame(content, style='TFrame')
        version_frame.pack(side=tk.BOTTOM, fill=tk.X)
        ttk.Label(version_frame, text="v2.0 | Developed by Dr. Esam Mejrab", 
                 style='Version.TLabel').pack(pady=10)

        # Add new style configurations
        self.style.configure('Dark.TFrame', background='#2a3b4c')
        self.style.configure('Logo.TLabel', font=('Helvetica', 20, 'bold'), 
                            foreground='#ffffff', background='#2a3b4c')
        self.style.configure('SubLogo.TLabel', font=('Helvetica', 10), 
                            foreground='#bdc3c7', background='#2a3b4c')
        self.style.configure('User Info.TLabel', font=('Helvetica', 9), 
                            foreground='#95a5a6', background='#2a3b4c')
        self.style.configure('Username.TLabel', font=('Helvetica', 12, 'bold'), 
                            foreground='#ecf0f1', background='#2a3b4c')
        self.style.configure('Role.TLabel', font=('Helvetica', 9), 
                            foreground='#7f8c8d', background='#2a3b4c')
        self.style.configure('Section.TLabelframe', font=('Helvetica', 10, 'bold'), 
                            borderwidth=2, relief='solid', labelmargins=10)
        self.style.configure('Section.TLabelframe.Label', foreground='#2c3e50')
        self.style.map('Accent.TButton',
            background=[('active', '#3498db'), ('!active', '#2980b9')],
            foreground=[('active', 'white'), ('!active', 'white')]
        )
        self.style.configure('Accent.TButton', font=('Helvetica', 10), 
                            background='#2980b9', foreground='white', padding=8)
        self.style.configure('Clinical.TButton', background='#27ae60', 
                            foreground='white', padding=8)
        self.style.configure('Chemo.TButton', background='#f39c12', 
                            foreground='white', padding=8)
        self.style.configure('Data.TButton', background='#8e44ad', 
                            foreground='white', padding=8)
        self.style.configure('Admin.TButton', background='#34495e', 
                            foreground='white', padding=8)
        self.style.configure('Logout.TButton', background='#e74c3c', 
                            foreground='white', padding=8)
        self.style.configure('Sidebar.TButton', font=('Helvetica', 10), 
                            background='#34495e', foreground='white', padding=6)
        self.style.configure('Version.TLabel', font=('Helvetica', 8), 
                            foreground='#7f8c8d', background=self.light_color)

    def get_dropdown_options(self):
        """Properly load dropdown values from JSON structure"""
        try:
            with open(DROPDOWN_FILE, 'r') as f:
                dropdown_data = json.load(f)
                return {k: v['values'] for k, v in dropdown_data.get('data', {}).items()}
        except Exception as e:
            print(f"Error loading dropdowns: {e}")
            return DROPDOWN_OPTIONS  # Fallback to defaults
    
    def manage_dropdowns(self):
        """Fixed dropdown manager with legacy data support"""
        if hasattr(self, 'dropdown_win') and self.dropdown_win.winfo_exists():
            self.dropdown_win.lift()
            return

        self.dropdown_win = tk.Toplevel(self.root)
        self.dropdown_win.title("Manage Drop-down Lists")
        self.dropdown_win.geometry("500x400")

        # Main container
        main_frame = ttk.Frame(self.dropdown_win)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Dropdown selection
        ttk.Label(main_frame, text="Select Drop-down List:").pack(pady=5)
        
        self.current_dropdown = tk.StringVar()
        self.dropdown_combo = ttk.Combobox(
            main_frame,
            textvariable=self.current_dropdown,
            state="readonly"
        )
        self.dropdown_combo.pack(fill=tk.X, pady=5)
        
        # Values list
        list_frame = ttk.Frame(main_frame)
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        self.values_listbox = tk.Listbox(
            list_frame,
            selectmode=tk.SINGLE,
            font=('Helvetica', 12),
            height=10
        )
        scrollbar = ttk.Scrollbar(list_frame)
        self.values_listbox.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.values_listbox.yview)
        
        self.values_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Entry field
        self.value_var = tk.StringVar()
        self.value_entry = ttk.Entry(main_frame, textvariable=self.value_var, font=('Helvetica', 12))
        self.value_entry.pack(fill=tk.X, pady=10)

        # Buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X)
        
        ttk.Button(btn_frame, text="Add", command=self._add_value).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Update", command=self._update_value).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Delete", command=self._delete_value).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Close", command=self.dropdown_win.destroy).pack(side=tk.RIGHT)

        # Load data and initialize
        self._load_dropdown_options()
        self.dropdown_combo.bind("<<ComboboxSelected>>", lambda e: self._update_values_list())
        if self.dropdown_combo['values']:
            self.dropdown_combo.current(0)
            self._update_values_list()

    def _load_dropdown_options(self):
        """Load dropdown names with legacy data support"""
        options = self._load_dropdown_data()
        dropdown_names = []
        
        for name, values in options.items():
            # Convert legacy list format to new dictionary format
            if isinstance(values, list):
                options[name] = {
                    'values': values,
                    'LAST_MODIFIED': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                }
                self._save_dropdown_data(options)
            
            dropdown_names.append(name)
        
        self.dropdown_combo['values'] = sorted(dropdown_names)

    def _update_values_list(self):
        """Update values listbox with proper data structure"""
        self.values_listbox.delete(0, tk.END)
        selected = self.current_dropdown.get()
        options = self._load_dropdown_data()
        
        if selected in options:
            # Handle both old list format and new dictionary format
            values = options[selected]['values'] if isinstance(options[selected], dict) else options[selected]
            for value in values:
                self.values_listbox.insert(tk.END, value)

    def _add_value(self):
        """Add new value to current dropdown"""
        dropdown = self.current_dropdown.get()
        new_value = self.value_var.get().strip()
        options = self._load_dropdown_data()

        if not dropdown:
            messagebox.showwarning("Warning", "Select a dropdown list first")
            return
        if not new_value:
            messagebox.showwarning("Warning", "Enter a value")
            return

        # Initialize dropdown if not exists
        if dropdown not in options:
            options[dropdown] = {
                'values': [],
                'LAST_MODIFIED': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }

        # Check if value exists
        existing_values = [v.lower() for v in options[dropdown]['values']]
        if new_value.lower() in existing_values:
            messagebox.showwarning("Warning", "Value already exists")
            return

        options[dropdown]['values'].append(new_value)
        options[dropdown]['LAST_MODIFIED'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        self._save_dropdown_data(options)
        self._update_values_list()
        self.value_var.set("")

    def _update_value(self):
        """Update existing value"""
        dropdown = self.current_dropdown.get()
        new_value = self.value_var.get().strip()
        options = self._load_dropdown_data()

        try:
            index = self.values_listbox.curselection()[0]
            old_value = self.values_listbox.get(index)
        except IndexError:
            messagebox.showwarning("Warning", "Select a value to edit")
            return

        if new_value == old_value:
            return

        # Check if new value already exists
        existing_values = [v.lower() for v in options[dropdown]['values'] if v != old_value]
        if new_value.lower() in existing_values:
            messagebox.showwarning("Warning", "Value already exists")
            return

        options[dropdown]['values'][index] = new_value
        options[dropdown]['LAST_MODIFIED'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        self._save_dropdown_data(options)
        self._update_values_list()

    def _delete_value(self):
        """Delete selected value"""
        dropdown = self.current_dropdown.get()
        options = self._load_dropdown_data()

        try:
            index = self.values_listbox.curselection()[0]
        except IndexError:
            messagebox.showwarning("Warning", "Select a value to delete")
            return

        if messagebox.askyesno("Confirm", "Delete this value?"):
            del options[dropdown]['values'][index]
            options[dropdown]['LAST_MODIFIED'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            self._save_dropdown_data(options)
            self._update_values_list()

    def _populate_entry_from_selection(self):
        """Fill entry field with selected value"""
        try:
            index = self.values_listbox.curselection()[0]
            self.value_var.set(self.values_listbox.get(index))
            self.value_entry.focus()
        except IndexError:
            pass

    def _load_dropdown_data(self):
        """Load dropdown data with version handling"""
        try:
            with open(DROPDOWN_FILE, 'r') as f:
                data = json.load(f)
                # Handle both old and new formats
                if 'data' in data:
                    return data['data']
                return data
        except (FileNotFoundError, json.JSONDecodeError):
            return {"GENDER": {"values": ["MALE", "FEMALE"], "LAST_MODIFIED": datetime.now().strftime('%Y-%m-%d %H:%M:%S')}}
        
    def _save_dropdown_data(self, data=None):
        """Save dropdown data with versioning"""
        if data is None:
            data = self.dropdown_data
            
        # Ensure proper format
        save_data = {'data': data} if not isinstance(data, dict) or 'data' not in data else data
        
        # Add/modify timestamps
        for name, dropdown in save_data['data'].items():
            if 'LAST_MODIFIED' not in dropdown:
                dropdown['LAST_MODIFIED'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Atomic write
        temp_file = os.path.join(os.path.dirname(__file__), DROPDOWN_FILE + '.tmp')
        with open(temp_file, 'w') as f:
            json.dump(save_data, f, indent=2)
        
        os.replace(temp_file, DROPDOWN_FILE)
        
    def start_sync(self):
        """Initiate synchronization process"""
        if self.sync_in_progress:
            messagebox.showinfo("Sync Running", "Synchronization is already in progress")
            return

        if not self.internet_connected:
            messagebox.showerror("No Internet", "Cannot sync without internet connection")
            return

        self.create_sync_window()
        self.sync_in_progress = True

    def open_protocols_folder(self):
        """Open the Protocols folder"""
        try:
            protocols_path = get_resource_path("Protocols")
            if os.path.exists(protocols_path):
                os.startfile(protocols_path)
            else:
                messagebox.showerror("Error", f"Protocols folder not found at: {protocols_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Could not open Protocols folder: {e}")

    def setup_protocols_config(self):
        """Setup configuration for Protocols folder"""
        self.protocols_config = ConfigParser()
        self.protocols_config_file = os.path.join(os.environ['APPDATA'], 'protocols_config.ini')
        
        # Get proper default path for EXE or dev environment
        try:
            default_path = get_resource_path("Protocols")
            if not os.path.exists(default_path):
                default_path = os.path.join(os.path.expanduser('~'), 'Documents', 'Protocols')
        except Exception:
            default_path = os.path.join(os.path.expanduser('~'), 'Documents', 'Protocols')
        
        if not os.path.exists(self.protocols_config_file):
            self.protocols_config['PROTOCOLS'] = {'path': default_path}
            with open(self.protocols_config_file, 'w', encoding='utf-8') as configfile:
                self.protocols_config.write(configfile)
            os.makedirs(default_path, exist_ok=True)
        else:
            self.protocols_config.read(self.protocols_config_file)
            current_path = self.protocols_config.get('PROTOCOLS', 'path', fallback=default_path)
            os.makedirs(current_path, exist_ok=True)

    def get_protocols_path(self):
        """Get the path to Protocols folder"""
        self.protocols_config.read(self.protocols_config_file)
        return self.protocols_config['PROTOCOLS']['path']

    def set_protocols_path(self, new_path):
        """Set new path for Protocols folder"""
        self.protocols_config['PROTOCOLS']['path'] = new_path
        with open(self.protocols_config_file, 'w') as configfile:
            self.protocols_config.write(configfile)

    def setup_chemo_sheets_config(self):
        """Setup configuration for Chemo Sheets file"""
        self.chemo_sheets_config = ConfigParser()
        self.chemo_sheets_config_file = os.path.join(os.environ['APPDATA'], 'chemo_sheets_config.ini')
        
        # Get proper default path for EXE or dev environment
        try:
            default_path = get_resource_path("FULL CHEMO SHEET.xlsm")
            if not os.path.exists(default_path):
                default_path = os.path.join(os.path.expanduser('~'), 'Documents', 'FULL CHEMO SHEET.xlsm')
        except Exception:
            default_path = os.path.join(os.path.expanduser('~'), 'Documents', 'FULL CHEMO SHEET.xlsm')
        
        if not os.path.exists(self.chemo_sheets_config_file):
            self.chemo_sheets_config['CHEMO_SHEETS'] = {'path': default_path}
            with open(self.chemo_sheets_config_file, 'w', encoding='utf-8') as configfile:
                self.chemo_sheets_config.write(configfile)
        else:
            self.chemo_sheets_config.read(self.chemo_sheets_config_file)

    def get_chemo_sheet_path(self):
        """Get the path to Chemo Sheets file"""
        self.chemo_sheets_config.read(self.chemo_sheets_config_file)
        return self.chemo_sheets_config['CHEMO_SHEETS']['path']

    def set_chemo_sheet_path(self, new_path):
        """Set new path for Chemo Sheets file"""
        self.chemo_sheets_config['CHEMO_SHEETS']['path'] = new_path
        with open(self.chemo_sheets_config_file, 'w') as configfile:
            self.chemo_sheets_config.write(configfile)

    def handle_protocols(self):
        """Handle Protocols button click with improved UI similar to F&N docs"""
        try:
            folder_path = self.get_protocols_path()
            
            dialog = tk.Toplevel(self.root)
            dialog.title("Protocols")
            dialog.geometry("400x200")
            
            msg = f"Current Protocols folder:\n{folder_path}"
            ttk.Label(dialog, text=msg, wraplength=380).pack(pady=20)
            
            btn_frame = ttk.Frame(dialog)
            btn_frame.pack(pady=10)
            
            ttk.Button(btn_frame, text="Open Folder",
                      command=lambda: self.open_folder(folder_path, dialog),
                     style='Green.TButton').pack(side=tk.LEFT, padx=10)
            
            if self.current_user.lower() == "mej.esam":
                ttk.Button(btn_frame, text="Change Path",
                          command=lambda: self.change_protocols_path(dialog),
                         style='Blue.TButton').pack(side=tk.LEFT, padx=10)
            
            ttk.Button(btn_frame, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT, padx=10)
        except Exception as e:
            messagebox.showerror("Error", f"Protocols Error: {str(e)}")

    def open_protocols_folder(self, dialog=None):
        """Open the Protocols folder"""
        try:
            folder_path = self.get_protocols_path()
            if sys.platform == 'win32':
                os.startfile(folder_path)
            else:
                subprocess.run(['open', folder_path] if sys.platform == 'darwin' 
                              else ['xdg-open', folder_path], check=True)
            if dialog:
                dialog.destroy()
        except Exception as e:
            messagebox.showerror("Error", f"Could not open Protocols folder: {str(e)}")

    def change_protocols_path(self, dialog):
        """Change path with comprehensive error handling"""
        new_path = filedialog.askdirectory(title="Select Protocols Folder")
        if new_path:
            try:
                # Normalize and validate path
                new_path = os.path.normpath(new_path)
                if len(new_path) < 2 or not isinstance(new_path, str):
                    raise ValueError("Invalid path format")
                
                # Update config with proper encoding
                self.protocols_config.set('PROTOCOLS', 'path', new_path)
                with open(self.protocols_config_file, 'w', encoding='utf-8') as configfile:
                    self.protocols_config.write(configfile)
                
                # Create directory with safe permissions
                os.makedirs(new_path, exist_ok=True, mode=0o755)
                
                messagebox.showinfo("Success", f"Path updated to:\n{new_path}")
                dialog.destroy()
                
            except Exception as e:
                messagebox.showerror("Error",
                    f"Failed to update path: {str(e)}\n"
                    "Please choose a different location with:\n"
                    "- Standard Latin characters\n"
                    "- No special symbols\n"
                    "- Short path length")
                
    def handle_chemo_sheet(self):
        """Handle Chemo Sheet button click with improved UI similar to F&N docs"""
        try:
            file_path = self.get_chemo_sheet_path()
            
            dialog = tk.Toplevel(self.root)
            dialog.title("Chemo Sheets")
            dialog.geometry("400x200")
            
            msg = f"Current Chemo Sheet file:\n{file_path}"
            ttk.Label(dialog, text=msg, wraplength=380).pack(pady=20)
            
            btn_frame = ttk.Frame(dialog)
            btn_frame.pack(pady=10)
            
            ttk.Button(btn_frame, text="Open File",
                      command=lambda: self.open_file(file_path, dialog),
                     style='Green.TButton').pack(side=tk.LEFT, padx=10)
            
            if self.current_user.lower() == "mej.esam":
                ttk.Button(btn_frame, text="Change Path",
                          command=lambda: self.change_chemo_sheet_path(dialog),
                         style='Blue.TButton').pack(side=tk.LEFT, padx=10)
            
            ttk.Button(btn_frame, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT, padx=10)
        except Exception as e:
            messagebox.showerror("Error", f"Chemo Sheets Error: {str(e)}")

    def open_file(self, path, dialog=None):
        """Generic file opening function"""
        try:
            if sys.platform == 'win32':
                os.startfile(path)
            else:
                subprocess.run(['open', path] if sys.platform == 'darwin'
                              else ['xdg-open', path], check=True)
            if dialog:
                dialog.destroy()
        except Exception as e:
            messagebox.showerror("Error", f"Could not open file: {str(e)}")

    def open_chemo_sheet(self, dialog=None):
        """Open the Chemo Sheet file"""
        try:
            file_path = self.get_chemo_sheet_path()
            if sys.platform == 'win32':
                os.startfile(file_path)
            else:
                subprocess.run(['open', file_path] if sys.platform == 'darwin' 
                              else ['xdg-open', file_path], check=True)
            if dialog:
                dialog.destroy()
        except Exception as e:
            messagebox.showerror("Error", f"Could not open Chemo Sheet: {str(e)}")

    def change_chemo_sheet_path(self, dialog):
        """Change path with comprehensive error handling"""
        new_path = filedialog.askopenfilename(
            title="Select Chemo Sheet File",
            filetypes=[("Excel Files", "*.xlsm"), ("Excel Files", "*.xlsx"), ("All Files", "*.*")]
        )
        if new_path:
            try:
                # Normalize and validate path
                new_path = os.path.normpath(new_path)
                if len(new_path) < 2 or not isinstance(new_path, str):
                    raise ValueError("Invalid path format")
                
                # Update config with proper encoding
                self.chemo_sheets_config.set('CHEMO_SHEETS', 'path', new_path)
                with open(self.chemo_sheets_config_file, 'w', encoding='utf-8') as configfile:
                    self.chemo_sheets_config.write(configfile)
                
                messagebox.showinfo("Success", f"Path updated to:\n{new_path}")
                dialog.destroy()
                
            except Exception as e:
                messagebox.showerror("Error",
                    f"Failed to update path: {str(e)}\n"
                    "Please choose a different file with:\n"
                    "- Standard Latin characters\n"
                    "- No special symbols\n"
                    "- Short path length")
                
    def show_chemo_stocks(self):
        """Open the Chemo Stocks Google Sheet"""
        try:
            webbrowser.open(CHEMO_STOCKS_URL)
        except Exception as e:
            messagebox.showerror("Error", f"Could not open Chemo Stocks: {e}")        
    
    def show_lab_ef_window(self, patient_data=None):
        """Show lab window - with patient selection if needed"""
        # Close existing window if open
        if hasattr(self, 'lab_ef_window') and self.lab_ef_window.winfo_exists():
            self.lab_ef_window.destroy()
            
        self.lab_ef_window = tk.Toplevel(self.root)
        self.lab_ef_window.title("Lab Investigations & EF Documentation")
        
        if patient_data:
            # Direct mode - open full lab interface
            self.lab_ef_window.attributes('-fullscreen', True)
            self.current_lab_patient = patient_data
            self.current_file_number = patient_data.get("FILE NUMBER", "")
            self.setup_lab_interface()
        else:
            # Selection mode - show patient selection
            self.lab_ef_window.geometry("1100x650")
            self.main_container = ttk.Frame(self.lab_ef_window)
            self.main_container.pack(fill=tk.BOTH, expand=True)
            self.show_patient_selection_screen()

    def setup_lab_interface(self):
        """Setup lab interface components with history controls"""
        # Clear existing widgets
        for widget in self.lab_ef_window.winfo_children():
            widget.destroy()
        
        # Main container
        main_frame = ttk.Frame(self.lab_ef_window)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Patient info header
        patient_frame = ttk.Frame(main_frame)
        patient_frame.pack(fill=tk.X, padx=10, pady=10)
        
        info_text = f"Patient: {self.current_lab_patient.get('NAME', '')} (File#: {self.current_file_number})"
        ttk.Label(patient_frame, text=info_text, font=('Helvetica', 12, 'bold')).pack(side=tk.LEFT)
        
        # History controls
        hist_frame = ttk.Frame(main_frame)
        hist_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.hist_data_label = ttk.Label(hist_frame, text="Viewing: New Entry", font=('Helvetica', 10))
        self.hist_data_label.pack(side=tk.LEFT)
        
        self.hist_prev_btn = ttk.Button(hist_frame, text="◀ Previous", state='disabled',
                                      command=lambda: self.navigate_lab_ef_history(-1))
        self.hist_prev_btn.pack(side=tk.LEFT, padx=5)
        
        self.hist_next_btn = ttk.Button(hist_frame, text="Next ▶", state='disabled',
                                      command=lambda: self.navigate_lab_ef_history(1))
        self.hist_next_btn.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(hist_frame, text="New Entry", command=self.create_new_lab_ef_entry).pack(side=tk.RIGHT)
        
        # Notebook with tabs
        notebook = ttk.Notebook(main_frame)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Lab tab
        lab_tab = ttk.Frame(notebook)
        notebook.add(lab_tab, text="Lab Investigations")
        self.setup_lab_tab(lab_tab)
        
        # EF tab
        ef_tab = ttk.Frame(notebook)
        notebook.add(ef_tab, text="Ejection Fraction")
        self.setup_ef_tab(ef_tab)
        
        # Control buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(btn_frame, text="Save", command=self.save_lab_ef_data).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="Close", command=self.lab_ef_window.destroy).pack(side=tk.RIGHT)
        
        # Load existing lab data for this patient
        self.load_lab_ef_history_controls()
        
    def show_patient_selection_screen(self):
        """First screen - patient selection"""
        self.selection_frame = ttk.Frame(self.main_container)
        self.selection_frame.pack(expand=True, fill=tk.BOTH, padx=20, pady=20)
        
        # Header
        ttk.Label(self.selection_frame, 
                 text="Select Patient First", 
                 font=('Helvetica', 16, 'bold')).pack(pady=20)
        
        # Search frame
        search_frame = ttk.Frame(self.selection_frame)
        search_frame.pack(pady=10)
        
        ttk.Label(search_frame, text="Search:").pack(side=tk.LEFT)
        self.search_entry = ttk.Entry(search_frame, width=40)
        self.search_entry.pack(side=tk.LEFT, padx=5)
        self.search_entry.bind('<KeyRelease>', lambda e: self.search_patients())
        
        # Patient list
        self.tree_frame = ttk.Frame(self.selection_frame)
        self.tree_frame.pack(expand=True, fill=tk.BOTH)
        
        columns = ("file_number", "name", "age", "gender")
        self.patient_tree = ttk.Treeview(self.tree_frame, columns=columns, show="headings", height=15)
        
        # Configure columns
        col_widths = {"file_number": 100, "name": 200, "age": 80, "gender": 80}
        for col in columns:
            self.patient_tree.heading(col, text=col.replace("_", " ").title())
            self.patient_tree.column(col, width=col_widths[col], anchor='center')
        
        # Double-click binding
        self.patient_tree.bind('<Double-1>', self.on_patient_selected)
        
        scrollbar = ttk.Scrollbar(self.tree_frame, orient="vertical", command=self.patient_tree.yview)
        self.patient_tree.configure(yscrollcommand=scrollbar.set)
        
        self.patient_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Initial population
        self.search_patients()

    def show_main_content(self, patient_data):
        """Show the main lab/EF interface"""
        # Switch to fullscreen after patient selection
        self.lab_ef_window.attributes('-fullscreen', True)
        
        # Destroy previous content
        for widget in self.main_container.winfo_children():
            widget.destroy()
        
        # Store patient reference
        self.current_lab_patient = patient_data
        self.current_file_number = patient_data.get("FILE NUMBER", "")
        
        # Rebuild main interface
        self.setup_lab_ef_content()
        self.load_lab_ef_history_controls()
        
        # Set patient info
        self.patient_info_label.config(
            text=f"Patient: {patient_data.get('NAME', '')} (File#: {self.current_file_number})"
        )

    def setup_lab_ef_content(self):
        """Setup all the lab/EF content in the container"""
        self.current_lab_data_index = -1
        
        # Patient info frame
        patient_frame = ttk.Frame(self.main_container, style='TFrame')
        patient_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.patient_info_label = ttk.Label(patient_frame, text="", font=('Helvetica', 12, 'bold'))
        self.patient_info_label.pack(side=tk.LEFT)
        
        # Historical data controls
        hist_frame = ttk.Frame(self.main_container)
        hist_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.hist_data_label = ttk.Label(hist_frame, text="Viewing: New Entry", font=('Helvetica', 10))
        self.hist_data_label.pack(side=tk.LEFT)
        
        self.hist_prev_btn = ttk.Button(hist_frame, text="◀ Previous", state='disabled',
                                       command=lambda: self.navigate_lab_ef_history(-1))
        self.hist_prev_btn.pack(side=tk.LEFT, padx=5)
        
        self.hist_next_btn = ttk.Button(hist_frame, text="Next ▶", state='disabled',
                                       command=lambda: self.navigate_lab_ef_history(1))
        self.hist_next_btn.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(hist_frame, text="New Entry", command=self.create_new_lab_ef_entry).pack(side=tk.RIGHT)
        
        # Notebook for tabs
        notebook = ttk.Notebook(self.main_container)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Lab Investigations Tab
        self.lab_frame = ttk.Frame(notebook)
        notebook.add(self.lab_frame, text="Lab Investigations")
        self.setup_lab_tab(self.lab_frame)
        
        # EF Documentation Tab
        self.ef_frame = ttk.Frame(notebook)
        notebook.add(self.ef_frame, text="Ejection Fraction")
        self.setup_ef_tab(self.ef_frame)
        
        # Button frame
        btn_frame = ttk.Frame(self.main_container, style='TFrame')
        btn_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(btn_frame, text="Save", command=self.save_lab_ef_data,
                  style='Blue.TButton').pack(side=tk.LEFT, padx=5)
        
        ttk.Button(btn_frame, text="Print", command=self.print_lab_ef_report,
                  style='Green.TButton').pack(side=tk.LEFT, padx=5)
        
        ttk.Button(btn_frame, text="Close", command=self.lab_ef_window.destroy).pack(side=tk.RIGHT, padx=5)

    def search_patients(self, search_term=None):
        """Search and populate patient list"""
        search_term = search_term or self.search_entry.get()
        self.patient_tree.delete(*self.patient_tree.get_children())
        
        for patient in self.patient_data:
            if (search_term.lower() in patient.get("NAME", "").lower() or 
                search_term.lower() in patient.get("FILE NUMBER", "").lower()):
                self.patient_tree.insert("", "end", values=(
                    patient.get("FILE NUMBER", ""),
                    patient.get("NAME", ""),
                    patient.get("AGE", ""),
                    patient.get("GENDER", "")
                ))

    def on_patient_selected(self, event):
        """Handle double-click patient selection"""
        selected = self.patient_tree.focus()
        if not selected:
            return
        
        values = self.patient_tree.item(selected, "values")
        file_number = values[0]
        
        for patient in self.patient_data:
            if patient.get("FILE NUMBER") == file_number:
                self.selection_frame.pack_forget()
                self.show_main_content(patient)
                break

    def load_lab_ef_history_controls(self):
        """Enable and setup historical data navigation controls"""
        if not hasattr(self, 'current_lab_patient') or not self.current_lab_patient:
            return
        
        lab_count = len(self.current_lab_patient.get("lab_results", []))
        ef_count = len(self.current_lab_patient.get("ef_data", []))
        total_entries = max(lab_count, ef_count)
        
        if total_entries > 0:
            self.hist_prev_btn.config(state='normal')
            self.hist_next_btn.config(state='normal')
            self.current_lab_data_index = total_entries - 1  # Start with most recent
            self.display_lab_ef_data(self.current_lab_data_index)
        else:
            self.hist_prev_btn.config(state='disabled')
            self.hist_next_btn.config(state='disabled')
            self.current_lab_data_index = -1
            self.hist_data_label.config(text="Viewing: New Entry")

    def navigate_lab_ef_history(self, direction):
        """Navigate through historical lab/EF data"""
        if not hasattr(self, 'current_lab_patient') or not self.current_lab_patient:
            return
        
        lab_count = len(self.current_lab_patient.get("lab_results", []))
        ef_count = len(self.current_lab_patient.get("ef_data", []))
        total_entries = max(lab_count, ef_count)
        
        new_index = self.current_lab_data_index + direction
        
        if 0 <= new_index < total_entries:
            self.current_lab_data_index = new_index
            self.display_lab_ef_data(self.current_lab_data_index)
        elif new_index == -1:  # Show new entry
            self.create_new_lab_ef_entry()
        else:
            return  # Don't go beyond available entries
        
        # Update button states
        self.hist_prev_btn.config(state='normal' if new_index > 0 else 'disabled')
        self.hist_next_btn.config(state='normal' if new_index < total_entries - 1 else 'disabled')

    def display_lab_ef_data(self, index):
        """Display historical lab and EF data at the given index"""
        if not hasattr(self, 'current_lab_patient') or not self.current_lab_patient:
            return
        
        lab_data = self.current_lab_patient.get("lab_results", [])
        ef_data = self.current_lab_patient.get("ef_data", [])
        
        # Clear current entries
        self.clear_lab_ef_entries()
        
        # Load lab data if available
        if index < len(lab_data):
            data = lab_data[index]
            self.lab_date_entry.insert(0, data.get("date", ""))
            
            values = data.get("values", {})
            for test, entry in self.lab_entries.items():
                if test in values:
                    entry.insert(0, values[test])
                    # Trigger critical value check
                    entry.event_generate("<FocusOut>")
        
        # Load EF data if available
        if index < len(ef_data):
            data = ef_data[index]
            baseline = data.get("baseline", {})
            self.ef_baseline_date.insert(0, baseline.get("date", ""))
            self.ef_baseline_value.insert(0, baseline.get("value", ""))
            
            serial = data.get("serial", [])
            for i, entry in enumerate(self.ef_serial_entries):
                if i < len(serial):
                    entry["date"].insert(0, serial[i].get("date", ""))
                    entry["value"].insert(0, serial[i].get("value", ""))
                    entry["change_label"].config(text=serial[i].get("change", ""))
            
            self.ef_notes.delete("1.0", tk.END)
            self.ef_notes.insert("1.0", data.get("notes", ""))
        
        # Update status label
        date_str = ""
        if index < len(lab_data) and "date" in lab_data[index]:
            date_str = lab_data[index]["date"]
        elif index < len(ef_data) and "baseline" in ef_data[index] and "date" in ef_data[index]["baseline"]:
            date_str = ef_data[index]["baseline"]["date"]
        
        self.hist_data_label.config(text=f"Viewing: Entry from {date_str if date_str else 'Unknown date'}")
        self.current_lab_data_index = index

    def clear_lab_ef_entries(self):
        """Clear all lab and EF entries"""
        self.lab_date_entry.delete(0, tk.END)
        for entry in self.lab_entries.values():
            entry.delete(0, tk.END)
        
        self.ef_baseline_date.delete(0, tk.END)
        self.ef_baseline_value.delete(0, tk.END)
        for entry in self.ef_serial_entries:
            entry["date"].delete(0, tk.END)
            entry["value"].delete(0, tk.END)
            entry["change_label"].config(text="")
        self.ef_notes.delete("1.0", tk.END)

    def create_new_lab_ef_entry(self):
        """Prepare the window for a new lab/EF entry"""
        self.clear_lab_ef_entries()
        self.current_lab_data_index = -1
        self.hist_data_label.config(text="Viewing: New Entry")
        
        # Set today's date as default
        today = datetime.now().strftime("%d/%m/%Y")
        self.lab_date_entry.insert(0, today)
        self.ef_baseline_date.insert(0, today)
        
        # Update button states
        lab_count = len(self.current_lab_patient.get("lab_results", [])) if hasattr(self, 'current_lab_patient') else 0
        ef_count = len(self.current_lab_patient.get("ef_data", [])) if hasattr(self, 'current_lab_patient') else 0
        total_entries = max(lab_count, ef_count)
        
        self.hist_prev_btn.config(state='normal' if total_entries > 0 else 'disabled')
        self.hist_next_btn.config(state='disabled')
        
    def setup_lab_tab(self, parent):
        """Setup the lab investigations tab"""
        # Create scrollable frame
        canvas = tk.Canvas(parent, highlightthickness=0)
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
        
        # Add mouse wheel scrolling
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        canvas.bind_all("<MouseWheel>", on_mousewheel)
        scrollable_frame.bind("<MouseWheel>", on_mousewheel)
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        
        # Lab entries dictionary
        self.lab_entries = {}
        self.lab_crit_labels = {}  # To store critical value labels
        
        # Date entry
        date_frame = ttk.Frame(scrollable_frame)
        date_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(date_frame, text="Date:").pack(side=tk.LEFT, padx=5)
        self.lab_date_entry = ttk.Entry(date_frame, width=15)
        self.lab_date_entry.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(date_frame, text="📅", width=3,
                  command=lambda: self.show_calendar(self.lab_date_entry)).pack(side=tk.LEFT, padx=5)
        
        # Add lab tests
        for test, ranges in LAB_RANGES.items():
            frame = ttk.Frame(scrollable_frame)
            frame.pack(fill=tk.X, padx=10, pady=5)
            
            # Test name and normal range
            info_label = f"{test} ({ranges['normal_range']})"
            ttk.Label(frame, text=info_label, width=30, anchor="w").pack(side=tk.LEFT, padx=5)
            
            # Value entry
            entry = ttk.Entry(frame, width=10)
            entry.pack(side=tk.LEFT, padx=5)
            self.lab_entries[test] = entry
            
            # Critical value indicator
            crit_label = ttk.Label(frame, text="", width=20)
            crit_label.pack(side=tk.LEFT, padx=5)
            self.lab_crit_labels[test] = crit_label
            
            # Bind validation
            entry.bind("<FocusOut>", lambda e, t=test: self.check_lab_critical_value(t))

    def setup_ef_tab(self, parent):
        """Setup the ejection fraction documentation tab"""
        # Create scrollable frame
        canvas = tk.Canvas(parent, highlightthickness=0)
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        
        # EF documentation fields
        self.ef_entries = {}
        
        # Baseline EF
        baseline_frame = ttk.LabelFrame(scrollable_frame, text="Baseline EF (Before Anthracyclines)")
        baseline_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(baseline_frame, text="Date:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.ef_baseline_date = ttk.Entry(baseline_frame, width=15)
        self.ef_baseline_date.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        ttk.Button(baseline_frame, text="📅", width=3,
                  command=lambda: self.show_calendar(self.ef_baseline_date)).grid(row=0, column=2, padx=5)
        
        ttk.Label(baseline_frame, text="EF (%):").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.ef_baseline_value = ttk.Entry(baseline_frame, width=10)
        self.ef_baseline_value.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        
        # Serial EF measurements
        serial_frame = ttk.LabelFrame(scrollable_frame, text="Serial EF Measurements")
        serial_frame.pack(fill=tk.X, padx=10, pady=10)
        
        # Table headers
        ttk.Label(serial_frame, text="Date", font=('Helvetica', 10, 'bold')).grid(row=0, column=0, padx=5, pady=5)
        ttk.Label(serial_frame, text="EF (%)", font=('Helvetica', 10, 'bold')).grid(row=0, column=1, padx=5, pady=5)
        ttk.Label(serial_frame, text="Change from Baseline", font=('Helvetica', 10, 'bold')).grid(row=0, column=2, padx=5, pady=5)
        
        # Add 5 measurement rows
        self.ef_serial_entries = []
        for i in range(1, 6):
            date_entry = ttk.Entry(serial_frame, width=15)
            date_entry.grid(row=i, column=0, padx=5, pady=5)
            
            ef_entry = ttk.Entry(serial_frame, width=10)
            ef_entry.grid(row=i, column=1, padx=5, pady=5)
            
            change_label = ttk.Label(serial_frame, text="", width=20)
            change_label.grid(row=i, column=2, padx=5, pady=5)
            
            self.ef_serial_entries.append({
                "date": date_entry,
                "value": ef_entry,
                "change_label": change_label
            })
            
            # Add calendar buttons
            ttk.Button(serial_frame, text="📅", width=3,
                      command=lambda e=date_entry: self.show_calendar(e)).grid(row=i, column=3, padx=5)
        
        # Add button to calculate changes
        ttk.Button(serial_frame, text="Calculate Changes", 
                  command=self.calculate_ef_changes).grid(row=6, column=0, columnspan=4, pady=10)
        
        # EF monitoring notes
        notes_frame = ttk.LabelFrame(scrollable_frame, text="Monitoring Notes")
        notes_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.ef_notes = tk.Text(notes_frame, height=5, wrap=tk.WORD)
        self.ef_notes.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Add predefined alerts
        alert_frame = ttk.Frame(notes_frame)
        alert_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(alert_frame, text="Quick Alerts:").pack(side=tk.LEFT, padx=5)
        
        for alert in EF_NOTES_TEMPLATES:
            ttk.Button(alert_frame, text=alert, width=20,
                      command=lambda a=alert: self.add_ef_alert(a)).pack(side=tk.LEFT, padx=2)

    def check_lab_critical_value(self, test):
        """Check if lab value is critical and update label"""
        try:
            value = float(self.lab_entries[test].get())
            ranges = LAB_RANGES[test]
            label = self.lab_crit_labels[test]
            
            # Parse normal range (simplified parsing)
            normal_range = ranges["normal_range"]
            if "-" in normal_range:
                low, high = map(float, normal_range.split("-")[:2])
            elif "<" in normal_range:
                low, high = 0, float(normal_range[1:])
            else:
                low, high = 0, float('inf')
            
            # Check critical values
            if ranges["critical_low"] and value < float(ranges["critical_low"][1:]):
                label.config(text="CRITICALLY LOW!", foreground="red")
            elif ranges["critical_high"] and value > float(ranges["critical_high"][1:]):
                label.config(text="CRITICALLY HIGH!", foreground="red")
            elif value < low or value > high:
                label.config(text="Abnormal", foreground="orange")
            else:
                label.config(text="Normal", foreground="green")
                
        except ValueError:
            self.lab_crit_labels[test].config(text="", foreground="black")

    def calculate_ef_changes(self):
        """Calculate EF changes from baseline"""
        try:
            baseline = float(self.ef_baseline_value.get())
            
            for entry in self.ef_serial_entries:
                if entry["value"].get():
                    current = float(entry["value"].get())
                    change = current - baseline
                    percent_change = (change / baseline) * 100
                    
                    alert = ""
                    if current < 50:
                        alert = "EF < 50%!"
                    elif percent_change < -10:
                        alert = ">10% drop!"
                    
                    entry["change_label"].config(
                        text=f"{change:.1f} ({percent_change:.1f}%) {alert}",
                        foreground="red" if alert else "black"
                    )
                    
        except ValueError:
            messagebox.showerror("Error", "Please enter valid numeric values")

    def add_ef_alert(self, alert_text):
        """Add a predefined alert to the notes"""
        self.ef_notes.insert(tk.END, alert_text + "\n")

    def save_lab_ef_data(self):
        """Save lab and EF data to patient record"""
        if not hasattr(self, 'current_file_number') or not self.current_file_number:
            messagebox.showwarning("Warning", "Please select a patient first")
            return
        
        # Find patient in data
        patient_found = False
        for patient in self.patient_data:
            if patient.get("FILE NUMBER") == self.current_file_number:
                patient_found = True
                
                # Prepare lab data
                lab_data = {
                    "date": self.lab_date_entry.get(),
                    "values": {test: entry.get() for test, entry in self.lab_entries.items()}
                }
                
                # Prepare EF data
                ef_data = {
                    "baseline": {
                        "date": self.ef_baseline_date.get(),
                        "value": self.ef_baseline_value.get()
                    },
                    "serial": [
                        {
                            "date": entry["date"].get(),
                            "value": entry["value"].get(),
                            "change": entry["change_label"].cget("text")
                        }
                        for entry in self.ef_serial_entries
                    ],
                    "notes": self.ef_notes.get("1.0", tk.END)
                }
                
                # Add to patient record
                if "lab_results" not in patient:
                    patient["lab_results"] = []
                if "ef_data" not in patient:
                    patient["ef_data"] = []
                
                if self.current_lab_data_index == -1:  # New entry
                    patient["lab_results"].append(lab_data)
                    patient["ef_data"].append(ef_data)
                else:  # Update existing entry
                    if self.current_lab_data_index < len(patient["lab_results"]):
                        patient["lab_results"][self.current_lab_data_index] = lab_data
                    else:
                        patient["lab_results"].append(lab_data)
                    
                    if self.current_lab_data_index < len(patient["ef_data"]):
                        patient["ef_data"][self.current_lab_data_index] = ef_data
                    else:
                        patient["ef_data"].append(ef_data)
                
                # Update last modified info
                patient["LAST_MODIFIED_BY"] = self.current_user
                patient["LAST_MODIFIED_DATE"] = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
                
                # Update the history controls
                self.load_lab_ef_history_controls()
                break
        
        if not patient_found:
            messagebox.showerror("Error", f"Patient with file number {self.current_file_number} not found")
            return
        
        self.save_patient_data()
        messagebox.showinfo("Success", f"Lab and EF data saved for patient {self.current_file_number}")

    def print_lab_ef_report(self):
        """Generate and open timeline report in Word with DD/MM/YY dates"""
        if not hasattr(self, 'current_lab_patient') or not self.current_lab_patient:
            messagebox.showwarning("Warning", "No patient selected")
            return

        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import tempfile
            import os
            from datetime import datetime

            # Create new document
            doc = Document()
            
            # Set document styles
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(10)

            # Add patient header
            patient = self.current_lab_patient
            doc.add_heading('Medical Results Timeline Report', level=1)
            patient_info = doc.add_paragraph()
            patient_info.add_run("Patient: ").bold = True
            patient_info.add_run(f"{patient.get('NAME', 'N/A')}\n")
            patient_info.add_run("File #: ").bold = True
            patient_info.add_run(f"{patient.get('FILE NUMBER', 'N/A')}\n")
            patient_info.add_run("Report Date: ").bold = True
            patient_info.add_run(datetime.now().strftime('%d/%m/%y'))

            # Process lab data
            lab_results = patient.get("lab_results", [])
            lab_dates = set()
            lab_tests = {test: {} for test in LAB_RANGES.keys()}
            for entry in lab_results:
                date_str = entry.get('date', 'Unknown Date')
                lab_dates.add(date_str)
                for test, value in entry.get('values', {}).items():
                    lab_tests[test][date_str] = value

            # Process EF data
            ef_data = patient.get("ef_data", [])
            ef_dates = set()
            ef_values = {'Baseline EF': {}, **{f"Serial EF {i+1}": {} for i in range(10)}}
            for entry in ef_data:
                baseline_date = entry.get('baseline', {}).get('date')
                if baseline_date:
                    ef_dates.add(baseline_date)
                    ef_values['Baseline EF'][baseline_date] = entry['baseline'].get('value')
                for i, serial in enumerate(entry.get('serial', [])):
                    key = f"Serial EF {i+1}"
                    serial_date = serial.get('date')
                    if serial_date:
                        ef_dates.add(serial_date)
                        ef_values[key][serial_date] = serial.get('value')

            # Date sorting function
            def sort_dates(date_str):
                try:
                    return datetime.strptime(date_str, "%d/%m/%Y")
                except:
                    return datetime.min

            # Create lab section
            sorted_lab_dates = sorted(lab_dates, key=sort_dates)
            if lab_dates:
                doc.add_heading('Laboratory Results Timeline', level=2)
                self.add_timeline_table(
                    doc=doc,
                    dates=sorted_lab_dates,
                    data=lab_tests,
                    row_header="Lab Test",
                    normal_ranges=LAB_RANGES
                )

            # Create EF section
            sorted_ef_dates = sorted(ef_dates, key=sort_dates)
            if ef_dates:
                doc.add_heading('Ejection Fraction Timeline', level=2)
                self.add_timeline_table(
                    doc=doc,
                    dates=sorted_ef_dates,
                    data=ef_values,
                    row_header="Measurement",
                    notes_column=True
                )

            # Save to temp file
            with tempfile.NamedTemporaryFile(suffix="_Medical_Report.docx", delete=False) as tmp_file:
                temp_path = tmp_file.name
                doc.save(temp_path)

            # Open in Word
            os.startfile(temp_path)
            
            # Show information message
            messagebox.showinfo(
                "Report Ready",
                "Document opened in Microsoft Word\n\n"
                "Please use Word's built-in options to:\n"
                "1. Review the report\n"
                "2. Save to permanent location if needed\n"
                "3. Print directly from Word\n\n"
                "Temporary file location:\n" + temp_path
            )

        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report: {str(e)}")
            if 'temp_path' in locals() and os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass

    def add_timeline_table(self, doc, dates, data, row_header, normal_ranges=None, notes_column=False):
        """Create timeline table with DD/MM/YY dates"""
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.shared import Pt, Inches

        if not dates:
            return

        # Create table
        col_count = len(dates) + 1 + (1 if notes_column else 0)
        table = doc.add_table(rows=1, cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Configure columns
        for col in table.columns:
            col.width = Inches(1.5) if col == 0 else Inches(0.8)
        if notes_column:
            table.columns[-1].width = Inches(2.0)

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = row_header
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

        for idx, date in enumerate(dates, 1):
            cell = hdr_cells[idx]
            cell.text = self.format_date(date)
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

        if notes_column:
            hdr_cells[-1].text = "Clinical Notes"

        # Data rows
        for key, values in data.items():
            if any(values.values()) or notes_column:
                row_cells = table.add_row().cells
                row_cells[0].text = key
                
                for idx, date in enumerate(dates, 1):
                    value = values.get(date, '')
                    cell = row_cells[idx]
                    cell.text = str(value) if value else '-'
                    
                    if normal_ranges and key in normal_ranges:
                        self.apply_value_formatting(cell, value, normal_ranges[key])

                if notes_column:
                    notes_cell = row_cells[-1]
                    if key == 'Baseline EF':
                        notes_cell.text = "Pre-treatment baseline measurement"
                    elif 'Serial' in key:
                        notes_cell.text = "Post-treatment follow-up measurement"

        # Formatting
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_VERTICAL.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    def apply_value_formatting(self, cell, value, ranges):
        """Apply color coding to abnormal values"""
        from docx.shared import RGBColor
        try:
            num_value = float(value)
            normal = ranges['normal_range']
            crit_low = ranges['critical_low']
            crit_high = ranges['critical_high']
            
            if crit_low and num_value < float(crit_low[1:]):
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            elif crit_high and num_value > float(crit_high[1:]):
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            else:
                if '-' in normal:
                    low, high = map(float, normal.split('-')[:2])
                    if num_value < low or num_value > high:
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
        except:
            pass

    def format_date(self, date_str):
        """Convert date to DD/MM/YY format with robust parsing"""
        from datetime import datetime
        
        formats_to_try = [
            "%d/%m/%Y",    # Original format (25/12/2023)
            "%Y-%m-%d",    # ISO format (2023-12-25)
            "%m/%d/%Y",    # US format (12/25/2023)
            "%d-%m-%Y",    # Alternate format (25-12-2023)
            "%d%m%Y"      # Compact format (25122023)
        ]
        
        for fmt in formats_to_try:
            try:
                dt = datetime.strptime(date_str, fmt)
                return dt.strftime("%d/%m/%y")  # Format as DD/MM/YY
            except ValueError:
                continue
        
        if len(date_str) > 0:
            return date_str  # Return original if parsing fails
        return "??/??/??"

    def show_statistics(self):
        """Show statistics window"""
        self.open_statistics_window()

    def show_extravasation_management(self):
        """Display chemotherapy extravasation management with improved layout and print function"""
        self.clear_frame()

        # Main container
        main_frame = tk.Frame(self.root, bg='white')
        main_frame.pack(expand=True, fill=tk.BOTH)
        
        # Left side with logo
        left_frame = tk.Frame(main_frame, bg='#3498db')
        left_frame.pack(side=tk.LEFT, fill=tk.Y, expand=False)
        
        logo_frame = tk.Frame(left_frame, bg='#3498db')
        logo_frame.pack(expand=True, fill=tk.BOTH, padx=40, pady=40)
        
        tk.Label(logo_frame, text="OncoCare", font=('Helvetica', 24, 'bold'), 
                bg='#3498db', fg='white').pack(pady=(0, 10))
        tk.Label(logo_frame, text="Extravasation Management", 
                font=('Helvetica', 14), bg='#3498db', fg='white').pack(pady=(0, 40))
        
        # Right side with content
        right_frame = tk.Frame(main_frame, bg='white')
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        # Create scrollable frame
        canvas = tk.Canvas(right_frame, highlightthickness=0)
        scrollbar = ttk.Scrollbar(right_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        
        # Mouse wheel scrolling
        def _on_mouse_wheel(event):
            canvas.yview_scroll(-1 * (event.delta // 120), "units")
        
        canvas.bind_all("<MouseWheel>", _on_mouse_wheel)
        
        # Add content with new layout
        self.add_extravasation_content(scrollable_frame)
        
        # Print button (outside scrollable area)
        btn_frame = ttk.Frame(right_frame, padding="10 10 10 10", style='TFrame')
        btn_frame.pack(fill=tk.X)
        
        print_btn = ttk.Button(btn_frame, text="Print Protocol", 
                             command=self.print_extravasation_protocol,
                             style='Blue.TButton')
        print_btn.pack(side=tk.RIGHT, padx=10)
        
        back_btn = ttk.Button(btn_frame, text="Back to Menu", 
                            command=self.main_menu,
                            style='Blue.TButton')
        back_btn.pack(side=tk.LEFT, padx=10)

    def add_extravasation_content(self, parent):
        """Add extravasation content with improved layout"""
        # Overview frame (top left)
        overview_frame = ttk.LabelFrame(parent, text="Chemotherapy Extravasation Overview", 
                                      style='TFrame')
        overview_frame.pack(fill=tk.X, padx=10, pady=5, anchor='nw')
        
        overview_text = (
            "Extravasation is the leakage of chemotherapy drugs into surrounding tissues. "
            "In pediatric patients, risk is higher due to smaller veins and movement.\n\n"
            "SEVERITY CLASSIFICATION:\n"
            "• Grade 1: Mild (erythema, swelling <1cm)\n"
            "• Grade 2: Moderate (pain, swelling 1-5cm)\n"
            "• Grade 3: Severe (ulceration, necrosis)\n"
            "• Grade 4: Life-threatening (compartment syndrome)"
        )
        ttk.Label(overview_frame, text=overview_text, justify='left', 
                 wraplength=700).pack(padx=10, pady=5)
        
        # General protocol frame (under overview)
        general_frame = ttk.LabelFrame(parent, text="General Management Protocol", 
                                     style='TFrame')
        general_frame.pack(fill=tk.X, padx=10, pady=5, anchor='nw')
        
        steps = [
        ("1. STOP INFUSION IMMEDIATELY", 
         "• Clamp the tubing closest to the IV site\n"
         "• Discontinue the infusion pump\n"
         "• Note exact time of recognition"),
        
        ("2. LEAVE CATHETER IN PLACE INITIALLY",
         "• Attempt to aspirate residual drug (use 1-3mL syringe for better suction)\n"
         "• Gently aspirate for at least 1 minute before removing\n"
         "• If unsuccessful, remove catheter after aspiration attempt"),
        
        ("3. ASSESS EXTENT OF EXTRAVASATION",
         "• Measure area of induration/erythema in cm\n"
         "• Photograph the site if possible\n"
         "• Grade severity (1-4) based on symptoms"),
        
        ("4. ADMINISTER SPECIFIC ANTIDOTE",
         "• See drug-specific protocols below\n"
         "• Prepare antidote within 15 minutes of recognition\n"
         "• Use 25-27G needle for subcutaneous administration"),
        
        ("5. APPLY TOPICAL MANAGEMENT",
         "• Warm compress (40-42°C) for vinca alkaloids\n"
         "• Cold compress (ice pack wrapped in cloth) for anthracyclines\n"
         "• Elevate extremity above heart level"),
        
        ("6. DOCUMENT THOROUGHLY",
         "• Complete incident report form\n"
         "• Record drug, concentration, volume extravasated\n"
         "• Document patient response and interventions"),
        
        ("7. NOTIFY PHYSICIAN/ONCOLOGIST",
         "• Immediate notification for grade 3-4 extravasation\n"
         "• Consider surgical consult for severe cases\n"
         "• Report to pharmacy for quality improvement")
        ]
        
        for step in steps:
            ttk.Label(general_frame, text=step, anchor='w').pack(fill=tk.X, padx=10, pady=2)
        
        # Drug selection frame
        drug_frame = ttk.LabelFrame(parent, text="Drug-Specific Protocol Selection",
                                  style='TFrame')
        drug_frame.pack(fill=tk.X, padx=10, pady=10)
        
        # Drug selection dropdown
        ttk.Label(drug_frame, text="Select Chemotherapy Drug:").pack(side=tk.LEFT, padx=5)

        self.drug_var = tk.StringVar()
        drug_dropdown = ttk.Combobox(drug_frame, textvariable=self.drug_var, 
                                   state="readonly", width=40)  # Increased width from 30 to 40

        # Configure the dropdown list style to be wider
        style = ttk.Style()
        style.configure('TCombobox', postoffset=(0,0,400,0))  # Adjust the 400 value as needed

        drug_dropdown.pack(side=tk.LEFT, padx=5)   
             
        # Initialize drug data with detailed protocols
        self.drug_protocols = {
            "Anthracyclines (Doxorubicin, Daunorubicin, Epirubicin)": {
                "Risk": "Vesicant (High Risk)\n• Severe tissue necrosis\n• Delayed onset (weeks to months)",
                "Antidote": (
                    "Dexrazoxane (Totect)\n"
                    "• Dose: 10:1 ratio to anthracycline dose (max 1000mg/m²/day)\n"
                    "• Schedule:\n"
                    "  - First dose: Within 6 hours of extravasation\n"
                    "  - Second dose: 24 hours after first dose\n"
                    "  - Third dose: 24 hours after second dose\n"
                    "• Infuse over 1-2 hours in large vein\n"
                    "• Preparation: Reconstitute with 25mL sterile water (50mg/mL)"
                ),
                "Local Care": (
                    "• Apply cold compress immediately (15-20 minutes QID x 3 days)\n"
                    "• Avoid pressure on affected area\n"
                    "• Consider topical DMSO 99% (apply thin layer Q8H x 7 days)"
                ),
                "Monitoring": (
                    "• Daily wound assessment for 7 days\n"
                    "• Monitor for delayed necrosis (may appear 1-4 weeks post)\n"
                    "• Consider MRI for deep tissue assessment"
                ),
                "Special Notes": (
                    "• Delayed tissue damage common (weeks to months later)\n"
                    "• Surgical debridement often required for necrosis\n"
                    "• Higher risk in infants and young children"
                )
            },
            "Vinca Alkaloids (Vincristine, Vinblastine, Vinorelbine)": {
                "Risk": "Vesicant (Moderate Risk)\n• Neurotoxicity possible\n• More severe with vincristine",
                "Antidote": (
                    "Hyaluronidase (Vitrase, Amphadase)\n"
                    "• Dose: 150-1500 units (1mL of 150 units/mL solution)\n"
                    "• Administration:\n"
                    "  - Inject subcutaneously around extravasation site\n"
                    "  - Use 25G needle, change puncture sites for multiple injections\n"
                    "  - May repeat after 1 hour if symptoms persist"
                ),
                "Local Care": (
                    "• Apply WARM compress (40-42°C) for 15-20 minutes QID x 3 days\n"
                    "• Elevate extremity\n"
                    "• Gentle massage to promote dispersion"
                ),
                "Monitoring": (
                    "• Assess neurovascular status Q2H x 24 hours\n"
                    "• Monitor for compartment syndrome\n"
                    "• Document sensory/motor function"
                ),
                "Special Notes": (
                    "• Neurotoxicity may occur without visible skin changes\n"
                    "• Consider nerve conduction studies if neurological symptoms\n"
                    "• More severe in patients with pre-existing neuropathy"
                )
            },
            "Platinum Compounds (Cisplatin, Carboplatin, Oxaliplatin)": {
                "Risk": "Irritant (Low-Moderate Risk)\n• Cisplatin more toxic than carboplatin\n• Oxaliplatin may cause cold-induced neuropathy",
                "Antidote": (
                    "No specific antidote\n"
                    "• Sodium thiosulfate 1/6M may be used (off-label)\n"
                    "• For cisplatin: Consider systemic sodium thiosulfate"
                ),
                "Local Care": (
                    "• Apply cold compress for 15-20 minutes QID x 24 hours\n"
                    "• Elevate extremity\n"
                    "• For oxaliplatin: Avoid cold exposure to affected area"
                ),
                "Monitoring": (
                    "• Daily assessment for 3 days\n"
                    "• Monitor for hypersensitivity reactions\n"
                    "• Check renal function if systemic absorption suspected"
                ),
                "Special Notes": (
                    "• Tissue damage usually resolves within 2 weeks\n"
                    "• Oxaliplatin extravasation may cause prolonged cold sensitivity\n"
                    "• Higher risk with concentrated solutions"
                )
            },
            "Alkylating Agents (Cyclophosphamide, Ifosfamide)": {
                "Risk": "Irritant (Low Risk)\n• Usually mild tissue damage\n• Higher risk with concentrated solutions",
                "Antidote": "No specific antidote",
                "Local Care": (
                    "• Cold compress for 15-20 minutes QID x 24 hours\n"
                    "• Topical corticosteroids may reduce inflammation\n"
                    "• Saline flush of affected area (off-label use)"
                ),
                "Monitoring": (
                    "• Assess site Q8H x 48 hours\n"
                    "• Monitor for infection\n"
                    "• Check urine for hemorrhagic cystitis (ifosfamide)"
                ),
                "Special Notes": (
                    "• Mesna not effective for local tissue damage\n"
                    "• Healing typically occurs within 1-2 weeks\n"
                    "• Rarely requires surgical intervention"
                )
            },
            "Taxanes (Paclitaxel, Docetaxel)": {
                "Risk": "Irritant (Moderate Risk)\n• Cremophor-containing formulations more irritating\n• Delayed reactions may occur 3-10 days post",
                "Antidote": "No specific antidote\n• Hyaluronidase may be considered (off-label)",
                "Local Care": (
                    "• Cold compress for 15-20 minutes QID x 24 hours\n"
                    "• Topical hydrocortisone 1% cream BID\n"
                    "• For severe reactions: Consider oral corticosteroids"
                ),
                "Monitoring": (
                    "• Assess for hypersensitivity reactions\n"
                    "• Monitor for neuropathic pain\n"
                    "• Document resolution of erythema"
                ),
                "Special Notes": (
                    "• More severe in patients with prior radiation to site\n"
                    "• Albumin-bound paclitaxel less likely to cause severe reactions\n"
                    "• May cause recall reactions at previous extravasation sites"
                )
            },
            "Etoposide": {
                "Risk": "Irritant (Low Risk)\n• Usually mild reactions\n• Rarely causes tissue necrosis",
                "Antidote": "No specific antidote",
                "Local Care": (
                    "• Cold compress for 15-20 minutes QID x 24 hours\n"
                    "• Topical corticosteroids for persistent inflammation"
                ),
                "Monitoring": "• Monitor site for 48 hours\n• Watch for hypersensitivity reactions",
                "Special Notes": "• Reactions typically resolve within 3-5 days\n• Higher risk with concentrated solutions"
            },
            "Methotrexate": {
                "Risk": "Irritant (Low Risk)\n• Mild local reaction\n• Rarely causes tissue damage",
                "Antidote": "Consider systemic leucovorin\n• Dose based on methotrexate exposure",
                "Local Care": (
                    "• Cold compress for 15-20 minutes QID x 24 hours\n"
                    "• Elevate extremity\n"
                    "• Alkalinization of urine if systemic absorption suspected"
                ),
                "Monitoring": (
                    "• Assess site Q8H x 48h\n"
                    "• Monitor renal function if significant absorption\n"
                    "• Check methotrexate levels if concern for systemic exposure"
                ),
                "Special Notes": (
                    "• Usually mild tissue damage\n"
                    "• Rarely requires surgical intervention\n"
                    "• Higher risk in patients with third spacing"
                )
            },
            "Bleomycin": {
                "Risk": "Non-vesicant\n• Rarely causes tissue damage\n• Minimal local reaction",
                "Antidote": "Not applicable",
                "Local Care": "• Observation only\n• Routine wound care if skin breakdown occurs",
                "Monitoring": "• Minimal monitoring required\n• Assess for infection if skin breaks",
                "Special Notes": "• Does not typically require specific treatment\n• Rarely causes significant tissue damage"
            },
            "Asparaginase": {
                "Risk": "Non-vesicant\n• Primary concern is hypersensitivity\n• Minimal local tissue effects",
                "Antidote": "Not applicable\n• Treat hypersensitivity reactions if they occur",
                "Local Care": "• Routine wound care\n• Cold compress if inflammation present",
                "Monitoring": "• Watch for allergic reactions\n• Monitor for infection",
                "Special Notes": "• Tissue damage extremely rare\n• More concern for systemic reactions than local effects"
            }
        }
        
        drug_dropdown['values'] = list(self.drug_protocols.keys())
        drug_dropdown.bind("<<ComboboxSelected>>", self.display_drug_protocol)
        
        # Protocol display frame
        self.protocol_display = ttk.LabelFrame(parent, text="Selected Drug Protocol",
                                             style='TFrame')
        self.protocol_display.pack(fill=tk.BOTH, padx=10, pady=5, expand=True)
        
        # Initialize with empty display
        self.display_drug_protocol()

    def display_drug_protocol(self, event=None):
        """Display protocol for selected drug"""
        selected_drug = self.drug_var.get()
        
        # Clear previous content
        for widget in self.protocol_display.winfo_children():
            widget.destroy()
        
        if not selected_drug:
            ttk.Label(self.protocol_display, 
                     text="Please select a chemotherapy drug from the dropdown above",
                     style='TLabel').pack(padx=10, pady=20)
            return
        
        protocol = self.drug_protocols.get(selected_drug, {})
        
        # Create notebook for different sections
        notebook = ttk.Notebook(self.protocol_display)
        notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Risk Assessment tab
        risk_frame = ttk.Frame(notebook)
        notebook.add(risk_frame, text="Risk Assessment")
        ttk.Label(risk_frame, text=protocol.get("Risk", "N/A"), 
                 justify='left', wraplength=650).pack(padx=10, pady=10)
        
        # Antidote tab
        antidote_frame = ttk.Frame(notebook)
        notebook.add(antidote_frame, text="Antidote")
        ttk.Label(antidote_frame, text=protocol.get("Antidote", "N/A"), 
                 justify='left', wraplength=650).pack(padx=10, pady=10)
        
        # Local Care tab
        local_frame = ttk.Frame(notebook)
        notebook.add(local_frame, text="Local Care")
        ttk.Label(local_frame, text=protocol.get("Local Care", "N/A"), 
                 justify='left', wraplength=650).pack(padx=10, pady=10)
        
        # Monitoring tab
        monitor_frame = ttk.Frame(notebook)
        notebook.add(monitor_frame, text="Monitoring")
        ttk.Label(monitor_frame, text=protocol.get("Monitoring", "N/A"), 
                 justify='left', wraplength=650).pack(padx=10, pady=10)
        
        # Special Notes tab
        notes_frame = ttk.Frame(notebook)
        notebook.add(notes_frame, text="Special Notes")
        ttk.Label(notes_frame, text=protocol.get("Special Notes", "N/A"), 
                 justify='left', wraplength=650).pack(padx=10, pady=10)

    def print_extravasation_protocol(self):
        """Print the current extravasation protocol"""
        selected_drug = self.drug_var.get()
        if not selected_drug:
            messagebox.showwarning("Print", "Please select a chemotherapy drug first")
            return
        
        protocol = self.drug_protocols.get(selected_drug, {})
        
        # Create a printable document
        doc = Document()
        doc.add_heading(f'Chemotherapy Extravasation Protocol: {selected_drug}', level=1)
        
        # Add overview
        doc.add_heading('Overview', level=2)
        doc.add_paragraph(
            "Extravasation is the leakage of chemotherapy drugs into surrounding tissues. "
            "Immediate recognition and proper management are crucial to minimize complications."
        )
        
        # Add general protocol
        doc.add_heading('General Management Protocol', level=2)
        general_steps = doc.add_paragraph()
        general_steps.add_run("1. STOP INFUSION IMMEDIATELY\n")
        general_steps.add_run("• Clamp the tubing closest to the IV site\n")
        general_steps.add_run("• Discontinue the infusion pump\n")
        general_steps.add_run("• Note exact time of recognition\n\n")
        
        general_steps.add_run("2. LEAVE CATHETER IN PLACE INITIALLY\n")
        general_steps.add_run("• Attempt to aspirate residual drug (use 1-3mL syringe for better suction)\n")
        general_steps.add_run("• Gently aspirate for at least 1 minute before removing\n")
        general_steps.add_run("• If unsuccessful, remove catheter after aspiration attempt\n\n")
        
        general_steps.add_run("3. ASSESS EXTENT OF EXTRAVASATION\n")
        general_steps.add_run("• Measure area of induration/erythema in cm\n")
        general_steps.add_run("• Photograph the site if possible\n")
        general_steps.add_run("• Grade severity (1-4) based on symptoms\n\n")
        
        general_steps.add_run("4. ADMINISTER SPECIFIC ANTIDOTE\n")
        general_steps.add_run("• See drug-specific protocols below\n")
        general_steps.add_run("• Prepare antidote within 15 minutes of recognition\n")
        general_steps.add_run("• Use 25-27G needle for subcutaneous administration\n\n")
        
        general_steps.add_run("5. APPLY TOPICAL MANAGEMENT\n")
        general_steps.add_run("• Warm compress (40-42°C) for vinca alkaloids\n")
        general_steps.add_run("• Cold compress (ice pack wrapped in cloth) for anthracyclines\n")
        general_steps.add_run("• Elevate extremity above heart level\n\n")
        
        general_steps.add_run("6. DOCUMENT THOROUGHLY\n")
        general_steps.add_run("• Complete incident report form\n")
        general_steps.add_run("• Record drug, concentration, volume extravasated\n")
        general_steps.add_run("• Document patient response and interventions\n\n")
        
        general_steps.add_run("7. NOTIFY PHYSICIAN/ONCOLOGIST\n")
        general_steps.add_run("• Immediate notification for grade 3-4 extravasation\n")
        general_steps.add_run("• Consider surgical consult for severe cases\n")
        general_steps.add_run("• Report to pharmacy for quality improvement")

        # Add drug-specific protocol
        doc.add_heading(f'{selected_drug} Specific Protocol', level=2)
        
        doc.add_heading('Risk Assessment', level=3)
        doc.add_paragraph(protocol.get("Risk", "N/A"))
        
        doc.add_heading('Antidote', level=3)
        doc.add_paragraph(protocol.get("Antidote", "N/A"))
        
        doc.add_heading('Local Care', level=3)
        doc.add_paragraph(protocol.get("Local Care", "N/A"))
        
        doc.add_heading('Monitoring', level=3)
        doc.add_paragraph(protocol.get("Monitoring", "N/A"))
        
        doc.add_heading('Special Notes', level=3)
        doc.add_paragraph(protocol.get("Special Notes", "N/A"))
        
        # Add footer
        doc.add_paragraph("\nGenerated by OncoCare Pediatric Oncology System")
        doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M"))
        
        # Save the document to a temporary file
        temp_file = "temp_extravasation_protocol.docx"
        doc.save(temp_file)
        
        # Open the document for the user to view and save if desired
        try:
            os.startfile(temp_file)
            messagebox.showinfo("Document Opened", "Protocol document opened. You can save it if you wish.")
        except Exception as e:
            messagebox.showerror("Open Error", f"Could not open document: {str(e)}")
        finally:
            # Clean up after a delay
            def delete_temp_file():
                try:
                    os.remove(temp_file)
                except:
                    pass
            
            # Schedule file deletion after a delay
            self.root.after(5000, delete_temp_file)

    def show_calculators(self):
        """Display all calculators in tabs"""
        self.clear_frame()
        
        main_frame = ttk.Frame(self.root)
        main_frame.pack(expand=True, fill=tk.BOTH)
        
        notebook = ttk.Notebook(main_frame)
        notebook.pack(fill=tk.BOTH, expand=True)
        
        # Add all calculator tabs
        calculators = [
            ("BSA Calculator", self.setup_bsa_calculator),
            ("IV Fluid Calculator", self.setup_iv_calculator),
            ("Chemo Dosage Calculator", self.setup_dosage_calculator),
            ("Antibiotics Calculator", self.setup_antibiotics_calculator)
        ]
        
        for text, setup_func in calculators:
            frame = ttk.Frame(notebook)
            notebook.add(frame, text=text)
            setup_func(frame)
        
        # Back button
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=10)
        ttk.Button(btn_frame, text="Back to Menu", command=self.main_menu).pack()
        
    def setup_antibiotics_calculator(self, parent):
        """Setup the antibiotics calculator interface for pediatric oncology patients with enhanced features"""
        # Create main container with scrollable canvas
        main_frame = ttk.Frame(parent)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        canvas = tk.Canvas(main_frame, borderwidth=0)
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind_all("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))
        
        # Configure age groups
        self.age_groups = [
            "0-30 days (Neonate)",
            "1 month-1 year (Infant)",
            "1-6 years (Young child)",
            "6-12 years (Older child)",
            "12+ years (Adolescent)"
        ]
        
        # Oncology clinical protocols
        self.protocols = [
            "Empiric Febrile Neutropenia",
            "Central Line Infection",
            "Mucositis Coverage",
            "Post-Chemotherapy Sepsis",
            "CNS Infection Coverage"
        ]
        
        # Expanded antibiotics data with oncology-specific information
        self.antibiotics_data = {
            "MEROPENEM": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 20, "max_dose": 40, "max_total": 2000, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 20, "max_dose": 40, "max_total": 2000, "frequency": "Every 8 hours"},
                    "1-6 years (Young child)": {"min_dose": 20, "max_dose": 40, "max_total": 2000, "frequency": "Every 8 hours"},
                    "6-12 years (Older child)": {"min_dose": 20, "max_dose": 40, "max_total": 2000, "frequency": "Every 8 hours"},
                    "12+ years (Adolescent)": {"min_dose": 20, "max_dose": 40, "max_total": 2000, "frequency": "Every 8 hours"}
                },
                "unit": "mg",
                "incompatible": ["Dextrose >5%", "Other beta-lactams", "Aminoglycosides"],
                "interactions": {
                    "Probenecid": "Decreases renal clearance",
                    "Valproic acid": "Reduces levels (avoid combination)",
                    "Warfarin": "Increased INR risk"
                },
                "oncology_notes": {
                    "penetration": "Excellent CNS penetration at higher doses",
                    "neutropenia": "First-line in febrile neutropenia protocols",
                    "mucositis": "Covers Gram-negatives including Pseudomonas",
                    "renal_adjust": "Reduce dose by 50% if eGFR <30",
                    "compatibility": "Safe with most chemotherapy regimens",
                    "monitoring": "Monitor seizure risk with high doses"
                },
                "notes": "First-line for febrile neutropenia. Higher doses for CNS infections."
            },

            "AMIKACIN": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 15, "max_dose": 20, "max_total": 1500, "frequency": "Every 24 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 15, "max_dose": 22.5, "max_total": 1500, "frequency": "Every 24 hours"},
                    "1-6 years (Young child)": {"min_dose": 15, "max_dose": 22.5, "max_total": 1500, "frequency": "Every 24 hours"},
                    "6-12 years (Older child)": {"min_dose": 15, "max_dose": 22.5, "max_total": 1500, "frequency": "Every 24 hours"},
                    "12+ years (Adolescent)": {"min_dose": 15, "max_dose": 22.5, "max_total": 1500, "frequency": "Every 24 hours"}
                },
                "unit": "mg",
                "incompatible": ["Penicillins", "Cephalosporins", "Heparin"],
                "interactions": {
                    "Vancomycin": "Increased nephrotoxicity",
                    "Diuretics": "Ototoxicity risk",
                    "Cisplatin": "Increased toxicity"
                },
                "oncology_notes": {
                    "penetration": "Poor CNS penetration",
                    "neutropenia": "Synergistic with beta-lactams for Gram-negative coverage",
                    "mucositis": "No direct activity",
                    "renal_adjust": "Extend interval to 36-48h if CrCl <30",
                    "compatibility": "Avoid concurrent nephrotoxic chemo",
                    "monitoring": "Trough <5 mg/L, peak 20-30 mg/L"
                },
                "notes": "Monitor levels. Once daily dosing preferred."
            },

            "CEFTRIAXONE (ROCEPHINE)": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 50, "max_dose": 80, "max_total": 4000, "frequency": "Every 24 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 50, "max_dose": 100, "max_total": 4000, "frequency": "Every 12-24 hours"},
                    "1-6 years (Young child)": {"min_dose": 50, "max_dose": 100, "max_total": 4000, "frequency": "Every 12-24 hours"},
                    "6-12 years (Older child)": {"min_dose": 50, "max_dose": 100, "max_total": 4000, "frequency": "Every 12-24 hours"},
                    "12+ years (Adolescent)": {"min_dose": 50, "max_dose": 100, "max_total": 4000, "frequency": "Every 12-24 hours"}
                },
                "unit": "mg",
                "incompatible": ["Calcium-containing solutions", "Aminoglycosides"],
                "interactions": {
                    "Warfarin": "Increased effect",
                    "Calcium": "Precipitation risk (especially neonates)",
                    "Probenecid": "Increased levels"
                },
                "oncology_notes": {
                    "penetration": "Good CSF penetration in meningitis",
                    "neutropenia": "Alternative for stable patients",
                    "mucositis": "Covers viridans streptococci",
                    "renal_adjust": "No adjustment needed",
                    "compatibility": "Avoid with calcium-containing fluids",
                    "monitoring": "Watch for biliary pseudolithiasis"
                },
                "notes": "Avoid in hyperbilirubinemic neonates."
            },

            "CIPROFLOXACIN": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 10, "max_dose": 15, "max_total": 800, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 10, "max_dose": 20, "max_total": 800, "frequency": "Every 12 hours"},
                    "1-6 years (Young child)": {"min_dose": 10, "max_dose": 20, "max_total": 800, "frequency": "Every 12 hours"},
                    "6-12 years (Older child)": {"min_dose": 10, "max_dose": 20, "max_total": 800, "frequency": "Every 12 hours"},
                    "12+ years (Adolescent)": {"min_dose": 10, "max_dose": 20, "max_total": 800, "frequency": "Every 12 hours"}
                },
                "unit": "mg",
                "incompatible": ["Divalent cation solutions", "TPN"],
                "interactions": {
                    "Antacids": "Reduced absorption",
                    "Theophylline": "Increased levels",
                    "Warfarin": "Increased INR"
                },
                "oncology_notes": {
                    "penetration": "Good tissue penetration",
                    "neutropenia": "Avoid as monotherapy",
                    "mucositis": "Gram-negative coverage only",
                    "renal_adjust": "Reduce dose by 50% if CrCl <30",
                    "compatibility": "Separate from oral supplements",
                    "monitoring": "Watch for tendonitis"
                },
                "notes": "Reserve for resistant gram-negative infections."
            },

            "TAZOCIN": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 80, "max_dose": 100, "max_total": 4000, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 80, "max_dose": 100, "max_total": 4000, "frequency": "Every 8 hours"},
                    "1-6 years (Young child)": {"min_dose": 80, "max_dose": 100, "max_total": 4000, "frequency": "Every 6 hours"},
                    "6-12 years (Older child)": {"min_dose": 80, "max_dose": 100, "max_total": 4000, "frequency": "Every 6 hours"},
                    "12+ years (Adolescent)": {"min_dose": 80, "max_dose": 100, "max_total": 4000, "frequency": "Every 6 hours"}
                },
                "unit": "mg (piperacillin component)",
                "incompatible": ["Aminoglycosides", "Vancomycin"],
                "interactions": {
                    "Probenecid": "Increased levels",
                    "Anticoagulants": "Increased bleeding risk",
                    "Methotrexate": "Increased toxicity"
                },
                "oncology_notes": {
                    "penetration": "Good tissue penetration",
                    "neutropenia": "Alternative antipseudomonal agent",
                    "mucositis": "Broad Gram-negative coverage",
                    "renal_adjust": "Adjust frequency based on CrCl",
                    "compatibility": "Monitor platelets with chemotherapy",
                    "monitoring": "Check electrolytes for hypokalemia"
                },
                "notes": "Piperacillin/tazobactam 8:1 ratio."
            },

            "VANCOMYCIN": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 10, "max_dose": 15, "max_total": 1000, "frequency": "Every 24 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 10, "max_dose": 15, "max_total": 1000, "frequency": "Every 12 hours"},
                    "1-6 years (Young child)": {"min_dose": 10, "max_dose": 15, "max_total": 1000, "frequency": "Every 8 hours"},
                    "6-12 years (Older child)": {"min_dose": 10, "max_dose": 15, "max_total": 1000, "frequency": "Every 6 hours"},
                    "12+ years (Adolescent)": {"min_dose": 10, "max_dose": 15, "max_total": 1000, "frequency": "Every 6 hours"}
                },
                "unit": "mg",
                "incompatible": ["Alkaline solutions", "Beta-lactams", "Heparin"],
                "interactions": {
                    "Aminoglycosides": "Nephrotoxicity risk",
                    "Loop diuretics": "Ototoxicity risk",
                    "NSAIDs": "Nephrotoxicity"
                },
                "oncology_notes": {
                    "penetration": "Poor CSF penetration",
                    "neutropenia": "Add for suspected MRSA",
                    "mucositis": "Covers coagulase-negative staph",
                    "renal_adjust": "Use trough-guided dosing",
                    "compatibility": "Monitor with nephrotoxic chemo",
                    "monitoring": "Trough 15-20 mg/L for serious infections"
                },
                "notes": "Pre-medicate for red man syndrome."
            },

            "METRONIDAZOLE (FLAGYL)": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 7.5, "max_dose": 10, "max_total": 500, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 7.5, "max_dose": 10, "max_total": 500, "frequency": "Every 8 hours"},
                    "1-6 years (Young child)": {"min_dose": 7.5, "max_dose": 10, "max_total": 500, "frequency": "Every 8 hours"},
                    "6-12 years (Older child)": {"min_dose": 7.5, "max_dose": 10, "max_total": 500, "frequency": "Every 8 hours"},
                    "12+ years (Adolescent)": {"min_dose": 7.5, "max_dose": 10, "max_total": 500, "frequency": "Every 8 hours"}
                },
                "unit": "mg",
                "incompatible": ["Aluminum-containing solutions"],
                "interactions": {
                    "Alcohol": "Disulfiram-like reaction",
                    "Warfarin": "Increased effect",
                    "Phenytoin": "Increased levels"
                },
                "oncology_notes": {
                    "penetration": "Excellent anaerobic coverage",
                    "neutropenia": "Add for abdominal symptoms",
                    "mucositis": "Covers anaerobic overgrowth",
                    "renal_adjust": "No adjustment needed",
                    "compatibility": "Safe with most regimens",
                    "monitoring": "Neurological toxicity at high doses"
                },
                "notes": "IV form contains sodium."
            },

            "FLUCONAZOLE": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 6, "max_dose": 12, "max_total": 800, "frequency": "Every 48 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 6, "max_dose": 12, "max_total": 800, "frequency": "Every 24 hours"},
                    "1-6 years (Young child)": {"min_dose": 6, "max_dose": 12, "max_total": 800, "frequency": "Every 24 hours"},
                    "6-12 years (Older child)": {"min_dose": 6, "max_dose": 12, "max_total": 800, "frequency": "Every 24 hours"},
                    "12+ years (Adolescent)": {"min_dose": 6, "max_dose": 12, "max_total": 800, "frequency": "Every 24 hours"}
                },
                "unit": "mg",
                "incompatible": ["None known"],
                "interactions": {
                    "Phenytoin": "Increased levels",
                    "Warfarin": "Increased INR",
                    "Rifampin": "Decreased fluconazole levels"
                },
                "oncology_notes": {
                    "penetration": "Good CSF penetration",
                    "neutropenia": "Prophylaxis in high-risk patients",
                    "mucositis": "Candida coverage",
                    "renal_adjust": "Double dosing interval if CrCl <50",
                    "compatibility": "Avoid with vinca alkaloids",
                    "monitoring": "Check LFTs weekly"
                },
                "notes": "Loading dose recommended."
            },

            "AMPHOTERICIN B": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 0.5, "max_dose": 1, "max_total": 50, "frequency": "Every 24 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 0.5, "max_dose": 1, "max_total": 50, "frequency": "Every 24 hours"},
                    "1-6 years (Young child)": {"min_dose": 0.5, "max_dose": 1, "max_total": 50, "frequency": "Every 24 hours"},
                    "6-12 years (Older child)": {"min_dose": 0.5, "max_dose": 1, "max_total": 50, "frequency": "Every 24 hours"},
                    "12+ years (Adolescent)": {"min_dose": 0.5, "max_dose": 1, "max_total": 50, "frequency": "Every 24 hours"}
                },
                "unit": "mg/kg",
                "incompatible": ["Saline solutions", "Other antifungals"],
                "interactions": {
                    "Cyclosporine": "Increased nephrotoxicity",
                    "Diuretics": "Hypokalemia risk",
                    "Azoles": "Antagonism"
                },
                "oncology_notes": {
                    "penetration": "Poor CSF penetration",
                    "neutropenia": "Empiric fungal coverage",
                    "mucositis": "No direct activity",
                    "renal_adjust": "Liposomal form preferred in renal impairment",
                    "compatibility": "Nephrotoxic with cisplatin",
                    "monitoring": "Daily electrolytes, renal function"
                },
                "notes": "Pre-medicate with antipyretics/antihistamines."
            },

            "VORICONAZOLE": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 7, "max_dose": 8, "max_total": 400, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 7, "max_dose": 8, "max_total": 400, "frequency": "Every 12 hours"},
                    "1-6 years (Young child)": {"min_dose": 7, "max_dose": 8, "max_total": 400, "frequency": "Every 12 hours"},
                    "6-12 years (Older child)": {"min_dose": 7, "max_dose": 8, "max_total": 400, "frequency": "Every 12 hours"},
                    "12+ years (Adolescent)": {"min_dose": 7, "max_dose": 8, "max_total": 400, "frequency": "Every 12 hours"}
                },
                "unit": "mg",
                "incompatible": ["None known"],
                "interactions": {
                    "Rifampin": "Decreased levels",
                    "Carbamazepine": "Decreased levels",
                    "Warfarin": "Increased INR"
                },
                "oncology_notes": {
                    "penetration": "Good CSF penetration",
                    "neutropenia": "First-line for aspergillosis",
                    "mucositis": "No direct activity",
                    "renal_adjust": "IV form contraindicated in CrCl <50",
                    "compatibility": "Adjust vincristine doses",
                    "monitoring": "Therapeutic drug monitoring essential"
                },
                "notes": "Visual disturbances common."
            },

            "AUGMENTIN": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 25, "max_dose": 30, "max_total": 1750, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 25, "max_dose": 45, "max_total": 1750, "frequency": "Every 8 hours"},
                    "1-6 years (Young child)": {"min_dose": 25, "max_dose": 45, "max_total": 1750, "frequency": "Every 8 hours"},
                    "6-12 years (Older child)": {"min_dose": 25, "max_dose": 45, "max_total": 1750, "frequency": "Every 8 hours"},
                    "12+ years (Adolescent)": {"min_dose": 25, "max_dose": 45, "max_total": 1750, "frequency": "Every 8 hours"}
                },
                "unit": "mg (amoxicillin component)",
                "incompatible": ["None known"],
                "interactions": {
                    "Probenecid": "Increased levels",
                    "Allopurinol": "Increased rash risk",
                    "Warfarin": "Increased INR"
                },
                "oncology_notes": {
                    "penetration": "Good oral bioavailability",
                    "neutropenia": "Not for empiric therapy",
                    "mucositis": "Community-acquired infections",
                    "renal_adjust": "Adjust based on CrCl",
                    "compatibility": "Diarrhea risk with chemo",
                    "monitoring": "Watch for C.diff"
                },
                "notes": "Amoxicillin/clavulanate 7:1 ratio."
            },

            "ACYCLOVIR": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 10, "max_dose": 20, "max_total": 800, "frequency": "Every 8 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 10, "max_dose": 20, "max_total": 800, "frequency": "Every 8 hours"},
                    "1-6 years (Young child)": {"min_dose": 10, "max_dose": 20, "max_total": 800, "frequency": "Every 8 hours"},
                    "6-12 years (Older child)": {"min_dose": 10, "max_dose": 20, "max_total": 800, "frequency": "Every 8 hours"},
                    "12+ years (Adolescent)": {"min_dose": 10, "max_dose": 20, "max_total": 800, "frequency": "Every 8 hours"}
                },
                "unit": "mg",
                "incompatible": ["Alkaline solutions"],
                "interactions": {
                    "Probenecid": "Increased levels",
                    "Nephrotoxic drugs": "Increased toxicity",
                    "Zidovudine": "Neurotoxicity"
                },
                "oncology_notes": {
                    "penetration": "Good CSF penetration",
                    "neutropenia": "HSV prophylaxis in stem cell transplants",
                    "mucositis": "Herpetic lesions coverage",
                    "renal_adjust": "Adjust dose for CrCl <50",
                    "compatibility": "Hydrate well with cyclophosphamide",
                    "monitoring": "Renal function every 48h"
                },
                "notes": "Higher doses for HSV encephalitis."
            },

            "AMPICILLIN": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 25, "max_dose": 50, "max_total": 2000, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 25, "max_dose": 50, "max_total": 2000, "frequency": "Every 6 hours"},
                    "1-6 years (Young child)": {"min_dose": 25, "max_dose": 50, "max_total": 2000, "frequency": "Every 6 hours"},
                    "6-12 years (Older child)": {"min_dose": 25, "max_dose": 50, "max_total": 2000, "frequency": "Every 6 hours"},
                    "12+ years (Adolescent)": {"min_dose": 25, "max_dose": 50, "max_total": 2000, "frequency": "Every 6 hours"}
                },
                "unit": "mg",
                "incompatible": ["Aminoglycosides"],
                "interactions": {
                    "Probenecid": "Increased levels",
                    "Allopurinol": "Increased rash risk",
                    "Warfarin": "Increased INR"
                },
                "oncology_notes": {
                    "penetration": "Limited CSF penetration",
                    "neutropenia": "Listeria coverage",
                    "mucositis": "Gram-positive coverage",
                    "renal_adjust": "Adjust frequency for CrCl <10",
                    "compatibility": "Monitor with methotrexate",
                    "monitoring": "Rash common in EBV patients"
                },
                "notes": "Adjust dose in renal impairment."
            },

            "CLARITHROMYCIN": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 7.5, "max_dose": 10, "max_total": 1000, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 7.5, "max_dose": 15, "max_total": 1000, "frequency": "Every 12 hours"},
                    "1-6 years (Young child)": {"min_dose": 7.5, "max_dose": 15, "max_total": 1000, "frequency": "Every 12 hours"},
                    "6-12 years (Older child)": {"min_dose": 7.5, "max_dose": 15, "max_total": 1000, "frequency": "Every 12 hours"},
                    "12+ years (Adolescent)": {"min_dose": 7.5, "max_dose": 15, "max_total": 1000, "frequency": "Every 12 hours"}
                },
                "unit": "mg",
                "incompatible": ["None known"],
                "interactions": {
                    "Digoxin": "Increased levels",
                    "Theophylline": "Increased levels",
                    "Warfarin": "Increased INR"
                },
                "oncology_notes": {
                    "penetration": "Good intracellular penetration",
                    "neutropenia": "Atypical coverage in allergies",
                    "mucositis": "Mycoplasma coverage",
                    "renal_adjust": "Reduce dose by 50% if CrCl <30",
                    "compatibility": "QT prolongation with anthracyclines",
                    "monitoring": "ECG for QT interval"
                },
                "notes": "QT prolongation risk."
            },

            "CEFTAZIDIME": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 30, "max_dose": 50, "max_total": 2000, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 30, "max_dose": 50, "max_total": 2000, "frequency": "Every 8 hours"},
                    "1-6 years (Young child)": {"min_dose": 30, "max_dose": 50, "max_total": 2000, "frequency": "Every 8 hours"},
                    "6-12 years (Older child)": {"min_dose": 30, "max_dose": 50, "max_total": 2000, "frequency": "Every 8 hours"},
                    "12+ years (Adolescent)": {"min_dose": 30, "max_dose": 50, "max_total": 2000, "frequency": "Every 8 hours"}
                },
                "unit": "mg",
                "incompatible": ["Aminoglycosides", "Vancomycin"],
                "interactions": {
                    "Probenecid": "Increased levels",
                    "Loop diuretics": "Nephrotoxicity risk"
                },
                "oncology_notes": {
                    "penetration": "Limited CNS penetration",
                    "neutropenia": "Pseudomonas coverage",
                    "mucositis": "Gram-negative spectrum",
                    "renal_adjust": "Adjust for CrCl <50",
                    "compatibility": "Safe with most regimens",
                    "monitoring": "Watch for neurotoxicity"
                },
                "notes": "Adjust dose in renal impairment."
            },

            "CLOXACILLIN": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 25, "max_dose": 50, "max_total": 2000, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 25, "max_dose": 50, "max_total": 2000, "frequency": "Every 6 hours"},
                    "1-6 years (Young child)": {"min_dose": 25, "max_dose": 50, "max_total": 2000, "frequency": "Every 6 hours"},
                    "6-12 years (Older child)": {"min_dose": 25, "max_dose": 50, "max_total": 2000, "frequency": "Every 6 hours"},
                    "12+ years (Adolescent)": {"min_dose": 25, "max_dose": 50, "max_total": 2000, "frequency": "Every 6 hours"}
                },
                "unit": "mg",
                "incompatible": ["Aminoglycosides"],
                "interactions": {
                    "Probenecid": "Increased levels",
                    "Allopurinol": "Increased rash risk"
                },
                "oncology_notes": {
                    "penetration": "Poor CSF penetration",
                    "neutropenia": "MSSA coverage",
                    "mucositis": "Skin/soft tissue infections",
                    "renal_adjust": "No adjustment needed",
                    "compatibility": "Monitor INR with warfarin",
                    "monitoring": "Hepatotoxicity rare"
                },
                "notes": "For MSSA infections."
            },

            "CLINDAMYCIN": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 5, "max_dose": 7.5, "max_total": 600, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 5, "max_dose": 10, "max_total": 600, "frequency": "Every 8 hours"},
                    "1-6 years (Young child)": {"min_dose": 5, "max_dose": 10, "max_total": 600, "frequency": "Every 6-8 hours"},
                    "6-12 years (Older child)": {"min_dose": 5, "max_dose": 10, "max_total": 600, "frequency": "Every 6-8 hours"},
                    "12+ years (Adolescent)": {"min_dose": 5, "max_dose": 10, "max_total": 600, "frequency": "Every 6-8 hours"}
                },
                "unit": "mg",
                "incompatible": ["Aminoglycosides"],
                "interactions": {
                    "Neuromuscular blockers": "Enhanced blockade",
                    "Erythromycin": "Antagonism"
                },
                "oncology_notes": {
                    "penetration": "Good bone penetration",
                    "neutropenia": "Anaerobic coverage",
                    "mucositis": "Necrotizing infections",
                    "renal_adjust": "No adjustment needed",
                    "compatibility": "High C.diff risk with chemo",
                    "monitoring": "Watch for diarrhea"
                },
                "notes": "C. diff risk."
            },

            "COLISTIN": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 2.5, "max_dose": 3, "max_total": 300, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 2.5, "max_dose": 5, "max_total": 300, "frequency": "Every 12 hours"},
                    "1-6 years (Young child)": {"min_dose": 2.5, "max_dose": 5, "max_total": 300, "frequency": "Every 12 hours"},
                    "6-12 years (Older child)": {"min_dose": 2.5, "max_dose": 5, "max_total": 300, "frequency": "Every 12 hours"},
                    "12+ years (Adolescent)": {"min_dose": 2.5, "max_dose": 5, "max_total": 300, "frequency": "Every 12 hours"}
                },
                "unit": "mg (CBA)",
                "incompatible": ["None known"],
                "interactions": {
                    "Aminoglycosides": "Increased nephrotoxicity",
                    "Neuromuscular blockers": "Enhanced blockade"
                },
                "oncology_notes": {
                    "penetration": "Limited CSF penetration",
                    "neutropenia": "MDR Gram-negative infections",
                    "mucositis": "No direct activity",
                    "renal_adjust": "Use ideal body weight",
                    "compatibility": "High nephrotoxicity risk",
                    "monitoring": "Daily renal function tests"
                },
                "notes": "For MDR gram-negative infections."
            },

            "GENTAMICIN": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 5, "max_dose": 7.5, "max_total": 400, "frequency": "Every 24 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 5, "max_dose": 7.5, "max_total": 400, "frequency": "Every 24 hours"},
                    "1-6 years (Young child)": {"min_dose": 5, "max_dose": 7.5, "max_total": 400, "frequency": "Every 24 hours"},
                    "6-12 years (Older child)": {"min_dose": 5, "max_dose": 7.5, "max_total": 400, "frequency": "Every 24 hours"},
                    "12+ years (Adolescent)": {"min_dose": 5, "max_dose": 7.5, "max_total": 400, "frequency": "Every 24 hours"}
                },
                "unit": "mg",
                "incompatible": ["Penicillins", "Cephalosporins", "Vancomycin"],
                "interactions": {
                    "Vancomycin": "Synergistic nephrotoxicity",
                    "Loop diuretics": "Ototoxicity",
                    "NSAIDs": "Nephrotoxicity"
                },
                "oncology_notes": {
                    "penetration": "Synergistic with beta-lactams",
                    "neutropenia": "Gram-negative synergy",
                    "mucositis": "No direct activity",
                    "renal_adjust": "Extended interval dosing",
                    "compatibility": "Avoid concurrent nephrotoxins",
                    "monitoring": "Trough <1 mg/L, peak 20-30 mg/L"
                },
                "notes": "Once daily dosing preferred."
            },

            "CEFEPIME": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 30, "max_dose": 50, "max_total": 2000, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 30, "max_dose": 50, "max_total": 2000, "frequency": "Every 8 hours"},
                    "1-6 years (Young child)": {"min_dose": 30, "max_dose": 50, "max_total": 2000, "frequency": "Every 8 hours"},
                    "6-12 years (Older child)": {"min_dose": 30, "max_dose": 50, "max_total": 2000, "frequency": "Every 8 hours"},
                    "12+ years (Adolescent)": {"min_dose": 30, "max_dose": 50, "max_total": 2000, "frequency": "Every 8 hours"}
                },
                "unit": "mg",
                "incompatible": ["Aminoglycosides", "Metronidazole"],
                "interactions": {
                    "Aminoglycosides": "Nephrotoxicity",
                    "Probenecid": "Increased levels"
                },
                "oncology_notes": {
                    "penetration": "Good CSF penetration",
                    "neutropenia": "Antipseudomonal coverage",
                    "mucositis": "Broad Gram-negative spectrum",
                    "renal_adjust": "Adjust for CrCl <60",
                    "compatibility": "Safe with most regimens",
                    "monitoring": "Neurotoxicity risk at high doses"
                },
                "notes": "4th gen cephalosporin."
            },

            "LINEZOLID": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 10, "max_dose": 12, "max_total": 600, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 10, "max_dose": 12, "max_total": 600, "frequency": "Every 8 hours"},
                    "1-6 years (Young child)": {"min_dose": 10, "max_dose": 12, "max_total": 600, "frequency": "Every 8 hours"},
                    "6-12 years (Older child)": {"min_dose": 10, "max_dose": 12, "max_total": 600, "frequency": "Every 8 hours"},
                    "12+ years (Adolescent)": {"min_dose": 10, "max_dose": 12, "max_total": 600, "frequency": "Every 12 hours"}
                },
                "unit": "mg",
                "incompatible": ["None known"],
                "interactions": {
                    "SSRIs": "Serotonin syndrome risk",
                    "MAOIs": "Hypertensive crisis",
                    "Adrenergics": "Increased pressor response"
                },
                "oncology_notes": {
                    "penetration": "Good tissue penetration",
                    "neutropenia": "VRE/MRSA coverage",
                    "mucositis": "Gram-positive coverage",
                    "renal_adjust": "No adjustment needed",
                    "compatibility": "Monitor platelets with chemo",
                    "monitoring": "Weekly CBC for myelosuppression"
                },
                "notes": "Limited course recommended."
            },

            "DAPTOMYCIN": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 4, "max_dose": 6, "max_total": 500, "frequency": "Every 24 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 4, "max_dose": 6, "max_total": 500, "frequency": "Every 24 hours"},
                    "1-6 years (Young child)": {"min_dose": 4, "max_dose": 6, "max_total": 500, "frequency": "Every 24 hours"},
                    "6-12 years (Older child)": {"min_dose": 4, "max_dose": 6, "max_total": 500, "frequency": "Every 24 hours"},
                    "12+ years (Adolescent)": {"min_dose": 4, "max_dose": 6, "max_total": 500, "frequency": "Every 24 hours"}
                },
                "unit": "mg/kg",
                "incompatible": ["None known"],
                "interactions": {
                    "Statins": "Increased rhabdomyolysis risk",
                    "Tobramycin": "Increased CPK"
                },
                "oncology_notes": {
                    "penetration": "Poor CSF penetration",
                    "neutropenia": "MRSA/VRE bacteremia",
                    "mucositis": "No direct activity",
                    "renal_adjust": "Adjust dose for CrCl <30",
                    "compatibility": "Avoid with myeloablative chemo",
                    "monitoring": "Weekly CPK levels"
                },
                "notes": "Not for pulmonary infections."
            },

            "CASPOFUNGIN": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 25, "max_dose": 50, "max_total": 70, "frequency": "Every 24 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 25, "max_dose": 50, "max_total": 70, "frequency": "Every 24 hours"},
                    "1-6 years (Young child)": {"min_dose": 25, "max_dose": 50, "max_total": 70, "frequency": "Every 24 hours"},
                    "6-12 years (Older child)": {"min_dose": 25, "max_dose": 50, "max_total": 70, "frequency": "Every 24 hours"},
                    "12+ years (Adolescent)": {"min_dose": 50, "max_dose": 70, "max_total": 70, "frequency": "Every 24 hours"}
                },
                "unit": "mg/m²",
                "incompatible": ["None known"],
                "interactions": {
                    "Cyclosporine": "Increased caspofungin levels",
                    "Tacrolimus": "Decreased tacrolimus levels"
                },
                "oncology_notes": {
                    "penetration": "Good for invasive candidiasis",
                    "neutropenia": "Aspergillosis treatment",
                    "mucositis": "No direct activity",
                    "renal_adjust": "No adjustment needed",
                    "compatibility": "Monitor with immunosuppressants",
                    "monitoring": "Liver function tests weekly"
                },
                "notes": "Loading dose required."
            },

            "MICAFUNGIN": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 2, "max_dose": 4, "max_total": 150, "frequency": "Every 24 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 2, "max_dose": 4, "max_total": 150, "frequency": "Every 24 hours"},
                    "1-6 years (Young child)": {"min_dose": 2, "max_dose": 4, "max_total": 150, "frequency": "Every 24 hours"},
                    "6-12 years (Older child)": {"min_dose": 2, "max_dose": 4, "max_total": 150, "frequency": "Every 24 hours"},
                    "12+ years (Adolescent)": {"min_dose": 2, "max_dose": 4, "max_total": 150, "frequency": "Every 24 hours"}
                },
                "unit": "mg/kg",
                "incompatible": ["None known"],
                "interactions": {
                    "Sirolimus": "Increased levels",
                    "Nifedipine": "Increased levels"
                },
                "oncology_notes": {
                    "penetration": "Good for candidemia",
                    "neutropenia": "Fungal prophylaxis",
                    "mucositis": "Candida coverage",
                    "renal_adjust": "No adjustment needed",
                    "compatibility": "Safe with TPN",
                    "monitoring": "Monitor liver enzymes"
                },
                "notes": "No loading dose needed."
            },

            "TRIMETHOPRIM-SULFAMETHOXAZOLE": {
                "age_groups": {
                    "0-30 days (Neonate)": {"min_dose": 5, "max_dose": 10, "max_total": 320, "frequency": "Every 12 hours"},
                    "1 month-1 year (Infant)": {"min_dose": 5, "max_dose": 10, "max_total": 320, "frequency": "Every 12 hours"},
                    "1-6 years (Young child)": {"min_dose": 5, "max_dose": 10, "max_total": 320, "frequency": "Every 12 hours"},
                    "6-12 years (Older child)": {"min_dose": 5, "max_dose": 10, "max_total": 320, "frequency": "Every 12 hours"},
                    "12+ years (Adolescent)": {"min_dose": 5, "max_dose": 10, "max_total": 320, "frequency": "Every 12 hours"}
                },
                "unit": "mg/kg (TMP component)",
                "incompatible": ["None known"],
                "interactions": {
                    "Warfarin": "Increased INR",
                    "Phenytoin": "Increased levels",
                    "Methotrexate": "Increased toxicity"
                },
                "oncology_notes": {
                    "penetration": "Good for PJP prophylaxis",
                    "neutropenia": "PJP prevention",
                    "mucositis": "No direct activity",
                    "renal_adjust": "Avoid if CrCl <15",
                    "compatibility": "High risk with methotrexate",
                    "monitoring": "CBC twice weekly"
                },
                "notes": "Myelosuppression risk."
            }
        }
        
        # UI Elements
        ttk.Label(scrollable_frame, text="Pediatric Oncology Antibiotics Calculator", 
                 font=('Helvetica', 16, 'bold')).grid(row=0, column=0, columnspan=6, pady=10)
        
        # Patient details frame
        input_frame = ttk.Frame(scrollable_frame)
        input_frame.grid(row=1, column=0, columnspan=6, pady=10, sticky="ew")
        
        # Weight input
        self.antibiotics_weight_var = tk.StringVar()
        ttk.Label(input_frame, text="Weight (kg):").grid(row=0, column=0, padx=5)
        ttk.Entry(input_frame, textvariable=self.antibiotics_weight_var, width=8).grid(row=0, column=1, padx=5)
        
        # Age group selection
        self.age_group_var = tk.StringVar()
        ttk.Label(input_frame, text="Age Group:").grid(row=0, column=2, padx=5)
        age_combo = ttk.Combobox(input_frame, textvariable=self.age_group_var, 
                                values=self.age_groups, state="readonly", width=25)
        age_combo.grid(row=0, column=3, padx=5)
        
        # Hematological parameters - adjust column positions
        ttk.Label(input_frame, text="ANC (x10⁹/L):").grid(row=1, column=0, padx=5)
        self.anc_var = tk.StringVar()
        ttk.Entry(input_frame, textvariable=self.anc_var, width=8).grid(row=1, column=1, padx=5)
        
        ttk.Label(input_frame, text="CRP (mg/L):").grid(row=1, column=2, padx=5)
        self.crp_var = tk.StringVar()
        ttk.Entry(input_frame, textvariable=self.crp_var, width=8).grid(row=1, column=3, padx=5)
        
        # Renal function inputs - adjust column span
        renal_frame = ttk.Frame(input_frame)
        renal_frame.grid(row=1, column=4, columnspan=2, padx=5)
        ttk.Label(renal_frame, text="Serum Cr (mg/dl):").grid(row=0, column=0)
        self.creatinine_var = tk.StringVar()
        ttk.Entry(renal_frame, textvariable=self.creatinine_var, width=8).grid(row=0, column=1)
        ttk.Label(renal_frame, text="Urine Output:").grid(row=0, column=2)
        self.urine_output_var = tk.StringVar()
        ttk.Combobox(renal_frame, textvariable=self.urine_output_var,
                    values=["Normal", "Oliguric (<1 mL/kg/hr)", "Anuric"], 
                    state="readonly", width=15).grid(row=0, column=3)
        
        # Antibiotic selection frames
        self.antibiotic_frames = []
        for i in range(5):
            abx_frame = ttk.LabelFrame(scrollable_frame, text=f"Antibiotic {i+1}")
            abx_frame.grid(row=2+i, column=0, columnspan=4, sticky="ew", padx=5, pady=5)
            
            # Drug selection
            drug_var = tk.StringVar()
            drug_combo = ttk.Combobox(abx_frame, textvariable=drug_var, 
                                     values=sorted(self.antibiotics_data.keys()), 
                                     state="readonly", width=25)
            drug_combo.grid(row=0, column=0, padx=5, pady=2)
            
            # Dose information
            dose_frame = ttk.Frame(abx_frame)
            dose_frame.grid(row=1, column=0, columnspan=3, sticky="ew")
            
            ttk.Label(dose_frame, text="Dose Range:").grid(row=0, column=0, padx=2)
            dose_range_var = tk.StringVar()
            ttk.Label(dose_frame, textvariable=dose_range_var, width=20).grid(row=0, column=1, padx=2)
            
            ttk.Label(dose_frame, text="Frequency:").grid(row=0, column=2, padx=2)
            frequency_var = tk.StringVar()
            ttk.Label(dose_frame, textvariable=frequency_var, width=15).grid(row=0, column=3, padx=2)
            
            # Dose input
            ttk.Label(abx_frame, text="Enter Dose:").grid(row=2, column=0, padx=5)
            dose_var = tk.StringVar()
            dose_entry = ttk.Entry(abx_frame, textvariable=dose_var, width=8)
            dose_entry.grid(row=2, column=1, padx=5)
            
            # Calculation results
            result_var = tk.StringVar()
            ttk.Label(abx_frame, textvariable=result_var, width=25).grid(row=2, column=2, padx=5)
            
            # Store variables
            self.antibiotic_frames.append({
                "frame": abx_frame,
                "drug_var": drug_var,
                "dose_range_var": dose_range_var,
                "frequency_var": frequency_var,
                "dose_var": dose_var,
                "result_var": result_var
            })
            
            # Bind drug selection to update fields
            drug_var.trace_add("write", lambda *args, i=i: self.update_antibiotic_fields(i))
        
        # Interaction panel
        interaction_frame = ttk.LabelFrame(scrollable_frame, text="Drug Interactions & Warnings")
        interaction_frame.grid(row=2, column=4, rowspan=5, sticky="nsew", padx=10, pady=5)
        
        self.interaction_text = tk.Text(interaction_frame, wrap=tk.WORD, width=50, height=25)
        self.interaction_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Buttons
        btn_frame = ttk.Frame(scrollable_frame)
        btn_frame.grid(row=7, column=0, columnspan=6, pady=10)
        
        ttk.Button(btn_frame, text="Calculate", command=self.calculate_antibiotics).grid(row=0, column=0, padx=5)
        ttk.Button(btn_frame, text="Clear", command=self.clear_antibiotics).grid(row=0, column=1, padx=5)
        
        # Configure grid weights
        scrollable_frame.columnconfigure(4, weight=1)
        scrollable_frame.rowconfigure(6, weight=1)

    def calculate_antibiotics(self):
        """Calculate doses and check interactions"""
        try:
            weight = float(self.antibiotics_weight_var.get())
            age_group = self.age_group_var.get()
            
            for abx in self.antibiotic_frames:
                drug = abx["drug_var"].get()
                if not drug or not age_group:
                    continue
                
                if drug not in self.antibiotics_data:
                    continue
                    
                dose_data = self.antibiotics_data[drug]["age_groups"].get(age_group, {})
                if not dose_data:
                    continue
                
                unit = self.antibiotics_data[drug]["unit"]
                
                try:
                    dose = float(abx["dose_var"].get())
                except ValueError:
                    abx["result_var"].set("Invalid dose input")
                    continue
                
                # Calculate total dose
                calculated_dose = weight * dose
                max_total = dose_data.get("max_total", float('inf'))
                
                if calculated_dose > max_total:
                    calculated_dose = max_total
                    note = f"Max exceeded! Using {max_total}{unit}"
                else:
                    note = f"Within safe range"
                
                # Update results with formatted output
                result_text = (f"Total Dose: {calculated_dose:.1f}{unit}\n"
                              f"Max Allowed: {max_total}{unit}\n"
                              f"Frequency: {dose_data.get('frequency', '')}\n"
                              f"Note: {note}")
                              
                abx["result_var"].set(result_text)
            
            self.check_interactions()
            
        except ValueError as e:
            messagebox.showerror("Error", f"Invalid input: {str(e)}")
        except Exception as e:
            messagebox.showerror("Error", f"Calculation failed: {str(e)}")

    def check_interactions(self):
        """Enhanced interaction check with detailed clinical notes"""
        selected_drugs = [abx["drug_var"].get() for abx in self.antibiotic_frames if abx["drug_var"].get()]
        interactions = []
        incompatibilities = set()
        clinical_notes = set()
        
        # ===== Clinical Notes Structure =====
        clinical_notes = []

        # Neutropenia Section
        try:
            anc = float(self.anc_var.get())
            if anc < 0.5:
                clinical_notes.append("🔴 Severe Neutropenia (ANC <0.5x10⁹/L)")
                clinical_notes.append("  A. Essential Antibiotic Coverage:")
                clinical_notes.append("     - Mandatory: Antipseudomonal β-lactam (Meropenem/Cefepime)")
                clinical_notes.append("     - Add Vancomycin if:")
                clinical_notes.append("       • Skin/soft tissue inflammation at line site")
                clinical_notes.append("       • Pneumonia signs (new infiltrate on CXR)")
                clinical_notes.append("       • Hemodynamic instability (SBP <90mmHg)")
                clinical_notes.append("       • MRSA colonization history")
                clinical_notes.append("       • CRP doubling in 24h with fever")
                clinical_notes.append("  B. Central Line Protocol:")
                clinical_notes.append("     - Draw paired blood cultures (peripheral + central)")
                clinical_notes.append("     - Change dressing using sterile technique")
                clinical_notes.append("     - Remove line if:")
                clinical_notes.append("       • Persistent bacteremia >48h")
                clinical_notes.append("       • S.aureus/Candida infection")
                clinical_notes.append("       • Tunnel infection/purulent exit site")
                clinical_notes.append("  C. Supportive Care:")
                clinical_notes.append("     - Start G-CSF if prolonged neutropenia expected")
                clinical_notes.append("     - Initiate fungal prophylaxis (Micafungin/Posaconazole)")
                clinical_notes.append("     - Daily chlorhexidine bathing")
        except: pass

        # Renal Section
        try:
            creatinine = float(self.creatinine_var.get())
            urine_output = self.urine_output_var.get()
            if creatinine > 2 or urine_output == "Anuric":
                clinical_notes.append("🔴 Renal Management")
                clinical_notes.append("  A. Dose Adjustments:")
                clinical_notes.append("     - Vancomycin: Target trough 10-15 mg/L")
                clinical_notes.append("     - Aminoglycosides: Single daily dose")
                clinical_notes.append("     - Penicillins: Reduce frequency by 50%")
                clinical_notes.append("  B. Avoid Combinations:")
                clinical_notes.append("     - NSAIDs or contrast agents")
                clinical_notes.append("     - Other nephrotoxic antibiotics")
        except: pass

        # CRP Section
        try:
            crp = float(self.crp_var.get())
            if crp > 100:
                clinical_notes.append("🔴 CRP >100 mg/L Actions:")
                clinical_notes.append("     - Escalate antibiotics")
                clinical_notes.append("     - Consider line removal")
                clinical_notes.append("     - Search for occult infection")
            elif crp > 50:
                clinical_notes.append("🟠 CRP >50 mg/L Actions:")
                clinical_notes.append("     - Repeat blood cultures")
                clinical_notes.append("     - Chest/abdominal imaging")
        except: pass

        # ===== Incompatibility Details =====
        detailed_incompatibilities = set()
        for drug in selected_drugs:
            for substance in self.antibiotics_data[drug]["incompatible"]:
                reason = self.get_incompatibility_reason(drug, substance)
                detailed_incompatibilities.add(f"• {substance}: {reason}")

        # ===== Build Structured Report =====
        report = []
        if clinical_notes:
            report.append("📋 CLINICAL GUIDANCE")
            report.extend(clinical_notes)
            report.append("")

        if interactions:
            report.append("🚨 DRUG INTERACTIONS:")
            report.extend(interactions)
            report.append("")

        if detailed_incompatibilities:
            report.append("🚫 INCOMPATIBILITIES (Detailed):")
            report.extend(sorted(detailed_incompatibilities))
            report.append("")

        self.interaction_text.delete(1.0, tk.END)
        self.interaction_text.insert(tk.END, "\n".join(report) if report else "✅ No significant interactions detected")

    def add_supportive_care(self):
        """Add supportive care recommendations"""
        care_recommendations = [
            "\n\n🛡️ Neutropenic Precautions:",
            "- Strict reverse isolation",
            "- Avoid fresh flowers/plants",
            "- Neutropenic diet",
            "\n💉 Access Considerations:",
            "- Change central line if persistent bacteremia",
            "- Consider line lock antibiotics",
            "\n🩺 Monitoring:",
            "- Daily CBC, CRP, cultures",
            "- TDM for vancomycin/aminoglycosides",
            "- Renal function q48h"
        ]
        
        self.interaction_text.insert(tk.END, "\n".join(care_recommendations))
        
    def clear_antibiotics(self):
        """Clear all inputs and results"""
        self.antibiotics_weight_var.set("")
        self.age_group_var.set("")
        self.anc_var.set("")
        self.crp_var.set("")
        self.creatinine_var.set("")
        self.urine_output_var.set("")
        self.interaction_text.delete(1.0, tk.END)
        
        for abx in self.antibiotic_frames:
            abx["drug_var"].set("")
            abx["dose_range_var"].set("")
            abx["frequency_var"].set("")
            abx["dose_var"].set("")
            abx["result_var"].set("")

    def update_antibiotic_fields(self, index):
        """Update fields when antibiotic is selected"""
        abx = self.antibiotic_frames[index]
        drug = abx["drug_var"].get()
        age_group = self.age_group_var.get()
        
        if drug and age_group and drug in self.antibiotics_data:
            try:
                data = self.antibiotics_data[drug]["age_groups"][age_group]
                unit = self.antibiotics_data[drug]["unit"]
                
                abx["dose_range_var"].set(f"{data['min_dose']}-{data['max_dose']} {unit}/kg")
                abx["frequency_var"].set(data["frequency"])
                abx["dose_var"].set(data["max_dose"])
                
            except KeyError:
                abx["dose_range_var"].set("N/A")
                abx["frequency_var"].set("N/A")
                abx["dose_var"].set("")
        else:
            abx["dose_range_var"].set("")
            abx["frequency_var"].set("")
            abx["dose_var"].set("")

    def get_incompatibility_reason(self, drug, substance):
        """Enhanced with drug-specific details"""
        drug_specific = {
            "Ceftriaxone (Rocephine)": {
                "Calcium-containing solutions": "Precipitates in neonates - risk of fatal pulmonary/renal deposits",
                "Vancomycin": "Physical incompatibility in same IV line"
            },
            "Amphotericin B": {
                "Saline solutions": "Forms colloidal aggregates - must use D5W",
                "Other antifungals": "Synergistic toxicity with azoles"
            }
        }
        
        # Check drug-specific reasons first
        if drug in drug_specific and substance in drug_specific[drug]:
            return drug_specific[drug][substance]
            
        # Existing general reasons
        reasons = { 
            # Solution-based incompatibilities
            "Dextrose >5%": "High dextrose concentrations alter pH balance, affecting drug stability",
            "Calcium-containing solutions": "Risk of precipitation (especially with ceftriaxone in neonates)",
            "Divalent cation solutions": "Cations (Ca²⁺, Mg²⁺, Zn²⁺) chelate drug reducing bioavailability",
            "TPN": "Complex parenteral nutrition formulations cause physical instability",
            "Saline solutions": "Osmolarity mismatch leads to precipitation (particularly Amphotericin B)",
            "Aluminum-containing solutions": "Aluminum binds drug molecules forming insoluble complexes",
            "Alkaline solutions": "High pH causes drug degradation through hydrolysis",
            
            # Drug class incompatibilities
            "Aminoglycosides": "Physical precipitation when mixed + synergistic nephrotoxicity risk",
            "Beta-lactams": "Shared beta-lactam structure increases cross-reactivity risk",
            "Penicillins": "Chemical degradation through nucleophilic reactions",
            "Cephalosporins": "Structural similarity leads to competitive inhibition",
            "Other antifungals": "Therapeutic antagonism of fungal coverage",
            "Vancomycin": "Physical incompatibility when co-infused + additive nephrotoxicity",
            
            # Specific drug incompatibilities
            "Heparin": "Polyanionic nature causes complex precipitation",
            "Loop diuretics": "Synergistic ototoxicity risk (e.g., with aminoglycosides)",
            "NSAIDs": "Additive nephrotoxicity through prostaglandin inhibition",
            "Probenecid": "Competitive renal tubular secretion blockade",
            "Warfarin": "Enhanced anticoagulation through protein binding displacement",
            "Phenytoin": "Hepatic enzyme induction alters metabolism",
            
            # Special cases
            "Calcium": "Ceftriaxone-calcium precipitates in biliary/urinary tract (neonates)",
            "Aminoglycosides (extended)": "Inactivation through pH changes when mixed in same line",
            "Vancomycin (extended)": "Nephrotoxicity potentiation through glomerular injury",
            "Cisplatin": "Additive renal tubular toxicity",
            "Methotrexate": "Competitive renal excretion increases toxicity",
            
            # Pharmacodynamic interactions
            "Neuromuscular blockers": "Enhanced paralysis through presynaptic calcium effects",
            "Statins": "Shared myotoxicity pathways increase rhabdomyolysis risk",
            "QT-prolonging agents": "Additive cardiac repolarization delay",
            
            # Pharmacokinetic interactions
            "Rifampin": "Hepatic enzyme induction accelerates drug metabolism",
            "Azoles": "CYP450 inhibition alters drug clearance",
            "Proton Pump Inhibitors": "Gastric pH changes reduce oral bioavailability"
        }
        
        # Try for exact match first
        if substance in reasons:
            return reasons[substance]
        
        # Partial matches for drug classes
        if "beta-lactam" in substance.lower():
            return reasons["Beta-lactams"]
        if "aminoglycoside" in substance.lower():
            return reasons["Aminoglycosides"]
        
        # Fallback to general mechanism
        return "Physical/Chemical incompatibility - avoid simultaneous administration"
    
    def setup_bsa_calculator(self, parent):
        """Setup the BSA calculator interface"""
        frame = ttk.Frame(parent, padding="10 10 10 10", style='TFrame')
        frame.pack(fill=tk.BOTH, expand=True)
        
        # Variables
        self.height_var = tk.StringVar()
        self.weight_var = tk.StringVar()
        self.bsa_result_var = tk.StringVar()
        self.bsa_formula_var = tk.StringVar(value="Mosteller")
        
        # Input fields
        ttk.Label(frame, text="Height (cm):").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        ttk.Entry(frame, textvariable=self.height_var).grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(frame, text="Weight (kg):").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        ttk.Entry(frame, textvariable=self.weight_var).grid(row=1, column=1, padx=5, pady=5)
        
        # Formula selection
        ttk.Label(frame, text="Formula:").grid(row=2, column=0, padx=5, pady=5, sticky="e")
        ttk.Combobox(frame, textvariable=self.bsa_formula_var, 
                    values=["Mosteller", "DuBois"], state="readonly").grid(row=2, column=1, padx=5, pady=5)
        
        # Calculate button
        ttk.Button(frame, text="Calculate BSA", command=self.calculate_bsa,
                  style='Blue.TButton').grid(row=3, column=0, columnspan=2, pady=10)
        
        # Result
        ttk.Label(frame, text="BSA (m²):").grid(row=4, column=0, padx=5, pady=5, sticky="e")
        ttk.Label(frame, textvariable=self.bsa_result_var, font=('Helvetica', 12, 'bold')).grid(row=4, column=1, padx=5, pady=5, sticky="w")

    def calculate_bsa(self):
        """Calculate Body Surface Area using selected formula"""
        try:
            height = float(self.height_var.get())
            weight = float(self.weight_var.get())
            formula = self.bsa_formula_var.get()
            
            if formula == "Mosteller":
                bsa = ((height * weight) / 3600) ** 0.5
            else:  # DuBois
                bsa = 0.007184 * (height ** 0.725) * (weight ** 0.425)
            
            self.bsa_result_var.set(f"{bsa:.4f}")
        except ValueError:
            messagebox.showerror("Error", "Please enter valid numeric values for height and weight")

    def setup_iv_calculator(self, parent):
        """Setup the IV fluid calculator interface"""
        frame = ttk.Frame(parent, padding="10 10 10 10", style='TFrame')
        frame.pack(fill=tk.BOTH, expand=True)
        
        # Variables
        self.iv_bsa_var = tk.StringVar()
        self.iv_weight_var = tk.StringVar()  # New weight variable
        self.iv_fluid_type_var = tk.StringVar(value="4000ml per BSA")
        self.iv_drug_dilution_var = tk.StringVar(value="0")
        self.iv_result_var = tk.StringVar()
        
        # Input fields
        ttk.Label(frame, text="BSA (m²):").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        ttk.Entry(frame, textvariable=self.iv_bsa_var).grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(frame, text="Weight (kg):").grid(row=1, column=0, padx=5, pady=5, sticky="e")  # New weight field
        ttk.Entry(frame, textvariable=self.iv_weight_var).grid(row=1, column=1, padx=5, pady=5)
        
        ttk.Label(frame, text="Fluid Type:").grid(row=2, column=0, padx=5, pady=5, sticky="e")
        ttk.Combobox(frame, textvariable=self.iv_fluid_type_var, 
                    values=["4000ml per BSA", "3000ml per BSA", "2000ml per BSA", 
                           "400ml per BSA", "Full Maintenance", "Half Maintenance"],
                    state="readonly").grid(row=2, column=1, padx=5, pady=5)
        
        ttk.Label(frame, text="Drug Dilution (ml):").grid(row=3, column=0, padx=5, pady=5, sticky="e")
        ttk.Entry(frame, textvariable=self.iv_drug_dilution_var).grid(row=3, column=1, padx=5, pady=5)
        
        # Calculate button
        ttk.Button(frame, text="Calculate IV Rate", command=self.calculate_iv_rate,
                  style='Blue.TButton').grid(row=4, column=0, columnspan=2, pady=10)
        
        # Result
        ttk.Label(frame, text="IV Rate (ml):").grid(row=5, column=0, padx=5, pady=5, sticky="e")
        ttk.Label(frame, textvariable=self.iv_result_var, font=('Helvetica', 12, 'bold')).grid(row=5, column=1, padx=5, pady=5, sticky="w")

    def calculate_iv_rate(self):
        """Calculate IV fluid rate based on BSA/weight and selected protocol"""
        try:
            bsa = float(self.iv_bsa_var.get())
            weight = float(self.iv_weight_var.get())  # Get weight from the new field
            fluid_type = self.iv_fluid_type_var.get()
            drug_dilution = float(self.iv_drug_dilution_var.get() or 0)
            
            if fluid_type == "4000ml per BSA":
                rate = 4000 * bsa - drug_dilution
            elif fluid_type == "3000ml per BSA":
                rate = 3000 * bsa - drug_dilution
            elif fluid_type == "2000ml per BSA":
                rate = 2000 * bsa - drug_dilution
            elif fluid_type == "400ml per BSA":
                rate = 400 * bsa - drug_dilution
            elif fluid_type == "Full Maintenance":
                # Holliday-Segar formula using actual weight
                if weight <= 10:
                    rate = weight * 100 - drug_dilution
                elif weight <= 20:
                    rate = 1000 + (weight - 10) * 50 - drug_dilution
                else:
                    rate = 1500 + (weight - 20) * 20 - drug_dilution
            else:  # Half Maintenance
                if weight <= 10:
                    rate = (weight * 100) / 2 - drug_dilution
                elif weight <= 20:
                    rate = (1000 + (weight - 10) * 50) / 2 - drug_dilution
                else:
                    rate = (1500 + (weight - 20) * 20) / 2 - drug_dilution
            
            self.iv_result_var.set(f"{rate:.2f}")
        except ValueError:
            messagebox.showerror("Error", "Please enter valid numeric values")

    def setup_dosage_calculator(self, parent):
        """Setup the dosage calculator with proper layout and fields"""
        # Create scrollable container
        scrollable_frame = self.create_scrollable_frame(parent)
        
        # Main content frame
        main_frame = ttk.Frame(scrollable_frame, padding="10 10 10 10", style='TFrame')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Create main container with left/right panels
        main_container = ttk.Frame(main_frame)
        main_container.pack(fill=tk.BOTH, expand=True)

        # Left panel - Calculator
        left_panel = ttk.Frame(main_container, padding="5 5 5 5")
        left_panel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Right panel - Drug Info
        right_panel = ttk.Frame(main_container, padding="5 5 5 5")
        right_panel.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        # Chemotherapy drugs data (same as your enhanced version)
        self.chemo_drugs = {
            "Vincristine": {
                "dose": [1.5, 2, "mg/m²"],
                "side_effects": ["Neurotoxicity", "Constipation", "Peripheral neuropathy", "SIADH", "Jaw pain"],
                "compatible_fluids": ["NS", "D5W"],
                "incompatible_fluids": {
                    "Bicarbonate solutions": "pH >7 causes precipitation",
                    "Furosemide": "Forms immediate precipitate",
                    "Ceftriaxone": "Forms insoluble complex",
                    "Heparin": "Physical incompatibility",
                    "Mitomycin": "Mutual inactivation"
                },
                "interactions": {
                    "L-asparaginase": "Increased neurotoxicity (administer vincristine first)",
                    "Phenytoin": "Reduces phenytoin levels by 50%",
                    "Itraconazole": "Increased neurotoxicity risk",
                    "Erythromycin": "Increases vincristine toxicity",
                    "CYP3A4 inhibitors": "Increased toxicity (avoid concurrent use)"
                },
                "photosensitive": False,
                "notes": "Must be administered through central line or free-flowing IV. Administer over 1-2 minutes. Severe vesicant."
            },
            "Dactinomycin": {
                "dose": [0.015, 2.5, "mg/m²"],
                "side_effects": ["Myelosuppression", "Mucositis", "Hepatotoxicity", "Radiation recall", "Extravasation injury"],
                "compatible_fluids": ["NS", "D5W"],
                "incompatible_fluids": {
                    "Dexamethasone": "Chemical degradation",
                    "Vancomycin": "Forms haze",
                    "Heparin": "Precipitate formation"
                },
                "interactions": {
                    "Radiation therapy": "Enhanced radiation recall effect",
                    "Live vaccines": "Avoid for 3 months post-treatment",
                    "Phenobarbital": "Reduced dactinomycin efficacy"
                },
                "photosensitive": False,
                "notes": "Severe vesicant. Administer through central line. Monitor for secondary malignancies."
            },
            "Doxorubicin": {
                "dose": [30, 60, "mg/m²"],
                "side_effects": ["Cardiotoxicity", "Myelosuppression", "Alopecia", "Radiation recall", "Red urine"],
                "compatible_fluids": ["NS"],
                "incompatible_fluids": {
                    "Heparin": "Forms precipitate",
                    "5-FU": "Chemical degradation",
                    "Alkaline solutions": "pH instability causes degradation",
                    "Dexamethasone": "Forms precipitate",
                    "Ceftriaxone": "Incompatible at Y-site"
                },
                "interactions": {
                    "Cyclophosphamide": "Increased cardiotoxicity risk",
                    "Verapamil": "Increased intracellular concentration",
                    "Trastuzumab": "Severe cardiac dysfunction risk",
                    "Paclitaxel": "Altered doxorubicin clearance"
                },
                "photosensitive": False,
                "notes": "Cardiac monitoring required (MUGA/ECHO). Cumulative lifetime dose limit 550mg/m². Severe vesicant."
            },
            "Cyclophosphamide": {
                "dose": [1000, 2000, "mg/m²"],
                "side_effects": ["Hemorrhagic cystitis", "Myelosuppression", "Nausea", "SIADH", "Cardiotoxicity"],
                "compatible_fluids": ["NS", "D5W"],
                "incompatible_fluids": {
                    "Amphotericin B": "Increased nephrotoxicity",
                    "Chloramphenicol": "Reduced efficacy"
                },
                "interactions": {
                    "Allopurinol": "Increased myelosuppression",
                    "Succinylcholine": "Prolonged neuromuscular blockade",
                    "Warfarin": "Increased anticoagulant effect",
                    "Digoxin": "Reduced absorption"
                },
                "photosensitive": False,
                "notes": "Mesna protection required for doses >1000mg/m². Aggressive hydration (3000mL/m²/day)."
            },
            "Methotrexate": {
                "dose": [200, 15000, "mg/m²"],
                "side_effects": ["Myelosuppression", "Mucositis", "Nephrotoxicity", "Hepatotoxicity", "Neurotoxicity"],
                "compatible_fluids": ["D5W"],
                "incompatible_fluids": {
                    "NS": "Increased precipitation risk in chloride solutions",
                    "Lactated Ringer's": "Calcium forms complex with MTX",
                    "Proton pump inhibitors": "Reduced renal clearance",
                    "Penicillins": "Increased MTX toxicity",
                    "NSAIDs": "Competitive tubular secretion"
                },
                "interactions": {
                    "NSAIDs": "Increased MTX toxicity (avoid)",
                    "Probenecid": "Reduced renal excretion",
                    "Sulfonamides": "Increased myelosuppression",
                    "Theophylline": "Increased theophylline levels"
                },
                "photosensitive": True,
                "notes": "Leucovorin rescue required for high-dose therapy. Urine pH monitoring (maintain >7)."
            },
            "Cisplatin": {
                "dose": [100, 120, "mg/m²"],
                "side_effects": ["Nephrotoxicity", "Ototoxicity", "Neuropathy", "Hypomagnesemia", "Anaphylaxis"],
                "compatible_fluids": ["NS"],
                "incompatible_fluids": {
                    "D5W": "Decreased stability in dextrose",
                    "Aluminum-containing solutions": "Forms black precipitate",
                    "Bicarbonate": "Forms precipitate",
                    "Paclitaxel": "Incompatible at Y-site"
                },
                "interactions": {
                    "Aminoglycosides": "Increased nephrotoxicity (avoid)",
                    "Phenytoin": "Reduced phenytoin absorption",
                    "Diuretics": "Increased ototoxicity risk",
                    "Live vaccines": "Avoid for 3 months"
                },
                "photosensitive": False,
                "notes": "Aggressive hydration required (3L/day). Magnesium supplementation often needed. Administer over 2-6 hours."
            },
            "Carboplatin": {
                "dose": [600, 800, "mg/m²"],
                "side_effects": ["Myelosuppression", "Nausea", "Peripheral neuropathy", "Ototoxicity", "Hypersensitivity"],
                "compatible_fluids": ["D5W", "NS"],
                "incompatible_fluids": {
                    "Aluminum-containing solutions": "Forms precipitate",
                    "Amifostine": "Chemical interaction"
                },
                "interactions": {
                    "Nephrotoxic drugs": "Additive renal toxicity",
                    "Live vaccines": "Avoid concurrent use",
                    "Phenytoin": "Reduced levels"
                },
                "photosensitive": False,
                "notes": "Dose based on AUC calculation preferred. Administer over 15-60 minutes."
            },
            "Etoposide": {
                "dose": [100, 500, "mg/m²"],
                "side_effects": ["Myelosuppression", "Hypotension", "Alopecia", "Anaphylaxis", "Secondary malignancies"],
                "compatible_fluids": ["NS", "D5W"],
                "incompatible_fluids": {
                    "Heparin": "Forms precipitate",
                    "Cefepime": "Physical incompatibility",
                    "Filgrastim": "Incompatible at Y-site"
                },
                "interactions": {
                    "Warfarin": "Increased anticoagulant effect",
                    "Cyclosporine": "Increased etoposide levels",
                    "St. John's Wort": "Reduced efficacy"
                },
                "photosensitive": True,
                "notes": "Administer over 30-60 minutes to prevent hypotension. Dilute to <0.4mg/mL concentration."
            },
            "Ifosfamide": {
                "dose": [1800, 3000, "mg/m²"],
                "side_effects": ["Hemorrhagic cystitis", "Neurotoxicity", "Nephrotoxicity", "SIADH", "Cardiotoxicity"],
                "compatible_fluids": ["D5W", "NS"],
                "incompatible_fluids": {
                    "Mesna": "Chemical interaction if mixed directly",
                    "Cisplatin": "Increased neurotoxicity"
                },
                "interactions": {
                    "CNS depressants": "Increased neurotoxicity risk",
                    "Warfarin": "Increased anticoagulant effect",
                    "Phenobarbital": "Increased metabolism"
                },
                "photosensitive": False,
                "notes": "Mesna protection required. Administer over 2-24 hours. Aggressive hydration (3L/m²/day)."
            },
            "Cytarabine": {
                "dose": [100, 3000, "mg/m²"],
                "side_effects": ["Myelosuppression", "Cerebellar toxicity", "Conjunctivitis", "Ara-C syndrome", "Anaphylaxis"],
                "compatible_fluids": ["D5W", "NS"],
                "incompatible_fluids": {
                    "Gentamicin": "Mutual inactivation",
                    "5-FU": "Antagonistic effect",
                    "Heparin": "Forms precipitate"
                },
                "interactions": {
                    "Digoxin": "Reduced digoxin absorption",
                    "Live vaccines": "Avoid concurrent use",
                    "Flucytosine": "Reduced efficacy"
                },
                "photosensitive": False,
                "notes": "High-dose regimen requires steroid eye drops prophylaxis. Administer over 1-3 hours."
            },
            "Daunorubicin": {
                "dose": [25, 45, "mg/m²"],
                "side_effects": ["Cardiotoxicity", "Myelosuppression", "Red urine", "Mucositis", "Radiation recall"],
                "compatible_fluids": ["D5W", "NS"],
                "incompatible_fluids": {
                    "Heparin": "Forms precipitate",
                    "Alkaline solutions": "Causes degradation",
                    "Dexamethasone": "Physical incompatibility"
                },
                "interactions": {
                    "Cyclophosphamide": "Increased myocardial damage",
                    "Trastuzumab": "Synergistic cardiotoxicity",
                    "CYP3A4 inhibitors": "Increased serum levels"
                },
                "photosensitive": False,
                "notes": "Cumulative lifetime dose limit 550mg/m². Cardiac monitoring required. Severe vesicant."
            },
            "Asparaginase": {
                "dose": [5000, 10000, "IU/m²"],
                "side_effects": ["Allergic reactions", "Pancreatitis", "Hyperglycemia", "Coagulopathy", "Hepatotoxicity"],
                "compatible_fluids": ["NS"],
                "incompatible_fluids": {
                    "Dextrose solutions": "Reduced stability",
                    "Dexamethasone": "Forms haze",
                    "Vancomycin": "Physical incompatibility"
                },
                "interactions": {
                    "Anticoagulants": "Increased bleeding risk",
                    "Methotrexate": "Reduced efficacy",
                    "Vaccines": "Avoid live vaccines"
                },
                "photosensitive": True,
                "notes": "Requires test dose for hypersensitivity. Monitor amylase/lipase. Administer over 1-2 hours."
            },
            "Melphalan": {
                "dose": [10, 20, "mg/m²"],
                "side_effects": ["Myelosuppression", "Nausea/vomiting", "Secondary malignancies", "Pulmonary fibrosis"],
                "compatible_fluids": ["NS"],
                "incompatible_fluids": {
                    "D5W": "Reduced stability",
                    "Ciprofloxacin": "Chemical interaction"
                },
                "interactions": {
                    "Cyclosporine": "Increased nephrotoxicity",
                    "Nalidixic acid": "Severe hemorrhagic colitis",
                    "Live vaccines": "Contraindicated"
                },
                "photosensitive": False,
                "notes": "Administer via central line. Premedicate with antiemetics. Handle with cytotoxic precautions."
            },
            "Thioguanine (6-TG)": {
                "dose": [75, 100, "mg/m²"],
                "side_effects": ["Hepatotoxicity", "Myelosuppression", "Hyperuricemia", "Mucositis"],
                "compatible_fluids": ["D5W"],
                "incompatible_fluids": {
                    "Allopurinol": "Increased toxicity risk",
                    "Cytarabine": "Synergistic toxicity"
                },
                "interactions": {
                    "Allopurinol": "Requires dose reduction",
                    "Warfarin": "Increased anticoagulation",
                    "Live vaccines": "Avoid concurrent use"
                },
                "photosensitive": True,
                "notes": "Monitor liver function tests. Usually administered orally. Adjust dose for TPMT deficiency."
            },
            "Bleomycin": {
                "dose": [10, 20, "units/m²"],
                "side_effects": ["Pulmonary fibrosis", "Fever/chills", "Skin toxicity", "Hypersensitivity"],
                "compatible_fluids": ["NS", "D5W"],
                "incompatible_fluids": {
                    "Aminophylline": "Forms precipitate",
                    "Dexamethasone": "Physical incompatibility"
                },
                "interactions": {
                    "Oxygen therapy": "Increased pulmonary toxicity",
                    "Cisplatin": "Synergistic toxicity",
                    "Radiation": "Enhanced skin reactions"
                },
                "photosensitive": False,
                "notes": "Cumulative lifetime dose limit 400 units. Monitor pulmonary function tests. Premedicate with steroids."
            },
            "Dacarbazine": {
                    "dose": [150, 1200, "mg/m²"],
                    "side_effects": ["Myelosuppression", "Severe nausea/vomiting", "Hepatotoxicity", "Photosensitivity", "Flu-like syndrome"],
                    "compatible_fluids": ["D5W", "NS"],
                    "incompatible_fluids": {
                            "Heparin": "Forms precipitate",
                            "Hydrocortisone": "Physical incompatibility",
                            "Cefepime": "Y-site incompatibility"
                    },
                    "interactions": {
                            "Phenytoin": "Reduces antiepileptic efficacy",
                            "Live vaccines": "Avoid during treatment",
                            "Allopurinol": "Increased myelosuppression risk"
                    },
                    "photosensitive": True,
                    "notes": "Administer IV over 15-30 minutes. Requires light protection during infusion. High emetic risk - use 5-HT3 antagonists."
            },
            "Vinblastine": {
                    "dose": [3.7, 11.1, "mg/m²"],
                    "side_effects": ["Myelosuppression", "Neurotoxicity", "Constipation", "Jaw pain", "SIADH"],
                    "compatible_fluids": ["NS", "D5W"],
                    "incompatible_fluids": {
                            "Furosemide": "Immediate precipitate",
                            "Ceftriaxone": "Forms insoluble complex",
                            "Bicarbonate solutions": "pH instability"
                    },
                    "interactions": {
                            "CYP3A4 inhibitors": "Increased neurotoxicity risk",
                            "Mitomycin": "Increased pulmonary reactions",
                            "Phenytoin": "Reduced serum levels"
                    },
                    "photosensitive": False,
                    "notes": "Vesicant. Dose reduce for hepatic dysfunction. Administer through central line over 1 min. Monitor neurologic status."
            },
            "Mesna": {
                    "dose": [1080, 3000, "mg/m²" "  (60 - 100 %) of ifosfamide dose"],
                    "side_effects": ["Nausea", "Headache", "Hypersensitivity", "Hematuria (if underdosed)"],
                    "compatible_fluids": ["D5W", "NS"],
                    "incompatible_fluids": {
                            "Cisplatin": "Incompatible in same line",
                            "Amphotericin B": "Forms precipitate"
                    },
                    "interactions": {
                            "Ifosfamide/Cyclophosphamide": "Prevents hemorrhagic cystitis",
                            "Cephalothin": "False positive ketonuria"
                    },
                    "photosensitive": False,
                    "notes": "Administer before and 4/8hrs after ifosfamide. IV or oral. Monitor urine for blood. Not protective against other toxicities."
            }
        }

        # Initialize variables
        self.drug_var = tk.StringVar()
        self.dose_var = tk.StringVar()
        self.max_dose_var = tk.StringVar()
        self.patient_weight_var = tk.StringVar()
        self.bsa_var = tk.StringVar()
        self.dosage_result_var = tk.StringVar()
        self.volume_result_var = tk.StringVar()
        self.amount_var = tk.StringVar()
        self.volume_var = tk.StringVar()
        self.patient_name_var = tk.StringVar()
        self.dob_var = tk.StringVar()
        self.diagnosis_var = tk.StringVar()
        self.protocol_var = tk.StringVar()
        self.cycle_var = tk.StringVar()
        self.day_var = tk.StringVar()
        
        # New variables for infusion parameters
        self.dilution_fluid_var = tk.StringVar()
        self.total_volume_var = tk.StringVar()
        self.infusion_duration_var = tk.StringVar()
        self.infusion_rate_var = tk.StringVar()

        # Left Panel Layout
        row = 0

        # Patient Information Section
        ttk.Label(left_panel, text="Patient Information", 
                  font=('Helvetica', 12, 'bold')).grid(row=row, column=0, columnspan=2, pady=5)
        row += 1

        ttk.Label(left_panel, text="Name:").grid(row=row, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(left_panel, textvariable=self.patient_name_var).grid(row=row, column=1, padx=5, pady=2, sticky="w")
        row += 1

        ttk.Label(left_panel, text="DOB (dd/mm/yyyy):").grid(row=row, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(left_panel, textvariable=self.dob_var).grid(row=row, column=1, padx=5, pady=2, sticky="w")
        row += 1

        ttk.Label(left_panel, text="Diagnosis:").grid(row=row, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(left_panel, textvariable=self.diagnosis_var).grid(row=row, column=1, padx=5, pady=2, sticky="w")
        row += 1

        ttk.Label(left_panel, text="Protocol:").grid(row=row, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(left_panel, textvariable=self.protocol_var).grid(row=row, column=1, padx=5, pady=2, sticky="w")
        row += 1

        ttk.Label(left_panel, text="Cycle:").grid(row=row, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(left_panel, textvariable=self.cycle_var).grid(row=row, column=1, padx=5, pady=2, sticky="w")
        row += 1

        ttk.Label(left_panel, text="Day:").grid(row=row, column=0, padx=5, pady=2, sticky="e")
        ttk.Entry(left_panel, textvariable=self.day_var).grid(row=row, column=1, padx=5, pady=2, sticky="w")
        row += 1

        # Treatment Parameters Section
        ttk.Label(left_panel, text="Treatment Parameters", 
                  font=('Helvetica', 12, 'bold')).grid(row=row, column=0, columnspan=2, pady=10)
        row += 1

        ttk.Label(left_panel, text="Drug:").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        drug_combo = ttk.Combobox(left_panel, textvariable=self.drug_var, 
                                  values=list(self.chemo_drugs.keys()), state="readonly")
        drug_combo.grid(row=row, column=1, padx=5, pady=5, sticky="w")
        drug_combo.bind("<<ComboboxSelected>>", self.update_drug_info)
        row += 1

        ttk.Label(left_panel, text="Dose per m²:").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        ttk.Entry(left_panel, textvariable=self.dose_var).grid(row=row, column=1, padx=5, pady=5, sticky="w")
        row += 1

        ttk.Label(left_panel, text="Max Dose:").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        ttk.Entry(left_panel, textvariable=self.max_dose_var).grid(row=row, column=1, padx=5, pady=5, sticky="w")
        row += 1

        ttk.Label(left_panel, text="Weight (kg):").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        ttk.Entry(left_panel, textvariable=self.patient_weight_var).grid(row=row, column=1, padx=5, pady=5, sticky="w")
        row += 1

        ttk.Label(left_panel, text="BSA (m²):").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        ttk.Entry(left_panel, textvariable=self.bsa_var).grid(row=row, column=1, padx=5, pady=5, sticky="w")
        row += 1

        ttk.Label(left_panel, text="Vial Amount (mg/IU):").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        ttk.Entry(left_panel, textvariable=self.amount_var).grid(row=row, column=1, padx=5, pady=5, sticky="w")
        row += 1

        ttk.Label(left_panel, text="Diluent Volume (mL):").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        ttk.Entry(left_panel, textvariable=self.volume_var).grid(row=row, column=1, padx=5, pady=5, sticky="w")
        row += 1

        # New fields for dilution and infusion parameters
        ttk.Label(left_panel, text="Type of Dilution Fluid:").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        fluid_combo = ttk.Combobox(left_panel, textvariable=self.dilution_fluid_var, 
                                 values=["NS", "D5W", "Sterile Water", "D5NS"], state="readonly")
        fluid_combo.grid(row=row, column=1, padx=5, pady=5, sticky="w")
        row += 1

        ttk.Label(left_panel, text="Total Infused Volume (mL):").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        ttk.Entry(left_panel, textvariable=self.total_volume_var).grid(row=row, column=1, padx=5, pady=5, sticky="w")
        row += 1

        ttk.Label(left_panel, text="Infusion Duration (hr):").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        ttk.Entry(left_panel, textvariable=self.infusion_duration_var).grid(row=row, column=1, padx=5, pady=5, sticky="w")
        row += 1

        ttk.Label(left_panel, text="Infusion Rate (mL/hr):").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        ttk.Label(left_panel, textvariable=self.infusion_rate_var, 
                font=('Helvetica', 10, 'bold')).grid(row=row, column=1, padx=5, pady=5, sticky="w")
        row += 1

        # Set up trace for automatic rate calculation
        self.total_volume_var.trace_add("write", self.calculate_infusion_rate)
        self.infusion_duration_var.trace_add("write", self.calculate_infusion_rate)

        # Calculate Button
        ttk.Button(left_panel, text="Calculate Dosage", command=self.calculate_dosage,
                   style='Blue.TButton').grid(row=row, column=0, columnspan=2, pady=15)
        row += 1

        # Results Section
        ttk.Label(left_panel, text="Calculated Dose:").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        ttk.Label(left_panel, textvariable=self.dosage_result_var, 
                  font=('Helvetica', 10, 'bold')).grid(row=row, column=1, padx=5, pady=5, sticky="w")
        row += 1

        ttk.Label(left_panel, text="Administer Volume:").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        ttk.Label(left_panel, textvariable=self.volume_result_var, 
                  font=('Helvetica', 10, 'bold')).grid(row=row, column=1, padx=5, pady=5, sticky="w")
        row += 1

        # Export Button
        ttk.Button(left_panel, text="Export to Word", command=self.export_dosage_to_word,
                   style='Green.TButton').grid(row=row, column=0, columnspan=2, pady=15)

        # Right Panel Drug Info
        info_label = ttk.Label(right_panel, text="Drug Information", 
                               font=('Helvetica', 14, 'bold'))
        info_label.pack(pady=5)
        
        self.info_text = tk.Text(right_panel, wrap=tk.WORD, width=70, height=25, 
                                 padx=10, pady=10, font=('Helvetica', 10), 
                                 bg='white', fg='black')
        scrollbar = ttk.Scrollbar(right_panel, command=self.info_text.yview)
        self.info_text.configure(yscrollcommand=scrollbar.set)
        
        self.info_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.update_drug_info()

        # Configure grid weights for left panel after creating all widgets
        for i in range(row):
            left_panel.grid_rowconfigure(i, weight=0)
        left_panel.grid_rowconfigure(row-1, weight=1)  # Let last row expand

        # Configure grid weights for main container
        main_container.columnconfigure(0, weight=1)
        main_container.columnconfigure(1, weight=1)
        main_container.rowconfigure(0, weight=1)

    def calculate_infusion_rate(self, *args):
        """Calculate infusion rate automatically when volume/duration change"""
        try:
            volume = float(self.total_volume_var.get())
            duration = float(self.infusion_duration_var.get())
            rate = volume / duration
            self.infusion_rate_var.set(f"{rate:.1f}")
        except (ValueError, ZeroDivisionError):
            self.infusion_rate_var.set("")

    def create_scrollable_frame(self, parent):
        """Create a scrollable frame with improved scroll region handling"""
        container = ttk.Frame(parent)
        canvas = tk.Canvas(container, highlightthickness=0)
        scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        # Configure canvas scrolling
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all")
            )
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Pack elements
        container.pack(fill="both", expand=True)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Add cross-platform mousewheel support
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        canvas.bind("<Enter>", lambda _: canvas.bind_all('<MouseWheel>', _on_mousewheel))
        canvas.bind("<Leave>", lambda _: canvas.unbind_all('<MouseWheel>'))

        return scrollable_frame
    
    def update_drug_info(self, event=None):
        """Update all drug information fields when drug is selected"""
        drug = self.drug_var.get()
        self.info_text.config(state=tk.NORMAL)
        self.info_text.delete(1.0, tk.END)
        
        if not drug or drug not in self.chemo_drugs:
            self.info_text.insert(tk.END, "Please select a chemotherapy drug from the dropdown to view detailed information.")
            self.info_text.config(state=tk.DISABLED)
            return
            
        data = self.chemo_drugs[drug]
        
        # Update dosage fields
        dose, max_dose, unit = data["dose"]
        self.dose_var.set(str(dose))
        self.max_dose_var.set(str(max_dose))
        
        # Format drug information with rich formatting
        self.info_text.tag_configure("title", font=('Helvetica', 14, 'bold'))
        self.info_text.tag_configure("header", font=('Helvetica', 12, 'bold'))
        self.info_text.tag_configure("bold", font=('Helvetica', 10, 'bold'))
        self.info_text.tag_configure("warning", foreground="red")
        
        # Extract the base unit without the "/m²" part
        base_unit = unit.split("/")[0]
        
        self.info_text.insert(tk.END, f"{drug.upper()}\n", "title")
        self.info_text.insert(tk.END, "\nSTANDARD DOSING\n", "header")
        self.info_text.insert(tk.END, f"Dose: {dose} {unit}\n", "bold")
        self.info_text.insert(tk.END, f"Max Dose: {max_dose} {base_unit}\n\n")
        
        self.info_text.insert(tk.END, "SIDE EFFECTS\n", "header")
        for effect in data["side_effects"]:
            self.info_text.insert(tk.END, f"• {effect}\n")
        
        self.info_text.insert(tk.END, "\nCOMPATIBLE FLUIDS\n", "header")
        for fluid in data["compatible_fluids"]:
            self.info_text.insert(tk.END, f"• {fluid}\n")
        
        self.info_text.insert(tk.END, "\nINCOMPATIBILITIES\n", "header")
        if data["incompatible_fluids"]:
            for fluid, reason in data["incompatible_fluids"].items():
                self.info_text.insert(tk.END, f"• {fluid}: ", "bold")
                self.info_text.insert(tk.END, f"{reason}\n")
        else:
            self.info_text.insert(tk.END, "• No significant incompatibilities\n")
        
        self.info_text.insert(tk.END, "\nDRUG INTERACTIONS\n", "header")
        if data["interactions"]:
            for drug_name, interaction in data["interactions"].items():
                self.info_text.insert(tk.END, f"• {drug_name}: ", "bold")
                self.info_text.insert(tk.END, f"{interaction}\n")
        else:
            self.info_text.insert(tk.END, "• No major interactions\n")
        
        self.info_text.insert(tk.END, "\nPHOTOSENSITIVITY\n", "header")
        self.info_text.insert(tk.END, f"{'Yes' if data['photosensitive'] else 'No'}\n")
    
        self.info_text.insert(tk.END, "\nSPECIAL NOTES\n", "header")
        self.info_text.insert(tk.END, f"{data['notes']}\n", "bold")
    
        self.info_text.config(state=tk.DISABLED)

    def calculate_dosage(self):
        """Calculate chemotherapy dosage with pediatric adjustments"""
        try:
            drug = self.drug_var.get()
            dose_per_kg = float(self.dose_var.get())
            max_dose = float(self.max_dose_var.get())
            weight = float(self.patient_weight_var.get())
            bsa = float(self.bsa_var.get())
            
            if drug not in self.chemo_drugs:
                messagebox.showerror("Error", "Please select a valid drug")
                return

            unit = self.chemo_drugs[drug]["dose"][2]
            base_unit = unit.split("/")[0]
            notes = []
            
            # BSA capping
            bsa_used = min(bsa, 1.7)
            if bsa > 1.7:
                notes.append("BSA capped at 1.7 m²")
            
            # Pediatric dosing logic
            if weight <= 12:
                calculated_dose = (weight * dose_per_kg) / 30
                notes.append("Weight ≤12 kg: dose calculated using (weight × dose) / 30")
            else:
                if unit.endswith("/m²"):
                    calculated_dose = dose_per_kg * bsa_used
                else:
                    calculated_dose = dose_per_kg * weight
            
            final_dose = min(calculated_dose, max_dose)
            
            # Volume calculation
            amount = self.amount_var.get()
            volume = self.volume_var.get()
            vol_warning = ""
            
            if amount and volume:
                try:
                    amount_val = float(amount)
                    volume_val = float(volume)
                    
                    if volume_val <= 0:
                        raise ValueError("Invalid volume")
                        
                    concentration = amount_val / volume_val
                    final_volume = final_dose / concentration
                    
                    # Volume formatting
                    if final_volume < 0.1:
                        vol_warning = "\n(WARNING: Volume <0.1 mL - consider dose adjustment)"
                        vol_str = f"{final_volume:.3f}"
                    elif final_volume < 1:
                        vol_str = f"{final_volume:.2f}"
                    else:
                        vol_str = f"{final_volume:.1f}"
                        
                    self.volume_result_var.set(f"{vol_str} mL {vol_warning}")
                except Exception as e:
                    self.volume_result_var.set(f"Error: {str(e)}")
            else:
                self.volume_result_var.set("Enter both amount and volume")
            
            # Store values for export
            self.export_final_dose = final_dose
            self.export_base_unit = base_unit

            # Clinical notes (GUI only)
            clinical_notes = [
                "Down syndrome: consider 15-25% dose reduction",
                "Adjust for prior toxicity/reactions"
            ]

            # Display in calculator
            result_text = f"{final_dose:.2f} {base_unit}"
            if notes or clinical_notes:
                all_notes = notes + clinical_notes
                result_text += "\n\nNOTES:\n• " + "\n• ".join(all_notes)
            
            self.dosage_result_var.set(result_text)

        except ValueError as ve:
            messagebox.showerror("Input Error", f"Invalid numeric value: {str(ve)}")
        except Exception as e:
            messagebox.showerror("Calculation Error", f"An error occurred: {str(e)}")

    def export_dosage_to_word(self):
        """Generate professional single-page chemotherapy order document"""
        try:
            from datetime import datetime
            import tempfile
            import os
            import subprocess
            import sys
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_ALIGN_VERTICAL

            # Create temporary file
            temp_dir = tempfile.mkdtemp()
            filename = os.path.join(temp_dir, "Chemo_Order_Final.docx")
            
            doc = Document()
            
            # Configure default styles
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(11)

            # Header Section
            header = doc.add_paragraph()
            header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_run = header.add_run("TRIPOLI CHILDREN TEACHING HOSPITAL\nONCOLOGY DEPARTMENT")
            header_run.font.name = 'Arial'
            header_run.font.size = Pt(14)
            header_run.bold = True
            doc.add_paragraph().add_run().add_break()

            # Document Title
            title = doc.add_paragraph()
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title.add_run("CHEMOTHERAPY ADMINISTRATION ORDER")
            title_run.font.size = Pt(12)
            title_run.bold = True
            title_run.underline = True

            # Generation Info
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}", 
                            style='Intense Quote').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Patient Information Table (Wider layout)
            doc.add_paragraph("\nPatient Information:", style='Heading 2')
            patient_data = [
                ["Patient Name:", self.patient_name_var.get(), "Date of Birth:", self.dob_var.get()],
                ["Weight:", f"{self.patient_weight_var.get()} kg", "BSA:", f"{self.bsa_var.get()} m²"],
                ["Diagnosis:", self.diagnosis_var.get(), "Protocol:", self.protocol_var.get()],
                ["Cycle:", self.cycle_var.get(), "Day:", self.day_var.get(), "", ""]
            ]
            
            table = doc.add_table(rows=0, cols=4)
            table.style = 'Light Shading'
            for row in patient_data:
                cells = table.add_row().cells
                for i in range(4):
                    cells[i].text = row[i] if i < len(row) else ""
                    if i % 2 == 0:  # Make labels bold
                        cells[i].paragraphs[0].runs[0].font.bold = True
                    cells[i].width = Inches(1.8)

            # Treatment Details Section (Expanded)
            doc.add_paragraph("\nTreatment Parameters:", style='Heading 2')
            treatment_data = [
                ["Drug Name:", self.drug_var.get()],
                ["Calculated Dose:", f"{self.export_final_dose:.2f} {self.export_base_unit}"],
                ["Max Dose:", f"{self.chemo_drugs[self.drug_var.get()]['dose'][1]} {self.export_base_unit}"],
                ["Vial Strength:", f"{self.amount_var.get()} {self.export_base_unit}"],
                ["Diluent Volume:", f"{self.volume_var.get()} mL"],
                ["Admin. Volume:", self.volume_result_var.get().split('\n')[0]]
            ]
            
            treatment_table = doc.add_table(rows=0, cols=2)
            treatment_table.style = 'Light Shading'
            for row in treatment_data:
                cells = treatment_table.add_row().cells
                cells[0].text = row[0]
                cells[1].text = row[1]
                cells[0].paragraphs[0].runs[0].font.bold = True
                cells[0].width = Inches(2.2)
                cells[1].width = Inches(3.2)

            # Drug Specifications Table (Revised)
            if self.drug_var.get() in self.chemo_drugs:
                drug = self.drug_var.get()
                data = self.chemo_drugs[drug]
                
                doc.add_paragraph("\nDrug Specifications:", style='Heading 2')
                spec_table = doc.add_table(rows=4, cols=2)
                spec_table.style = 'Light Grid'
                
                # Header row
                header_cells = spec_table.rows[0].cells
                header_cells[0].text = "Parameter"
                header_cells[1].text = "Details"
                for cell in header_cells:
                    cell.paragraphs[0].runs[0].bold = True
                
                # Data rows
                spec_data = [
                    ("Incompatibilities", "\n".join(
                        [f"• {k} ({v})" for k,v in list(data['incompatible_fluids'].items())[:3]]
                        if data['incompatible_fluids'] else "• None")),
                    ("Drug Interactions", "\n".join(
                        [f"• {k}: {v}" for k,v in list(data['interactions'].items())[:2]]
                        if data['interactions'] else "• None")),
                    ("Photosensitivity", "Yes" if data['photosensitive'] else "No"),
                    ("Administration Notes", data['notes'][:120] + "...")
                ]
                
                for i in range(1, 4):
                    cells = spec_table.rows[i].cells
                    cells[0].text = spec_data[i-1][0]
                    cells[1].text = spec_data[i-1][1]
                    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Compact Verification Section
            doc.add_paragraph("\nAdministration Details:", style='Heading 2')
            
            # Infusion parameters table (modified)
            infusion_table = doc.add_table(rows=1, cols=3)
            infusion_table.style = 'Light Shading'
            infusion_table.autofit = False
            
            # Set column widths
            for cell in infusion_table.columns[0].cells:
                cell.width = Inches(2)
            for cell in infusion_table.columns[1].cells:
                cell.width = Inches(2.5)
            for cell in infusion_table.columns[2].cells:
                cell.width = Inches(2.5)
            
            # Header
            header_cells = infusion_table.rows[0].cells
            header_cells[0].text = "Parameter"
            header_cells[1].text = "Calculation"
            header_cells[2].text = "Verification"
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True

            # Add data row with actual values
            row_cells = infusion_table.add_row().cells
            row_cells[0].text = ("Infusion Details\n"
                               f"Diluent: {self.dilution_fluid_var.get()}\n"
                               f"Total Volume: {self.total_volume_var.get()} mL")
            row_cells[1].text = (f"Infusion Duration: {self.infusion_duration_var.get()} hr\n"
                                f"Infusion Rate: {self.infusion_rate_var.get()} mL/hr")
            row_cells[2].text = "Verified By:\n\nSignature: ________________\nDate: __/__/____"
            
            # Format infusion table
            for cell in row_cells:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Enhanced Verification Section
            doc.add_paragraph("\nFinal Verification:", style='Heading 2')
            verification = doc.add_table(rows=2, cols=2)
            verification.style = 'Light Shading'
            
            # Set column widths
            verification.columns[0].width = Inches(3)
            verification.columns[1].width = Inches(3)
            
            # Header cells
            verification.cell(0,0).text = "PHARMACIST VERIFICATION"
            verification.cell(0,1).text = "NURSING VERIFICATION"
            for cell in verification.rows[0].cells:
                cell.paragraphs[0].runs[0].font.size = Pt(12)
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Content cells
            verification.cell(1,0).text = ("Name:\n\n\nSignature:\n\n\nDate:")
            verification.cell(1,1).text = ("Name:\n\n\nSignature:\n\n\nDate:")
            for cell in verification.rows[1].cells:
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Single-page formatting adjustments
            section = doc.sections[0]
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

            # Reduce paragraph spacing
            for paragraph in doc.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)

            doc.save(filename)
            
            # Open the document
            if os.name == 'nt':
                os.startfile(filename)
            elif os.name == 'posix':
                subprocess.run(['open', filename] if sys.platform == 'darwin' else ['xdg-open', filename])
        
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to create document: {str(e)}")
                        
    def add_patient(self):
        """Display the add patient form with scrollable fields"""
        if self.users[self.current_user]["role"] not in ["admin", "editor"]:
            messagebox.showerror("Access Denied", "Only admins and editors can add patients")
            return

        self.clear_frame()

        # Main container with gradient background
        main_frame = tk.Frame(self.root, bg='white')
        main_frame.pack(expand=True, fill=tk.BOTH)
        
        # Left side with logo and app name
        left_frame = tk.Frame(main_frame, bg='#3498db')
        left_frame.pack(side=tk.LEFT, fill=tk.Y, expand=False)
        
        # App logo and name
        logo_frame = tk.Frame(left_frame, bg='#3498db')
        logo_frame.pack(expand=True, fill=tk.BOTH, padx=40, pady=40)
        
        # App name with modern font
        tk.Label(logo_frame, text="OncoCare", font=('Helvetica', 24, 'bold'), 
                bg='#3498db', fg='white').pack(pady=(0, 10))
        
        tk.Label(logo_frame, text="Add New Patient", 
                font=('Helvetica', 14), bg='#3498db', fg='white').pack(pady=(0, 40))
        
        # Right side with form
        right_frame = tk.Frame(main_frame, bg='white')
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        # Create a container frame for the canvas and scrollbar
        container = ttk.Frame(right_frame)
        container.pack(fill=tk.BOTH, expand=True)
        
        # Create a canvas
        canvas = tk.Canvas(container, highlightthickness=0)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Add a scrollbar
        scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Configure the canvas
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Create a frame inside the canvas
        form_container = ttk.Frame(canvas)
        canvas.create_window((0, 0), window=form_container, anchor="nw")
        
        # Bind the canvas to the scroll region
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
        
        form_container.bind("<Configure>", on_frame_configure)
        
        # Enable mouse wheel scrolling
        def _on_mouse_wheel(event):
            canvas.yview_scroll(-1 * (event.delta // 120), "units")
        
        canvas.bind_all("<MouseWheel>", _on_mouse_wheel)

        # Malignancy selection
        malignancy_frame = ttk.LabelFrame(form_container, text="Select Malignancy Type", style='TFrame')
        malignancy_frame.pack(fill=tk.X, padx=10, pady=10)

        self.malignancy_var = tk.StringVar()
        
        # Create a bold combobox for malignancy selection
        ttk.Label(malignancy_frame, text="Malignancy Type:", 
                 font=('Helvetica', 12)).pack(pady=5, anchor="w")
        
        malignancy_combo = ttk.Combobox(malignancy_frame, textvariable=self.malignancy_var,
                                      values=MALIGNANCIES, state="readonly",
                                      font=('Helvetica', 12, 'bold'), style='Malignancy.TCombobox')
        malignancy_combo.pack(fill=tk.X, padx=5, pady=5)
        
        # Bind the selection to show malignancy fields
        malignancy_combo.bind("<<ComboboxSelected>>", lambda e: self.show_malignancy_fields(form_container))

        # Container for malignancy-specific fields
        self.malignancy_fields_frame = ttk.Frame(form_container, style='TFrame')
        self.malignancy_fields_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Button frame at the bottom (outside the scrollable area)
        btn_frame = ttk.Frame(right_frame, padding="10 10 10 10", style='TFrame')
        btn_frame.pack(fill=tk.X)

        ttk.Button(btn_frame, text="Back", command=self.main_menu).pack(side=tk.LEFT, padx=10)

    def show_malignancy_fields(self, form_container):
        """Original layout with corrected dropdown loading"""
        # Clear previous fields
        for widget in self.malignancy_fields_frame.winfo_children():
            widget.destroy()

        malignancy = self.malignancy_var.get()
        if not malignancy:
            return

        # Get dropdown options using corrected method
        DROPDOWN_OPTIONS = self.get_dropdown_options()

        # Common fields
        common_frame = ttk.LabelFrame(self.malignancy_fields_frame, 
                                    text="Common Information", 
                                    style='TFrame')
        common_frame.pack(fill=tk.X, padx=10, pady=10)

        self.entries = {}
        row = 0

        # Create common fields with FIXED dropdown values
        for field in COMMON_FIELDS:
            ttk.Label(common_frame, text=f"{field}:", anchor="w").grid(
                row=row, column=0, padx=5, pady=5, sticky="w")
            
            if field == "GENDER":
                var = tk.StringVar()
                combobox = ttk.Combobox(common_frame, textvariable=var, 
                                      values=DROPDOWN_OPTIONS.get("GENDER", []), 
                                      state="readonly")
                combobox.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                self.entries[field] = var
            elif field in ["EXAMINATION", "SYMPTOMS"]:
                listbox_frame = ttk.Frame(common_frame)
                listbox_frame.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                
                listbox = tk.Listbox(listbox_frame, selectmode=tk.MULTIPLE, height=4, exportselection=0)
                scrollbar = ttk.Scrollbar(listbox_frame, orient="vertical", command=listbox.yview)
                listbox.configure(yscrollcommand=scrollbar.set)
                
                options = DROPDOWN_OPTIONS.get(field, [])
                for option in options:
                    listbox.insert(tk.END, option)
                
                listbox.pack(side=tk.LEFT, fill=tk.X, expand=True)
                scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
                self.entries[field] = listbox
                
                if "OTHERS" in options:
                    ttk.Label(common_frame, text=f"{field} (Others):").grid(
                        row=row+1, column=0, padx=5, pady=5, sticky="w")
                    others_entry = ttk.Entry(common_frame)
                    others_entry.grid(row=row+1, column=1, padx=5, pady=5, sticky="ew")
                    self.entries[f"{field}_OTHERS"] = others_entry
                    row += 1
            elif field in ["DATE OF BIRTH", "AGE ON DIAGNOSIS"]:
                entry_frame = ttk.Frame(common_frame)
                entry_frame.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                
                entry = ttk.Entry(entry_frame)
                entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
                
                cal_btn = ttk.Button(entry_frame, text="📅", width=3,
                                    command=lambda e=entry: self.show_calendar(e))
                cal_btn.pack(side=tk.LEFT, padx=5)
                self.entries[field] = entry            
            else:
                entry = ttk.Entry(common_frame)
                entry.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                self.entries[field] = entry
            
            row += 1

        # Malignancy-specific fields (original structure)
        malignancy_frame = ttk.LabelFrame(self.malignancy_fields_frame,
                                        text=f"{malignancy} Specific Information")
        malignancy_frame.pack(fill=tk.X, padx=10, pady=10)

        row = 0
        for field in MALIGNANCY_FIELDS.get(malignancy, []):
            ttk.Label(malignancy_frame, text=f"{field}:", anchor="w").grid(
                row=row, column=0, padx=5, pady=5, sticky="w")
            
            if field in DROPDOWN_OPTIONS:
                var = tk.StringVar()
                combobox = ttk.Combobox(malignancy_frame, textvariable=var, 
                                      values=DROPDOWN_OPTIONS[field], state="readonly")
                combobox.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                self.entries[field] = var
            elif field.endswith("_DATE"):
                entry_frame = ttk.Frame(malignancy_frame)
                entry_frame.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                
                entry = ttk.Entry(entry_frame)
                entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
                
                cal_btn = ttk.Button(entry_frame, text="📅", width=3,
                                    command=lambda e=entry: self.show_calendar(e))
                cal_btn.pack(side=tk.LEFT, padx=5)
                self.entries[field] = entry
            elif "_SIDE_EFFECTS" in field:
                listbox_frame = ttk.Frame(malignancy_frame)
                listbox_frame.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                
                listbox = tk.Listbox(listbox_frame, selectmode=tk.MULTIPLE, height=4, exportselection=0)
                scrollbar = ttk.Scrollbar(listbox_frame, orient="vertical", command=listbox.yview)
                listbox.configure(yscrollcommand=scrollbar.set)
                
                options = DROPDOWN_OPTIONS.get(field, [])
                for option in options:
                    listbox.insert(tk.END, option)
                
                listbox.pack(side=tk.LEFT, fill=tk.X, expand=True)
                scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
                self.entries[field] = listbox
                
                if "OTHERS" in options:
                    ttk.Label(malignancy_frame, text=f"{field} (Others):").grid(
                        row=row+1, column=0, padx=5, pady=5, sticky="w")
                    others_entry = ttk.Entry(malignancy_frame)
                    others_entry.grid(row=row+1, column=1, padx=5, pady=5, sticky="ew")
                    self.entries[f"{field}_OTHERS"] = others_entry
                    row += 1
            else:
                entry = ttk.Entry(malignancy_frame)
                entry.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                self.entries[field] = entry
            
            row += 1

        # Button frame (original structure)
        btn_frame = ttk.Frame(self.malignancy_fields_frame)
        btn_frame.pack(fill=tk.X, padx=10, pady=10)

        ttk.Button(btn_frame, text="Save Patient", command=self.save_patient,
                  style='Blue.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Patient Folder", command=self.open_patient_folder,
                  style='Green.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Main Menu", command=self.main_menu).pack(side=tk.RIGHT, padx=5)
    
    def save_patient(self):
        """Save patient data to file and sync with Google Drive"""
        # Check if malignancy is selected
        malignancy = self.malignancy_var.get()
        if not malignancy:
            messagebox.showerror("Error", "Please select a malignancy type")
            return

        # Check mandatory fields
        mandatory_fields = [
            "FILE NUMBER", "NAME", "GENDER", "DATE OF BIRTH", "NATIONALITY",
            "DIAGNOSIS", "AGE ON DIAGNOSIS"
        ]
        
        missing_fields = []
        for field in mandatory_fields:
            value = self.entries[field].get() if isinstance(self.entries[field], (tk.Entry, tk.StringVar)) else None
            if not value:
                missing_fields.append(field)
        
        # Check at least one from examination, history, symptoms
        exam_selected = False
        if "EXAMINATION" in self.entries:
            exam_selected = len(self.entries["EXAMINATION"].curselection()) > 0 or \
                          ("EXAMINATION_OTHERS" in self.entries and self.entries["EXAMINATION_OTHERS"].get())
        
        history_selected = "HISTORY" in self.entries and self.entries["HISTORY"].get()
        symptoms_selected = False
        if "SYMPTOMS" in self.entries:
            symptoms_selected = len(self.entries["SYMPTOMS"].curselection()) > 0 or \
                              ("SYMPTOMS_OTHERS" in self.entries and self.entries["SYMPTOMS_OTHERS"].get())
        
        if not (exam_selected or history_selected or symptoms_selected):
            missing_fields.append("At least one from EXAMINATION and SYMPTOMS")
        
        if missing_fields:
            messagebox.showerror("Error", f"The following fields are mandatory:\n{', '.join(missing_fields)}")
            return

        # Check and validate file number
        file_no = self.entries["FILE NUMBER"].get().strip().upper()
        if not file_no:
            messagebox.showerror("Error", "File Number cannot be empty")
            return
            
        # Validate file number format
        import re
        file_pattern = re.compile(r'^[A-Z0-9\/\.\\*\-]*[0-9]+[A-Z0-9\/\.\\*\-]*$')
        if not file_pattern.match(file_no):
            messagebox.showerror("Error", 
                "Invalid File Number format. Must contain:\n"
                "- At least one number\n"
                "- Allowed characters: A-Z, 0-9, / . \\ * -")
            return
            
        # Check file number uniqueness
        for patient in self.patient_data:
            if patient.get("FILE NUMBER", "").upper() == file_no:
                messagebox.showerror("Error", "Patient with this File Number already exists")
                return

        # Validate date formats
        date_fields = {
            "DATE OF BIRTH": self.entries["DATE OF BIRTH"].get(),
            "AGE ON DIAGNOSIS": self.entries["AGE ON DIAGNOSIS"].get()
        }
        
        for field_name, date_value in date_fields.items():
            try:
                datetime.strptime(date_value, '%d/%m/%Y')
            except ValueError:
                messagebox.showerror("Error", 
                    f"Invalid date format for '{field_name}'. Please use dd/mm/yyyy")
                return

        # Collect patient data
        patient_data = {
            "MALIGNANCY": malignancy,
            "CREATED_BY": self.current_user,
            "CREATED_DATE": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
            "LAST_MODIFIED_BY": self.current_user,
            "LAST_MODIFIED_DATE": datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        }

        for field, widget in self.entries.items():
            if isinstance(widget, tk.Listbox):  # Multi-select fields
                selected = [widget.get(i) for i in widget.curselection()]
                if "OTHERS" in selected and f"{field}_OTHERS" in self.entries:
                    others_text = self.entries[f"{field}_OTHERS"].get()
                    if others_text:
                        selected[selected.index("OTHERS")] = f"OTHERS: {others_text}"
                patient_data[field] = ", ".join(selected) if selected else ""
            elif isinstance(widget, tk.StringVar):  # Combobox
                patient_data[field] = widget.get()
            else:  # Entry fields
                patient_data[field] = widget.get()

        # Save the patient data
        self.patient_data.append(patient_data)
        # Sort patient data by file number (alphanumeric)
        self.patient_data.sort(key=lambda x: str(x.get("FILE NUMBER", "")).upper())
        self.save_patient_data()

        # Create patient folder and report
        self.create_patient_folder(file_no)
        self.create_patient_report(file_no, create_only=True)

        # Upload to Google Drive in background
        self.executor.submit(self.upload_patient_to_drive, patient_data)

        messagebox.showinfo("Success", "Patient data saved successfully!")
        self.main_menu()

    def show_calendar(self, entry_widget):
        """Show a calendar popup for date selection"""
        top = tk.Toplevel(self.root)
        top.title("Select Date")
        top.geometry("300x300")
        
        cal = Calendar(top, selectmode='day', date_pattern='dd/mm/yyyy')
        cal.pack(pady=20)
        
        def set_date():
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, cal.get_date())
            top.destroy()
        
        ttk.Button(top, text="Select", command=set_date, style='Blue.TButton').pack(pady=10)
        
    def upload_patient_to_drive(self, patient_data):
        """Upload patient data and files to Google Drive"""
        if not self.drive.initialized:
            return False
        
        try:
            # Upload patient data
            file_no = patient_data["FILE NUMBER"]
            self.drive.upload_patient_data(patient_data)
            
            # Sync patient folder
            self.drive.sync_patient_files(file_no)
            
            return True
        except Exception as e:
            print(f"Error uploading patient to Google Drive: {e}")
            return False
    
    def create_patient_folder(self, file_no):
        """Create a folder for the patient's documents"""
        folder_name = f"Patient_{file_no}"
        if not os.path.exists(folder_name):
            os.makedirs(folder_name)
    
    def open_patient_folder(self):
        """Open the patient's folder in file explorer"""
        file_no = self.entries["FILE NUMBER"].get()
        if not file_no:
            messagebox.showerror("Error", "Please enter a file number first")
            return
        
        folder_name = f"Patient_{file_no}"
        self.create_patient_folder(file_no)
        
        # Try to sync with Google Drive
        if self.drive.initialized:
            self.executor.submit(self.drive.sync_patient_files, file_no)
        
        try:
            os.startfile(folder_name)
        except:
            messagebox.showerror("Error", f"Could not open folder: {folder_name}")
    
    def create_patient_report(self, file_no=None, create_only=False):
        """Create a Word document report for the patient"""
        try:
            if not file_no:
                file_no = self.entries["FILE NUMBER"].get()
                if not file_no:
                    messagebox.showerror("Error", "Please enter a file number first")
                    return

            doc_name = f"Patient_{file_no}_Report.docx"
            doc_path = os.path.join(f"Patient_{file_no}", doc_name)
            
            # Ensure patient folder exists
            os.makedirs(f"Patient_{file_no}", exist_ok=True)

            if not os.path.exists(doc_path) or not create_only:
                doc = Document()
                
                # Add title with error handling
                try:
                    doc.add_heading(f'Patient Report - File Number: {file_no}', 0)
                except ValueError:
                    doc.add_heading(f'Patient Report', 0)

                # Add basic info section
                doc.add_heading('Basic Information', level=1)
                
                if not create_only:
                    # Create from current form data
                    for field in COMMON_FIELDS:
                        value = ""
                        if field in self.entries:
                            widget = self.entries[field]
                            if isinstance(widget, tk.Listbox):
                                selected = widget.curselection()
                                value = ", ".join([widget.get(i) for i in selected])
                            elif isinstance(widget, tk.StringVar):
                                value = widget.get()
                            else:
                                value = widget.get()
                        doc.add_paragraph(f"{field}: {value}", style='List Bullet')
                else:
                    # Create from existing data with validation
                    patient = next((p for p in self.patient_data if str(p.get("FILE NUMBER")) == str(file_no)), None)
                    if patient:
                        for field in COMMON_FIELDS:
                            value = patient.get(field, "N/A")
                            doc.add_paragraph(f"{field}: {value}", style='List Bullet')

                # Save with error handling
                try:
                    doc.save(doc_path)
                except PermissionError:
                    messagebox.showerror("Error", f"Please close {doc_name} before saving")
                    return

            if not create_only:
                # Upload to Google Drive
                if self.drive.initialized:
                    self.executor.submit(
                        self.drive.upload_file,
                        doc_path,
                        doc_name,
                        self.drive.create_patient_folder(file_no)
                    )
                
                # Open document with error handling
                try:
                    if os.name == 'nt':
                        os.startfile(doc_path)
                    else:
                        subprocess.call(('xdg-open', doc_path))
                except Exception as e:
                    messagebox.showerror("Error", f"Could not open document: {str(e)}")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to create report: {str(e)}")
            self.main_menu()

    def search_patient(self):
        """Display the patient search form"""
        self.clear_frame()
        try:
            main_frame = tk.Frame(self.root, bg='white')
            main_frame.pack(expand=True, fill=tk.BOTH)

            # Left sidebar
            left_frame = tk.Frame(main_frame, bg='#3498db')
            left_frame.pack(side=tk.LEFT, fill=tk.Y, expand=False)
            
            # Logo section
            logo_frame = tk.Frame(left_frame, bg='#3498db')
            logo_frame.pack(padx=40, pady=40, fill=tk.BOTH, expand=True)
            tk.Label(logo_frame, text="OncoCare", font=('Helvetica', 24, 'bold'), 
                    bg='#3498db', fg='white').pack(pady=(0, 10))
            tk.Label(logo_frame, text="Search Patient", 
                    font=('Helvetica', 14), bg='#3498db', fg='white').pack()

            # Right side form
            right_frame = tk.Frame(main_frame, bg='white')
            right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

            # Form container with improved layout
            form_container = ttk.Frame(right_frame, style='TFrame')
            form_container.place(relx=0.5, rely=0.4, anchor='center')

            # Search fields with validation
            ttk.Label(form_container, text="File Number:", font=('Helvetica', 12)).grid(row=0, column=0, padx=5, pady=5)
            self.search_file_no_entry = ttk.Entry(form_container, width=25, font=('Helvetica', 12))
            self.search_file_no_entry.grid(row=0, column=1, padx=5, pady=5)

            ttk.Label(form_container, text="Name:", font=('Helvetica', 12)).grid(row=1, column=0, padx=5, pady=5)
            self.search_name_entry = ttk.Entry(form_container, width=25, font=('Helvetica', 12))
            self.search_name_entry.grid(row=1, column=1, padx=5, pady=5)

            # Button grid
            btn_frame = ttk.Frame(form_container)
            btn_frame.grid(row=2, column=0, columnspan=2, pady=20)

            ttk.Button(btn_frame, text="Search", command=self.perform_search,
                       style='Blue.TButton').grid(row=0, column=0, padx=10)
            ttk.Button(btn_frame, text="Back", command=self.main_menu).grid(row=0, column=1, padx=10)

        except Exception as e:
            messagebox.showerror("Error", f"Search form error: {str(e)}")
            self.main_menu()

    def perform_search(self):
        """Perform patient search with validation"""
        try:
            file_no = self.search_file_no_entry.get().strip()
            name = self.search_name_entry.get().strip().lower()

            if not self.patient_data:
                messagebox.showinfo("Info", "No patient data available")
                return

            # Convert all file numbers to strings for comparison
            str_patient_data = []
            for p in self.patient_data:
                p["FILE NUMBER"] = str(p.get("FILE NUMBER", ""))
                str_patient_data.append(p)

            # Filter patients
            matching = []
            for patient in str_patient_data:
                file_match = True
                name_match = True
                
                if file_no:
                    file_match = patient["FILE NUMBER"] == file_no
                if name:
                    name_match = name in patient.get("NAME", "").lower()
                
                if file_match and name_match:
                    matching.append(patient)

            if not matching:
                messagebox.showinfo("Not Found", "No matching patients found")
                return

            self.current_results = matching
            self.current_result_index = 0
            self.view_patient(self.current_results[0], len(matching) > 1)

        except Exception as e:
            messagebox.showerror("Error", f"Search failed: {str(e)}")
            self.search_patient()

    def view_patient(self, patient_data, multiple_results=False):
        """View patient details with error handling"""
        try:
            self.clear_frame()
            if not patient_data:
                raise ValueError("No patient data provided")

            # Main container setup
            main_frame = tk.Frame(self.root, bg='white')
            main_frame.pack(expand=True, fill=tk.BOTH)

            # Left sidebar
            left_frame = tk.Frame(main_frame, bg='#3498db')
            left_frame.pack(side=tk.LEFT, fill=tk.Y, expand=False)
            
            # Logo section
            logo_frame = tk.Frame(left_frame, bg='#3498db')
            logo_frame.pack(padx=40, pady=40, fill=tk.BOTH, expand=True)
            malignancy = patient_data.get("MALIGNANCY", "Unknown")
            tk.Label(logo_frame, text="OncoCare", font=('Helvetica', 24, 'bold'), 
                    bg='#3498db', fg='white').pack(pady=(0, 10))
            tk.Label(logo_frame, text=f"Patient Details\n({malignancy})", 
                    font=('Helvetica', 14), bg='#3498db', fg='white').pack()

            # Right side content
            right_frame = tk.Frame(main_frame, bg='white')
            right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

            # Scrollable content
            canvas = tk.Canvas(right_frame, bg='white')
            scrollbar = ttk.Scrollbar(right_frame, orient="vertical", command=canvas.yview)
            scrollable_frame = ttk.Frame(canvas)

            # Configure scrolling
            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(
                    scrollregion=canvas.bbox("all")
                )
            )
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)

            # Pack scroll elements
            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")

            # Patient information
            bg_color = MALIGNANCY_COLORS.get(malignancy, "#e6f2ff")
            
            # Metadata section
            meta_frame = ttk.Frame(scrollable_frame)
            meta_frame.pack(fill=tk.X, pady=10, padx=20)
            created = patient_data.get("CREATED_DATE", "Unknown")
            modified = patient_data.get("LAST_MODIFIED_DATE", "Unknown")
            ttk.Label(meta_frame, text=f"Created: {created} | Last Modified: {modified}", 
                     font=('Helvetica', 9)).pack()

            # Common fields section
            common_frame = ttk.LabelFrame(scrollable_frame, text="Basic Information")
            common_frame.pack(fill=tk.X, padx=20, pady=10)
            
            for i, field in enumerate(COMMON_FIELDS):
                row = ttk.Frame(common_frame)
                row.pack(fill=tk.X, pady=2)
                ttk.Label(row, text=f"{field}:", width=20, anchor='w').pack(side=tk.LEFT)
                value = patient_data.get(field, "N/A")
                ttk.Label(row, text=value, width=30, anchor='w').pack(side=tk.LEFT)

            # Malignancy-specific fields
            if malignancy in MALIGNANCY_FIELDS:
                specific_frame = ttk.LabelFrame(scrollable_frame, 
                    text=f"{malignancy} Specific Information")
                specific_frame.pack(fill=tk.X, padx=20, pady=10)
                
                for field in MALIGNANCY_FIELDS[malignancy]:
                    row = ttk.Frame(specific_frame)
                    row.pack(fill=tk.X, pady=2)
                    ttk.Label(row, text=f"{field}:", width=20, anchor='w').pack(side=tk.LEFT)
                    value = patient_data.get(field, "N/A")
                    ttk.Label(row, text=value, width=30, anchor='w').pack(side=tk.LEFT)

            # Action buttons
            btn_frame = ttk.Frame(right_frame)
            btn_frame.pack(fill=tk.X, pady=10)
            
            actions = [
                ("Edit", lambda: self.edit_patient(patient_data), "Blue.TButton"),
                ("Delete", lambda: self.delete_patient(patient_data), "Red.TButton"),
                ("Folder", lambda: self.open_existing_patient_folder(
                    patient_data["FILE NUMBER"]), "Green.TButton"),
                ("Report", lambda: self.open_existing_patient_report(
                    patient_data["FILE NUMBER"]), "Yellow.TButton"),
                # Add this new button entry
                ("Lab/EF", lambda: self.show_lab_ef_window(patient_data), "Cyan.TButton")
            ]
            
            for text, cmd, style in actions:
                if text == "Delete" and self.current_user != "mej.esam":
                    continue
                ttk.Button(btn_frame, text=text, command=cmd, style=style).pack(side=tk.LEFT, padx=5)

            if multiple_results:
                ttk.Button(btn_frame, text="Previous", 
                          command=self.show_previous_result).pack(side=tk.LEFT, padx=5)
                ttk.Button(btn_frame, text="Next", 
                          command=self.show_next_result).pack(side=tk.LEFT, padx=5)

            ttk.Button(btn_frame, text="Main Menu", 
                      command=self.main_menu).pack(side=tk.RIGHT, padx=5)

        except Exception as e:
            messagebox.showerror("Error", f"Failed to display patient: {str(e)}")
            self.main_menu()

    def open_existing_patient_folder(self, file_no):
        """Open an existing patient's folder"""
        folder_name = f"Patient_{file_no}"
        if not os.path.exists(folder_name):
            self.create_patient_folder(file_no)
        
        # Try to sync with Google Drive
        if self.drive.initialized:
            self.executor.submit(self.drive.sync_patient_files, file_no)
        
        try:
            os.startfile(folder_name)
        except:
            messagebox.showerror("Error", f"Could not open folder: {folder_name}")

    def open_existing_patient_report(self, file_no):
        """Open an existing patient's report"""
        doc_name = f"Patient_{file_no}_Report.docx"
        doc_path = os.path.join(f"Patient_{file_no}", doc_name)
        
        if not os.path.exists(doc_path) and self.drive.initialized:
            self.drive.download_file(doc_name, None, doc_path)
        
        try:
            os.startfile(doc_path)
        except:
            messagebox.showerror("Error", f"Could not open document: {doc_path}")

    def show_previous_result(self):
        """Show the previous search result"""
        if self.current_result_index > 0:
            self.current_result_index -= 1
            self.view_patient(self.current_results[self.current_result_index], multiple_results=True)
        else:
            messagebox.showinfo("Info", "This is the first result.")

    def show_next_result(self):
        """Show the next search result"""
        if self.current_result_index < len(self.current_results) - 1:
            self.current_result_index += 1
            self.view_patient(self.current_results[self.current_result_index], multiple_results=True)
        else:
            messagebox.showinfo("Info", "This is the last result.")

    def edit_patient(self, patient_data):
        """Edit patient form with editable malignancy type"""
        if self.users[self.current_user]["role"] not in ["admin", "editor"]:
            messagebox.showerror("Access Denied", "Only admins and editors can edit patients")
            return
            
        self.clear_frame()
        
        # Main container with gradient background
        main_frame = tk.Frame(self.root, bg='white')
        main_frame.pack(expand=True, fill=tk.BOTH)
        
        # Left side with logo and app name
        left_frame = tk.Frame(main_frame, bg='#3498db')
        left_frame.pack(side=tk.LEFT, fill=tk.Y, expand=False)
        
        # App logo and name
        logo_frame = tk.Frame(left_frame, bg='#3498db')
        logo_frame.pack(expand=True, fill=tk.BOTH, padx=40, pady=40)
        
        malignancy = patient_data.get("MALIGNANCY", "")
        title = f"Edit Patient - {malignancy}" if malignancy else "Edit Patient"
        
        # App name with modern font
        tk.Label(logo_frame, text="OncoCare", font=('Helvetica', 24, 'bold'), 
                bg='#3498db', fg='white').pack(pady=(0, 10))
        
        tk.Label(logo_frame, text=title, 
                font=('Helvetica', 14), bg='#3498db', fg='white').pack(pady=(0, 40))
        
        # Right side with form
        right_frame = tk.Frame(main_frame, bg='white')
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        # Form container
        form_frame = self.create_scrollable_frame(right_frame)
        
        # Editable malignancy type section
        malignancy_frame = ttk.LabelFrame(form_frame, text="Malignancy Type", style='TFrame')
        malignancy_frame.pack(fill=tk.X, padx=10, pady=10)
        
        current_malignancy = patient_data.get("MALIGNANCY", "")
        self.malignancy_var = tk.StringVar(value=current_malignancy)
        
        malignancy_combo = ttk.Combobox(malignancy_frame, 
                                      textvariable=self.malignancy_var,
                                      values=MALIGNANCIES,
                                      state="readonly")
        malignancy_combo.pack(pady=5)

        # Common fields
        common_frame = ttk.LabelFrame(form_frame, text="Common Information", style='TFrame')
        common_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.edit_entries = {}
        row = 0
        
        # Load dropdowns using public method
        dropdowns = self.get_dropdown_options()

        # Create common fields
        for field in COMMON_FIELDS:
            ttk.Label(common_frame, text=f"{field}:", anchor="w").grid(row=row, column=0, padx=5, pady=5, sticky="w")
            
            current_value = patient_data.get(field, "")
            
            if field == "GENDER":
                var = tk.StringVar(value=current_value)
                combobox = ttk.Combobox(common_frame, textvariable=var, 
                                      values=dropdowns.get("GENDER", []), 
                                      state="readonly")
                combobox.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                self.edit_entries[field] = var
            elif field in ["EXAMINATION", "SYMPTOMS"]:
                listbox_frame = ttk.Frame(common_frame)
                listbox_frame.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                
                listbox = tk.Listbox(listbox_frame, selectmode=tk.MULTIPLE, height=4, exportselection=0)
                scrollbar = ttk.Scrollbar(listbox_frame, orient="vertical", command=listbox.yview)
                listbox.configure(yscrollcommand=scrollbar.set)
                
                options = dropdowns.get(field, [])
                for option in options:
                    listbox.insert(tk.END, option)
                
                # Select previously selected items
                selected_values = patient_data.get(field, "").split(", ") if patient_data.get(field) else []
                for i, item in enumerate(options):
                    if item in selected_values:
                        listbox.select_set(i)
                
                listbox.pack(side=tk.LEFT, fill=tk.X, expand=True)
                scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
                self.edit_entries[field] = listbox
                
                # Add "Others" entry if needed
                if "OTHERS" in options:
                    others_value = ""
                    for val in selected_values:
                        if val.startswith("OTHERS:"):
                            others_value = val[8:]  # Remove "OTHERS: " prefix
                            break
                    
                    ttk.Label(common_frame, text=f"{field} (Others):").grid(row=row+1, column=0, padx=5, pady=5, sticky="w")
                    others_entry = ttk.Entry(common_frame)
                    others_entry.insert(0, others_value)
                    others_entry.grid(row=row+1, column=1, padx=5, pady=5, sticky="ew")
                    self.edit_entries[f"{field}_OTHERS"] = others_entry
                    row += 1
            elif field == "DATE OF BIRTH":
                entry_frame = ttk.Frame(common_frame)
                entry_frame.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                
                entry = ttk.Entry(entry_frame)
                entry.insert(0, current_value)
                entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
                
                cal_btn = ttk.Button(entry_frame, text="📅", width=3,
                                    command=lambda e=entry: self.show_calendar(e))
                cal_btn.pack(side=tk.LEFT, padx=5)
                self.edit_entries[field] = entry
            else:
                entry = ttk.Entry(common_frame)
                entry.insert(0, current_value)
                entry.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                self.edit_entries[field] = entry
            
            row += 1
        
        # Malignancy-specific fields (now based on selected malignancy)
        malignancy = self.malignancy_var.get()
        if malignancy in MALIGNANCY_FIELDS:
            specific_frame = tk.LabelFrame(form_frame, text=f"{malignancy} Specific Information", 
                                          bg=MALIGNANCY_COLORS.get(malignancy, "#FFFFFF"))
            specific_frame.pack(fill=tk.X, padx=10, pady=10)
            
            row = 0
            for field in MALIGNANCY_FIELDS[malignancy]:
                ttk.Label(specific_frame, text=f"{field}:", anchor="w").grid(row=row, column=0, padx=5, pady=5, sticky="w")
                
                current_value = patient_data.get(field, "")
                
                if field in dropdowns:
                    var = tk.StringVar(value=current_value)
                    combobox = ttk.Combobox(specific_frame, textvariable=var, 
                                          values=dropdowns[field], state="readonly")
                    combobox.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                    self.edit_entries[field] = var
                elif field.endswith("_DATE"):
                    entry_frame = ttk.Frame(specific_frame)
                    entry_frame.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                    
                    entry = ttk.Entry(entry_frame)
                    entry.insert(0, current_value)
                    entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
                    
                    cal_btn = ttk.Button(entry_frame, text="📅", width=3,
                                        command=lambda e=entry: self.show_calendar(e))
                    cal_btn.pack(side=tk.LEFT, padx=5)
                    self.edit_entries[field] = entry
                elif "_SIDE_EFFECTS" in field:
                    listbox_frame = ttk.Frame(specific_frame)
                    listbox_frame.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                    
                    listbox = tk.Listbox(listbox_frame, selectmode=tk.MULTIPLE, height=4, exportselection=0)
                    scrollbar = ttk.Scrollbar(listbox_frame, orient="vertical", command=listbox.yview)
                    listbox.configure(yscrollcommand=scrollbar.set)
                    
                    options = dropdowns.get(field, [])
                    for option in options:
                        listbox.insert(tk.END, option)
                    
                    # Select previously selected items
                    selected_values = patient_data.get(field, "").split(", ") if patient_data.get(field) else []
                    for i, item in enumerate(options):
                        if item in selected_values:
                            listbox.select_set(i)
                    
                    listbox.pack(side=tk.LEFT, fill=tk.X, expand=True)
                    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
                    self.edit_entries[field] = listbox
                    
                    if "OTHERS" in options:
                        others_value = ""
                        for val in selected_values:
                            if val.startswith("OTHERS:"):
                                others_value = val[8:]  # Remove "OTHERS: " prefix
                                break
                        
                        ttk.Label(specific_frame, text=f"{field} (Others):").grid(row=row+1, column=0, padx=5, pady=5, sticky="w")
                        others_entry = ttk.Entry(specific_frame)
                        others_entry.insert(0, others_value)
                        others_entry.grid(row=row+1, column=1, padx=5, pady=5, sticky="ew")
                        self.edit_entries[f"{field}_OTHERS"] = others_entry
                        row += 1
                else:
                    entry = ttk.Entry(specific_frame)
                    entry.insert(0, current_value)
                    entry.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
                    self.edit_entries[field] = entry
                
                row += 1
        
        # Button frame
        btn_frame = ttk.Frame(right_frame, padding="10 10 10 10", style='TFrame')
        btn_frame.pack(fill=tk.X)
        
        ttk.Button(btn_frame, text="Save Changes", 
                  command=lambda: self.save_edited_patient(patient_data),
                  style='Blue.TButton').pack(side=tk.LEFT, padx=10)
        
        ttk.Button(btn_frame, text="Cancel", 
                  command=lambda: self.view_patient(patient_data)).pack(side=tk.LEFT, padx=10)
        
        ttk.Button(btn_frame, text="Main Menu", command=self.main_menu,
                  style='Blue.TButton').pack(side=tk.RIGHT, padx=10)

    def save_edited_patient(self, original_data):
        """Save edited patient data with updated malignancy"""
        # Validate date format
        dob = self.edit_entries["DATE OF BIRTH"].get()
        try:
            datetime.strptime(dob, '%d/%m/%Y')
        except ValueError:
            messagebox.showerror("Error", "Invalid date format. Please use dd/mm/yyyy")
            return
            
        updated_data = {
            "MALIGNANCY": self.malignancy_var.get(),  # Get from combobox
            "CREATED_BY": original_data.get("CREATED_BY", ""),
            "CREATED_DATE": original_data.get("CREATED_DATE", ""),
            "LAST_MODIFIED_BY": self.current_user,
            "LAST_MODIFIED_DATE": datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        }
        
        for field, widget in self.edit_entries.items():
            if isinstance(widget, tk.Listbox):  # Multi-select fields
                selected = [widget.get(i) for i in widget.curselection()]
                if "OTHERS" in selected and f"{field}_OTHERS" in self.edit_entries:
                    others_text = self.edit_entries[f"{field}_OTHERS"].get()
                    if others_text:
                        selected[selected.index("OTHERS")] = f"OTHERS: {others_text}"
                updated_data[field] = ", ".join(selected) if selected else ""
            elif isinstance(widget, tk.StringVar):  # Combobox
                updated_data[field] = widget.get()
            else:  # Entry fields
                updated_data[field] = widget.get()
        
        # Update the record
        file_no = original_data.get("FILE NUMBER")
        for i, patient in enumerate(self.patient_data):
            if patient.get("FILE NUMBER") == file_no:
                self.patient_data[i] = updated_data
                break
        
        # Sort patient data by file number (numeric)
        self.patient_data.sort(key=lambda x: int(x.get("FILE NUMBER", 0)))
        self.save_patient_data()
        
        # Update patient report
        self.create_patient_report(file_no, create_only=True)
        
        # Upload to Google Drive in background
        self.executor.submit(self.upload_patient_to_drive, updated_data)
        
        # Show confirmation message
        messagebox.showinfo("Success", "Patient data updated successfully!")
        
        # Force the message box to display and process events
        self.root.update()
        
        # Now show the updated patient view
        self.view_patient(updated_data)
                
    def delete_patient(self, patient_data):
        """Delete a patient record"""
        if self.current_user != "mej.esam":
            messagebox.showerror("Access Denied", "Only 'mej.esam' can delete patients.")
            return
        
        confirm = messagebox.askyesno("Confirm Delete", "Are you sure you want to delete this patient?")
        if not confirm:
            return
        
        file_no = patient_data.get("FILE NUMBER")
        
        # Remove from local data
        self.patient_data = [patient for patient in self.patient_data if patient.get("FILE NUMBER") != file_no]
        self.save_patient_data()
        
        # Remove from Google Drive in background
        if self.drive.initialized:
            self.executor.submit(self.delete_patient_from_drive, file_no)
        
        # Delete local folder
        try:
            folder_name = f"Patient_{file_no}"
            if os.path.exists(folder_name):
                shutil.rmtree(folder_name)
        except Exception as e:
            print(f"Error deleting patient folder: {e}")
        
        messagebox.showinfo("Success", "Patient deleted successfully!")
        self.search_patient()
    
    def delete_patient_from_drive(self, file_no):
        """Delete patient data and folder from Google Drive"""
        if not self.drive.initialized:
            return False
        
        try:
            # Delete patient data file
            file_name = f"patient_{file_no}.json"
            query = f"name='{file_name}' and '{self.drive.app_folder_id}' in parents and trashed=false"
            results = self.drive.service.files().list(q=query, fields="files(id)").execute()
            items = results.get('files', [])
            
            if items:
                self.drive.service.files().delete(fileId=items[0]['id']).execute()
            
            # Delete patient folder
            folder_name = f"Patient_{file_no}"
            query = f"name='{folder_name}' and '{self.drive.patients_folder_id}' in parents and trashed=false"
            results = self.drive.service.files().list(q=query, fields="files(id)").execute()
            items = results.get('files', [])
            
            if items:
                self.drive.service.files().delete(fileId=items[0]['id']).execute()
            
            return True
        except Exception as e:
            print(f"Error deleting patient from Google Drive: {e}")
            return False
    
    def backup_data(self):
        """Create a backup of patient data"""
        if not self.patient_data:
            messagebox.showerror("Error", "No data available to back up.")
            return

        # Local backup
        backup_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
            title="Save Backup As"
        )
        if not backup_path:
            return

        try:
            with open(backup_path, 'w') as f:
                json.dump(self.patient_data, f, indent=4)
            
            # Upload backup to Google Drive
            if self.drive.initialized:
                backup_name = os.path.basename(backup_path)
                self.drive.upload_file(backup_path, backup_name, self.drive.app_folder_id)
            
            messagebox.showinfo("Success", f"Data backed up to {backup_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to back up data: {e}")
    
    def restore_data(self):
        """Restore patient data from backup"""
        if self.current_user != "mej.esam":
            messagebox.showerror("Access Denied", "Only 'mej.esam' can restore data.")
            return

        restore_path = filedialog.askopenfilename(
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
            title="Select Backup File"
        )
        if not restore_path:
            return

        try:
            with open(restore_path, 'r') as f:
                data = json.load(f)
            
            self.patient_data = data
            # Sort patient data by file number (numeric)
            self.patient_data.sort(key=lambda x: int(x.get("FILE NUMBER", 0)))
            self.save_patient_data()
            
            # Upload restored data to Google Drive
            if self.drive.initialized:
                for patient in self.patient_data:
                    self.drive.upload_patient_data(patient)
            
            messagebox.showinfo("Success", "Data restored successfully.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to restore data: {e}")
    
    def export_all_data(self):
        """Export all patient data to Excel"""
        if self.users[self.current_user]["role"] != "admin":
            messagebox.showerror("Access Denied", "Only admins can export all patient data.")
            return

        if not self.patient_data:
            messagebox.showerror("Error", "No data to export.")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
            title="Save All Patient Data As"
        )
        if not file_path:
            return

        try:
            # Create a Pandas Excel writer using XlsxWriter as the engine
            with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
                # Create a sheet for each malignancy type
                for malignancy in MALIGNANCIES:
                    malignancy_data = [patient for patient in self.patient_data if patient.get("MALIGNANCY") == malignancy]
                    if malignancy_data:
                        df = pd.DataFrame(malignancy_data)
                        df.to_excel(writer, sheet_name=malignancy, index=False)
                
                # Create a sheet with all patients
                df = pd.DataFrame(self.patient_data)
                df.to_excel(writer, sheet_name="ALL PATIENTS", index=False)
            
            # Upload to Google Drive
            if self.drive.initialized:
                self.drive.upload_file(file_path, os.path.basename(file_path), self.drive.app_folder_id)
            
            messagebox.showinfo("Success", f"All patient data exported to {file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export data: {e}")
    
    def view_all_patients(self):
        """View all patient records in a table"""
        if self.users[self.current_user]["role"] != "admin":
            messagebox.showerror("Access Denied", "Only admins can view all patient records.")
            return

        if not self.patient_data:
            messagebox.showerror("Error", "No data available to display.")
            return

        self.clear_frame()

        # Main container with gradient background
        main_frame = tk.Frame(self.root, bg='white')
        main_frame.pack(expand=True, fill=tk.BOTH)
        
        # Left side with logo and app name
        left_frame = tk.Frame(main_frame, bg='#3498db')
        left_frame.pack(side=tk.LEFT, fill=tk.Y, expand=False)
        
        # App logo and name
        logo_frame = tk.Frame(left_frame, bg='#3498db')
        logo_frame.pack(expand=True, fill=tk.BOTH, padx=40, pady=40)
        
        # App name with modern font
        tk.Label(logo_frame, text="OncoCare", font=('Helvetica', 24, 'bold'), 
                bg='#3498db', fg='white').pack(pady=(0, 10))
        
        tk.Label(logo_frame, text="All Patients", 
                font=('Helvetica', 14), bg='#3498db', fg='white').pack(pady=(0, 40))
        
        # Right side with content
        right_frame = tk.Frame(main_frame, bg='white')
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        # Filter frame
        filter_frame = ttk.Frame(right_frame, style='TFrame')
        filter_frame.pack(fill=tk.X, padx=20, pady=10)

        ttk.Label(filter_frame, text="Filter by Malignancy:", 
                 font=('Helvetica', 12)).pack(side=tk.LEFT, padx=5)
        
        self.malignancy_filter = ttk.Combobox(filter_frame, values=MALIGNANCIES, state="readonly")
        self.malignancy_filter.pack(side=tk.LEFT, padx=5)

        ttk.Button(filter_frame, text="Apply Filter", command=self.apply_malignancy_filter,
                  style='Blue.TButton').pack(side=tk.LEFT, padx=10)
        
        # View all button (outside the filter frame)
        ttk.Button(filter_frame, text="View All Patients", command=self.display_all_patients,
                  style='Blue.TButton').pack(side=tk.RIGHT, padx=10)

        # Main content
        content_frame = ttk.Frame(right_frame, style='TFrame')
        content_frame.pack(fill=tk.BOTH, expand=True)

        # Create a frame for the canvas and scrollbars
        container = ttk.Frame(content_frame, style='TFrame')
        container.pack(fill=tk.BOTH, expand=True)

        # Create a canvas for scrolling
        self.canvas = tk.Canvas(container, highlightthickness=0)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Add a vertical scrollbar
        y_scrollbar = ttk.Scrollbar(container, orient="vertical", command=self.canvas.yview)
        y_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Add a horizontal scrollbar with increased height
        x_scrollbar = ttk.Scrollbar(content_frame, orient="horizontal", command=self.canvas.xview)
        x_scrollbar.pack(side=tk.BOTTOM, fill=tk.X, pady=5, ipady=10)

        self.canvas.configure(xscrollcommand=x_scrollbar.set, yscrollcommand=y_scrollbar.set)

        # Create a frame inside the canvas to hold the table
        self.table_frame = ttk.Frame(self.canvas, style='TFrame')
        self.canvas.create_window((0, 0), window=self.table_frame, anchor="nw")

        def on_configure(event):
            self.canvas.configure(scrollregion=self.canvas.bbox("all"))

        self.table_frame.bind("<Configure>", on_configure)

        # ===== ADD MOUSE WHEEL SCROLLING =====
        def on_mouse_wheel(event):
            """Handle mouse wheel scrolling for all platforms"""
            if event.delta:
                self.canvas.yview_scroll(int(-1*(event.delta/120)), "units")
            elif event.num == 4:
                self.canvas.yview_scroll(-1, "units")
            elif event.num == 5:
                self.canvas.yview_scroll(1, "units")

        # Bind mouse wheel events
        self.canvas.bind_all("<MouseWheel>", on_mouse_wheel)
        self.canvas.bind_all("<Button-4>", on_mouse_wheel)
        self.canvas.bind_all("<Button-5>", on_mouse_wheel)
        # ===== END SCROLLING ADDITIONS =====

        # Load and display all data by default
        self.display_all_patients()

        # Button frame
        btn_frame = ttk.Frame(right_frame, padding="10 10 10 10", style='TFrame')
        btn_frame.pack(fill=tk.X)

        ttk.Button(btn_frame, text="Back to Menu", command=self.main_menu,
                  style='Blue.TButton').pack(fill=tk.X, pady=10)
    
    def display_all_patients(self, malignancy_filter=None):
        """Display all patients in a table, optionally filtered by malignancy"""
        # Clear previous data
        for widget in self.table_frame.winfo_children():
            widget.destroy()

        if not self.patient_data:
            ttk.Label(self.table_frame, text="No patient data available", style='TLabel').pack(pady=20)
            return

        # Filter by malignancy if needed and sort by file number (numeric)
        if malignancy_filter:
            data = sorted([patient for patient in self.patient_data 
                          if patient.get("MALIGNANCY") == malignancy_filter],
                         key=lambda x: int(x.get("FILE NUMBER", 0)))
        else:
            data = sorted(self.patient_data, key=lambda x: int(x.get("FILE NUMBER", 0)))

        # Create a table-like structure
        columns = ["FILE NUMBER", "NAME", "MALIGNANCY", "GENDER", "AGE ON DIAGNOSIS", "DIAGNOSIS"]

        # Set a fixed width for all cells
        cell_width = 20

        # Create header row
        header_row = ttk.Frame(self.table_frame, style='TFrame')
        header_row.pack(fill=tk.X)

        for col in columns:
            header_cell = ttk.Label(header_row, text=col, 
                                  font=('Helvetica', 10, 'bold'), 
                                  anchor="center", borderwidth=1, relief="solid", 
                                  width=cell_width)
            header_cell.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, ipadx=5, ipady=5)

        # Create data rows with alternating colors
        for i, patient in enumerate(data):
            row_frame = ttk.Frame(self.table_frame, style='TFrame')
            row_frame.pack(fill=tk.X)

            bg_color = '#f0f7ff' if i % 2 == 0 else '#e6f2ff'

            for col in columns:
                cell = tk.Label(row_frame, text=patient.get(col, ""), 
                              font=('Helvetica', 10), 
                              anchor="center", borderwidth=1, relief="solid", 
                              width=cell_width, bg=bg_color)
                cell.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, ipadx=5, ipady=5)
                # Bind double-click to view patient
                cell.bind("<Double-1>", lambda e, p=patient: self.view_patient(p))

    def apply_malignancy_filter(self):
        """Apply malignancy filter to the patient table"""
        malignancy = self.malignancy_filter.get()
        if malignancy:
            self.display_all_patients(malignancy)
    
    def open_statistics_window(self):
        """Open the enhanced statistics window with multiple analysis options"""
        if self.users[self.current_user]["role"] not in ["admin", "editor"]:
            messagebox.showerror("Access Denied", "Only admins and editors can access statistics.")
            return

        self.clear_frame()
        self.root.minsize(800, 600)

        # Main container with gradient background
        main_frame = tk.Frame(self.root, bg='white')
        main_frame.pack(expand=True, fill=tk.BOTH)
        
        # Left side with logo and app name
        left_frame = tk.Frame(main_frame, bg='#3498db')
        left_frame.pack(side=tk.LEFT, fill=tk.Y, expand=False)
        
        # App logo and name
        logo_frame = tk.Frame(left_frame, bg='#3498db')
        logo_frame.pack(expand=True, fill=tk.BOTH, padx=40, pady=40)
        
        # App name with modern font
        tk.Label(logo_frame, text="OncoCare", font=('Helvetica', 24, 'bold'), 
                bg='#3498db', fg='white').pack(pady=(0, 10))
        
        tk.Label(logo_frame, text="Advanced Statistics", 
                font=('Helvetica', 14), bg='#3498db', fg='white').pack(pady=(0, 40))
        
        # Right side with content
        right_frame = tk.Frame(main_frame, bg='white')
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        # Analysis options frame
        options_frame = ttk.Frame(right_frame, style='TFrame')
        options_frame.pack(fill=tk.X, padx=20, pady=10)

        ttk.Label(options_frame, text="Select Analysis:", 
                 font=('Helvetica', 12)).pack(side=tk.LEFT, padx=5)
        
        self.analysis_var = tk.StringVar(value="Malignancy Distribution")
        analysis_options = [
            "Malignancy Distribution",
            "Age at Diagnosis",
            "Gender Distribution",
            "Treatment Outcomes",
            "B-Symptoms Analysis",
            "Risk Group Analysis",
            "Diagnosis Timeline",
            "Survival Analysis"
        ]
        
        analysis_menu = ttk.Combobox(options_frame, textvariable=self.analysis_var,
                                   values=analysis_options, state="readonly")
        analysis_menu.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(options_frame, text="Generate", command=self.generate_analysis,
                  style='Blue.TButton').pack(side=tk.LEFT, padx=10)

        # Filter frame
        filter_frame = ttk.Frame(right_frame, style='TFrame')
        filter_frame.pack(fill=tk.X, padx=20, pady=10)

        ttk.Label(filter_frame, text="Filter by Malignancy:", 
                 font=('Helvetica', 12)).pack(side=tk.LEFT, padx=5)
        
        self.stats_malignancy_filter = ttk.Combobox(filter_frame, values=["All"] + MALIGNANCIES, 
                                                   state="readonly")
        self.stats_malignancy_filter.set("All")
        self.stats_malignancy_filter.pack(side=tk.LEFT, padx=5)

        # Date range filter
        ttk.Label(filter_frame, text="Date Range:", 
                 font=('Helvetica', 12)).pack(side=tk.LEFT, padx=5)
        
        self.start_date_var = tk.StringVar()
        ttk.Entry(filter_frame, textvariable=self.start_date_var, width=10).pack(side=tk.LEFT, padx=5)
        ttk.Button(filter_frame, text="📅", width=3,
                  command=lambda: self.show_calendar(self.start_date_var)).pack(side=tk.LEFT, padx=0)
        
        ttk.Label(filter_frame, text="to").pack(side=tk.LEFT, padx=5)
        
        self.end_date_var = tk.StringVar()
        ttk.Entry(filter_frame, textvariable=self.end_date_var, width=10).pack(side=tk.LEFT, padx=5)
        ttk.Button(filter_frame, text="📅", width=3,
                  command=lambda: self.show_calendar(self.end_date_var)).pack(side=tk.LEFT, padx=0)

        # Main content area for statistics
        self.stats_content_frame = ttk.Frame(right_frame, style='TFrame')
        self.stats_content_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

        # Button frame at bottom
        btn_frame = ttk.Frame(right_frame, style='TFrame')
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=20, pady=10)  # Changed to side=BOTTOM

        ttk.Button(btn_frame, text="Print Report", command=self.print_statistics_report,
                  style='Yellow.TButton').pack(side=tk.LEFT, padx=10)

        ttk.Button(btn_frame, text="Back to Menu", command=self.main_menu,
                  style='Blue.TButton').pack(side=tk.RIGHT, padx=10)
    
        # Generate initial analysis
        self.generate_analysis()

    def generate_analysis(self):
        """Generate the selected analysis with current filters"""
        # Clear previous content
        for widget in self.stats_content_frame.winfo_children():
            widget.destroy()

        if not self.patient_data:
            ttk.Label(self.stats_content_frame, text="No patient data available", 
                     style='TLabel').pack(pady=20)
            return

        # Apply filters
        filtered_data = self.apply_statistics_filters()

        if not filtered_data:
            ttk.Label(self.stats_content_frame, text="No data matches the selected filters", 
                     style='TLabel').pack(pady=20)
            return

        analysis_type = self.analysis_var.get()
        
        # Create main container
        container = ttk.Frame(self.stats_content_frame, style='TFrame')
        container.pack(fill=tk.BOTH, expand=True)

        # Create canvas with improved scrolling
        canvas = tk.Canvas(container, highlightthickness=0, bg='white')
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Add scrollbars
        y_scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
        y_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        x_scrollbar = ttk.Scrollbar(self.stats_content_frame, orient="horizontal", 
                                  command=canvas.xview)
        x_scrollbar.pack(side=tk.BOTTOM, fill=tk.X, pady=5, ipady=10)

        canvas.configure(yscrollcommand=y_scrollbar.set, xscrollcommand=x_scrollbar.set)

        # Create content frame
        content_frame = ttk.Frame(canvas, style='TFrame')
        canvas.create_window((0, 0), window=content_frame, anchor="nw", tags="content_frame")

        def on_frame_configure(event):
            """Update scroll region"""
            bbox = canvas.bbox("all")
            if bbox:
                canvas.configure(scrollregion=bbox)

        content_frame.bind("<Configure>", on_frame_configure)

        # UNIVERSAL MOUSE SCROLLING (NEW IMPROVEMENT)
        def on_mousewheel(event):
            """Handle mouse wheel scrolling from anywhere in the window"""
            # Get current scroll position
            x1, y1, x2, y2 = canvas.bbox("all")
            
            # Calculate scroll amount (more sensitive)
            scroll_speed = 2 if event.state & 1 else 1  # Faster with Shift key
            
            if event.delta > 0 or event.num == 4:  # Scroll up/left
                canvas.yview_scroll(-1 * scroll_speed, "units")
            elif event.delta < 0 or event.num == 5:  # Scroll down/right
                canvas.yview_scroll(1 * scroll_speed, "units")
            
            # Horizontal scrolling with Shift
            if event.state & 1:  # Shift key pressed
                if event.delta > 0:
                    canvas.xview_scroll(-1 * scroll_speed, "units")
                else:
                    canvas.xview_scroll(1 * scroll_speed, "units")

        # Bind mouse wheel events to the ENTIRE WINDOW
        self.root.bind_all("<MouseWheel>", on_mousewheel)  # Windows/Mac
        self.root.bind_all("<Button-4>", on_mousewheel)    # Linux up
        self.root.bind_all("<Button-5>", on_mousewheel)    # Linux down
        
        # Middle-click drag scrolling
        def start_drag(event):
            if event.num == 2:  # Middle mouse button
                canvas.scan_mark(event.x, event.y)
                canvas.config(cursor="fleur")

        def drag_scroll(event):
            if canvas.cget("cursor") == "fleur":
                canvas.scan_dragto(event.x, event.y, gain=1)

        def end_drag(event):
            canvas.config(cursor="")

        canvas.bind("<ButtonPress-2>", start_drag)
        canvas.bind("<B2-Motion>", drag_scroll)
        canvas.bind("<ButtonRelease-2>", end_drag)

        # Generate the selected analysis
        if analysis_type == "Malignancy Distribution":
            self.generate_malignancy_distribution(content_frame, filtered_data)
        elif analysis_type == "Age at Diagnosis":
            self.generate_age_at_diagnosis(content_frame, filtered_data)
        elif analysis_type == "Gender Distribution":
            self.generate_gender_distribution(content_frame, filtered_data)
        elif analysis_type == "Treatment Outcomes":
            self.generate_treatment_outcomes(content_frame, filtered_data)
        elif analysis_type == "B-Symptoms Analysis":
            self.generate_b_symptoms_analysis(content_frame, filtered_data)
        elif analysis_type == "Risk Group Analysis":
            self.generate_risk_group_analysis(content_frame, filtered_data)
        elif analysis_type == "Diagnosis Timeline":
            self.generate_diagnosis_timeline(content_frame, filtered_data)
        elif analysis_type == "Survival Analysis":
            self.generate_survival_analysis(content_frame, filtered_data)

        # Initial configuration
        content_frame.update_idletasks()
        canvas.configure(scrollregion=canvas.bbox("all"))
        canvas.itemconfigure("content_frame", width=max(
            content_frame.winfo_reqwidth(), 
            canvas.winfo_width()
        ))
                
    def apply_statistics_filters(self):
        """Apply the current filters to the patient data"""
        filtered_data = self.patient_data.copy()
        
        # Malignancy filter
        malignancy_filter = self.stats_malignancy_filter.get()
        if malignancy_filter != "All":
            filtered_data = [p for p in filtered_data if p.get("MALIGNANCY") == malignancy_filter]
        
        # Date range filter
        start_date = self.start_date_var.get()
        end_date = self.end_date_var.get()
        
        if start_date or end_date:
            try:
                if start_date:
                    start_date = datetime.strptime(start_date, "%d/%m/%Y")
                if end_date:
                    end_date = datetime.strptime(end_date, "%d/%m/%Y")
                
                def date_in_range(patient):
                    try:
                        # Try to get diagnosis date first, fall back to created date
                        date_str = patient.get("CREATED_DATE", "")
                        if not date_str:  # Handle empty or None
                            return False
                        
                        # Extract date part (assuming format is "dd/mm/YYYY HH:MM:SS")
                        date_part = date_str.split()[0]
                        patient_date = datetime.strptime(date_part, "%d/%m/%Y")
                        
                        if start_date and end_date:
                            return start_date <= patient_date <= end_date
                        elif start_date:
                            return patient_date >= start_date
                        elif end_date:
                            return patient_date <= end_date
                        return True
                    except:
                        return False
                
                filtered_data = [p for p in filtered_data if date_in_range(p)]
            except ValueError:
                messagebox.showerror("Error", "Invalid date format. Please use dd/mm/yyyy")
                return self.patient_data
        
        return filtered_data

    def generate_malignancy_distribution(self, parent, data):
        """Generate malignancy distribution analysis with enhanced visualizations"""
        # Calculate malignancy counts
        malignancy_counts = defaultdict(int)
        for patient in data:
            malignancy = patient.get("MALIGNANCY", "Unknown")
            malignancy_counts[malignancy] += 1

        if not malignancy_counts:
            ttk.Label(parent, text="No data available for malignancy distribution", 
                     style='TLabel').pack(pady=20)
            return

        total_patients = sum(malignancy_counts.values())
        
        # Create statistics text
        stats_text = "Malignancy Distribution Analysis:\n\n"
        stats_text += f"Total Patients: {total_patients}\n\n"
        
        for malignancy, count in sorted(malignancy_counts.items(), key=lambda x: x[1], reverse=True):
            percentage = (count / total_patients) * 100
            stats_text += f"{malignancy}: {count} patients ({percentage:.1f}%)\n"
        
        # Text summary
        text_frame = ttk.Frame(parent, style='TFrame')
        text_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(text_frame, text=stats_text, 
                 font=('Helvetica', 12), justify='left').pack(anchor='w')

        # Create visualizations frame
        viz_frame = ttk.Frame(parent, style='TFrame')
        viz_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Create pie chart
        fig1, ax1 = plt.subplots(figsize=(8, 6))
        labels = list(malignancy_counts.keys())
        sizes = list(malignancy_counts.values())
        colors = [MALIGNANCY_COLORS.get(m, '#999999') for m in labels]

        ax1.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
        ax1.axis('equal')
        ax1.set_title('Malignancy Distribution')
        fig1.tight_layout()

        # Create bar chart
        fig2, ax2 = plt.subplots(figsize=(10, 6))
        y_pos = range(len(labels))
        
        ax2.barh(y_pos, sizes, color=colors)
        ax2.set_yticks(y_pos)
        ax2.set_yticklabels(labels)
        ax2.invert_yaxis()
        ax2.set_xlabel('Number of Patients')
        ax2.set_title('Malignancy Distribution (Count)')
        fig2.tight_layout()

        # Display the plots in Tkinter
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        
        # Pie chart
        pie_canvas = FigureCanvasTkAgg(fig1, master=viz_frame)
        pie_canvas.draw()
        pie_canvas.get_tk_widget().pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Bar chart
        bar_canvas = FigureCanvasTkAgg(fig2, master=viz_frame)
        bar_canvas.draw()
        bar_canvas.get_tk_widget().pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Store figures for export/print
        self.current_figures = [fig1, fig2]

    def generate_age_at_diagnosis(self, parent, data):
        """Generate age at diagnosis analysis with distribution and statistics"""
        ages = []
        for patient in data:
            try:
                age = float(patient.get("AGE ON DIAGNOSIS", 0))
                if age > 0:
                    ages.append(age)
            except (ValueError, TypeError):
                continue

        if not ages:
            ttk.Label(parent, text="No valid age data available", 
                     style='TLabel').pack(pady=20)
            return

        # Calculate statistics
        min_age = min(ages)
        max_age = max(ages)
        mean_age = sum(ages) / len(ages)
        median_age = sorted(ages)[len(ages) // 2]
        
        # Create age groups
        age_groups = {
            "0-1 years": 0,
            "1-5 years": 0,
            "5-10 years": 0,
            "10-15 years": 0,
            "15+ years": 0
        }
        
        for age in ages:
            if age <= 1:
                age_groups["0-1 years"] += 1
            elif age <= 5:
                age_groups["1-5 years"] += 1
            elif age <= 10:
                age_groups["5-10 years"] += 1
            elif age <= 15:
                age_groups["10-15 years"] += 1
            else:
                age_groups["15+ years"] += 1

        # Create statistics text
        stats_text = "Age at Diagnosis Analysis:\n\n"
        stats_text += f"Total Patients with Age Data: {len(ages)}\n"
        stats_text += f"Minimum Age: {min_age:.1f} years\n"
        stats_text += f"Maximum Age: {max_age:.1f} years\n"
        stats_text += f"Mean Age: {mean_age:.1f} years\n"
        stats_text += f"Median Age: {median_age:.1f} years\n\n"
        stats_text += "Age Group Distribution:\n"
        
        for group, count in age_groups.items():
            percentage = (count / len(ages)) * 100
            stats_text += f"{group}: {count} patients ({percentage:.1f}%)\n"

        # Text summary
        text_frame = ttk.Frame(parent, style='TFrame')
        text_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(text_frame, text=stats_text, 
                 font=('Helvetica', 12), justify='left').pack(anchor='w')

        # Create visualizations frame
        viz_frame = ttk.Frame(parent, style='TFrame')
        viz_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Create histogram
        fig1, ax1 = plt.subplots(figsize=(8, 6))
        ax1.hist(ages, bins=20, color='#3498db', edgecolor='black')
        ax1.set_xlabel('Age at Diagnosis (years)')
        ax1.set_ylabel('Number of Patients')
        ax1.set_title('Age Distribution at Diagnosis')
        fig1.tight_layout()

        # Create age group pie chart
        fig2, ax2 = plt.subplots(figsize=(8, 6))
        ax2.pie(age_groups.values(), labels=age_groups.keys(), autopct='%1.1f%%', startangle=90)
        ax2.axis('equal')
        ax2.set_title('Age Group Distribution')
        fig2.tight_layout()

        # Display the plots in Tkinter
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        
        # Histogram
        hist_canvas = FigureCanvasTkAgg(fig1, master=viz_frame)
        hist_canvas.draw()
        hist_canvas.get_tk_widget().pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Pie chart
        pie_canvas = FigureCanvasTkAgg(fig2, master=viz_frame)
        pie_canvas.draw()
        pie_canvas.get_tk_widget().pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Store figures for export/print
        self.current_figures = [fig1, fig2]

    def generate_gender_distribution(self, parent, data):
        """Generate gender distribution analysis"""
        gender_counts = defaultdict(int)
        for patient in data:
            gender = patient.get("GENDER", "Unknown")
            gender_counts[gender] += 1

        if not gender_counts:
            ttk.Label(parent, text="No gender data available", 
                     style='TLabel').pack(pady=20)
            return

        total_patients = sum(gender_counts.values())
        
        # Create statistics text
        stats_text = "Gender Distribution Analysis:\n\n"
        stats_text += f"Total Patients: {total_patients}\n\n"
        
        for gender, count in gender_counts.items():
            percentage = (count / total_patients) * 100
            stats_text += f"{gender}: {count} patients ({percentage:.1f}%)\n"

        # Text summary
        text_frame = ttk.Frame(parent, style='TFrame')
        text_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(text_frame, text=stats_text, 
                 font=('Helvetica', 12), justify='left').pack(anchor='w')

        # Create visualizations frame
        viz_frame = ttk.Frame(parent, style='TFrame')
        viz_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Create pie chart
        fig1, ax1 = plt.subplots(figsize=(8, 6))
        ax1.pie(gender_counts.values(), labels=gender_counts.keys(), autopct='%1.1f%%', startangle=90)
        ax1.axis('equal')
        ax1.set_title('Gender Distribution')
        fig1.tight_layout()

        # Create bar chart
        fig2, ax2 = plt.subplots(figsize=(8, 6))
        ax2.bar(gender_counts.keys(), gender_counts.values(), color=['#3498db', '#e74c3c'])
        ax2.set_xlabel('Gender')
        ax2.set_ylabel('Number of Patients')
        ax2.set_title('Gender Distribution (Count)')
        fig2.tight_layout()

        # Display the plots in Tkinter
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        
        # Pie chart
        pie_canvas = FigureCanvasTkAgg(fig1, master=viz_frame)
        pie_canvas.draw()
        pie_canvas.get_tk_widget().pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Bar chart
        bar_canvas = FigureCanvasTkAgg(fig2, master=viz_frame)
        bar_canvas.draw()
        bar_canvas.get_tk_widget().pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Store figures for export/print
        self.current_figures = [fig1, fig2]

    def generate_treatment_outcomes(self, parent, data):
        """Generate treatment outcomes analysis"""
        outcome_counts = defaultdict(int)
        for patient in data:
            outcome = patient.get("STATE", "Unknown")
            outcome_counts[outcome] += 1

        if not outcome_counts:
            ttk.Label(parent, text="No treatment outcome data available", 
                     style='TLabel').pack(pady=20)
            return

        total_patients = sum(outcome_counts.values())
        
        # Create statistics text
        stats_text = "Treatment Outcomes Analysis:\n\n"
        stats_text += f"Total Patients: {total_patients}\n\n"
        
        for outcome, count in outcome_counts.items():
            percentage = (count / total_patients) * 100
            stats_text += f"{outcome}: {count} patients ({percentage:.1f}%)\n"

        # Text summary
        text_frame = ttk.Frame(parent, style='TFrame')
        text_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(text_frame, text=stats_text, 
                 font=('Helvetica', 12), justify='left').pack(anchor='w')

        # Create visualizations frame
        viz_frame = ttk.Frame(parent, style='TFrame')
        viz_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Create pie chart
        fig1, ax1 = plt.subplots(figsize=(8, 6))
        ax1.pie(outcome_counts.values(), labels=outcome_counts.keys(), autopct='%1.1f%%', startangle=90)
        ax1.axis('equal')
        ax1.set_title('Treatment Outcomes')
        fig1.tight_layout()

        # Create bar chart
        fig2, ax2 = plt.subplots(figsize=(8, 6))
        ax2.bar(outcome_counts.keys(), outcome_counts.values(), color=['#2ecc71', '#e74c3c', '#f39c12'])
        ax2.set_xlabel('Outcome')
        ax2.set_ylabel('Number of Patients')
        ax2.set_title('Treatment Outcomes (Count)')
        fig2.tight_layout()

        # Display the plots in Tkinter
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        
        # Pie chart
        pie_canvas = FigureCanvasTkAgg(fig1, master=viz_frame)
        pie_canvas.draw()
        pie_canvas.get_tk_widget().pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Bar chart
        bar_canvas = FigureCanvasTkAgg(fig2, master=viz_frame)
        bar_canvas.draw()
        bar_canvas.get_tk_widget().pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Store figures for export/print
        self.current_figures = [fig1, fig2]

    def generate_b_symptoms_analysis(self, parent, data):
        """Generate B-symptoms analysis"""
        symptom_counts = defaultdict(int)
        total_with_data = 0
        
        for patient in data:
            symptoms = patient.get("B_SYMPTOMS", "")
            if symptoms:
                total_with_data += 1
                for symptom in symptoms.split(", "):
                    symptom_counts[symptom] += 1

        if not symptom_counts:
            ttk.Label(parent, text="No B-symptoms data available", 
                     style='TLabel').pack(pady=20)
            return

        # Create statistics text
        stats_text = "B-Symptoms Analysis:\n\n"
        stats_text += f"Total Patients with B-Symptoms Data: {total_with_data}\n\n"
        
        for symptom, count in sorted(symptom_counts.items(), key=lambda x: x[1], reverse=True):
            percentage = (count / total_with_data) * 100
            stats_text += f"{symptom}: {count} patients ({percentage:.1f}%)\n"

        # Text summary
        text_frame = ttk.Frame(parent, style='TFrame')
        text_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(text_frame, text=stats_text, 
                 font=('Helvetica', 12), justify='left').pack(anchor='w')

        # Create visualizations frame
        viz_frame = ttk.Frame(parent, style='TFrame')
        viz_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Create horizontal bar chart
        fig, ax = plt.subplots(figsize=(10, 6))
        
        symptoms = list(symptom_counts.keys())
        counts = list(symptom_counts.values())
        
        y_pos = range(len(symptoms))
        ax.barh(y_pos, counts, color='#3498db')
        ax.set_yticks(y_pos)
        ax.set_yticklabels(symptoms)
        ax.invert_yaxis()
        ax.set_xlabel('Number of Patients')
        ax.set_title('B-Symptoms Distribution')
        fig.tight_layout()

        # Display the plot in Tkinter
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        
        canvas = FigureCanvasTkAgg(fig, master=viz_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # Store figure for export/print
        self.current_figures = [fig]

    def generate_risk_group_analysis(self, parent, data):
        """Generate risk group analysis"""
        risk_counts = defaultdict(int)
        total_with_data = 0
        
        for patient in data:
            risk_group = patient.get("RISK_GROUP", "")
            if risk_group:
                total_with_data += 1
                risk_counts[risk_group] += 1

        if not risk_counts:
            ttk.Label(parent, text="No risk group data available", 
                     style='TLabel').pack(pady=20)
            return

        # Create statistics text
        stats_text = "Risk Group Analysis:\n\n"
        stats_text += f"Total Patients with Risk Group Data: {total_with_data}\n\n"
        
        for group, count in sorted(risk_counts.items(), key=lambda x: x[0]):
            percentage = (count / total_with_data) * 100
            stats_text += f"{group}: {count} patients ({percentage:.1f}%)\n"

        # Text summary
        text_frame = ttk.Frame(parent, style='TFrame')
        text_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(text_frame, text=stats_text, 
                 font=('Helvetica', 12), justify='left').pack(anchor='w')

        # Create visualizations frame
        viz_frame = ttk.Frame(parent, style='TFrame')
        viz_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Create pie chart
        fig, ax = plt.subplots(figsize=(8, 6))
        ax.pie(risk_counts.values(), labels=risk_counts.keys(), autopct='%1.1f%%', startangle=90)
        ax.axis('equal')
        ax.set_title('Risk Group Distribution')
        fig.tight_layout()

        # Display the plot in Tkinter
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        
        canvas = FigureCanvasTkAgg(fig, master=viz_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # Store figure for export/print
        self.current_figures = [fig]

    def generate_diagnosis_timeline(self, parent, data):
        """Generate diagnosis timeline analysis"""
        diagnosis_dates = []
        
        for patient in data:
            try:
                date_str = patient.get("CREATED_DATE", "")
                if not date_str:
                    continue
                
                # Extract date part (assuming format is "dd/mm/YYYY HH:MM:SS")
                date_part = date_str.split()[0]
                date = datetime.strptime(date_part, "%d/%m/%Y")
                diagnosis_dates.append(date)
            except (ValueError, IndexError):
                continue

        if not diagnosis_dates:
            ttk.Label(parent, text="No valid diagnosis date data available", 
                     style='TLabel').pack(pady=20)
            return

        # Group by month
        monthly_counts = defaultdict(int)
        for date in diagnosis_dates:
            month_year = f"{date.year}-{date.month:02d}"
            monthly_counts[month_year] += 1

        # Sort by date
        sorted_months = sorted(monthly_counts.keys())
        sorted_counts = [monthly_counts[m] for m in sorted_months]

        # Create statistics text
        stats_text = "Diagnosis Timeline Analysis:\n\n"
        stats_text += f"Total Patients with Diagnosis Dates: {len(diagnosis_dates)}\n"
        stats_text += f"First Diagnosis Date: {min(diagnosis_dates).strftime('%d/%m/%Y')}\n"
        stats_text += f"Last Diagnosis Date: {max(diagnosis_dates).strftime('%d/%m/%Y')}\n\n"
        stats_text += "Monthly Diagnosis Counts:\n"
        
        for month, count in zip(sorted_months, sorted_counts):
            stats_text += f"{month}: {count} patients\n"

        # Text summary
        text_frame = ttk.Frame(parent, style='TFrame')
        text_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(text_frame, text=stats_text, 
                 font=('Helvetica', 12), justify='left').pack(anchor='w')

        # Create visualizations frame
        viz_frame = ttk.Frame(parent, style='TFrame')
        viz_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Create line chart
        fig, ax = plt.subplots(figsize=(12, 6))
        ax.plot(sorted_months, sorted_counts, marker='o', color='#3498db')
        ax.set_xlabel('Month-Year')
        ax.set_ylabel('Number of Diagnoses')
        ax.set_title('Monthly Diagnosis Counts')
        plt.xticks(rotation=45)
        fig.tight_layout()

        # Display the plot in Tkinter
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        
        canvas = FigureCanvasTkAgg(fig, master=viz_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # Store figure for export/print
        self.current_figures = [fig]

    def generate_survival_analysis(self, parent, data):
        """Generate basic survival analysis"""
        # This is a simplified survival analysis - in a real application you would need
        # more detailed data including diagnosis date, last follow-up date, and status
        
        outcomes = defaultdict(int)
        for patient in data:
            outcome = patient.get("STATE", "Unknown")
            outcomes[outcome] += 1

        if not outcomes:
            ttk.Label(parent, text="No outcome data available for survival analysis", 
                     style='TLabel').pack(pady=20)
            return

        total_patients = sum(outcomes.values())
        alive = outcomes.get("ALIVE", 0)
        deceased = outcomes.get("DECEASED", 0)
        lost = outcomes.get("MISSED FOLLOW UP", 0)
        
        # Calculate survival rate (simplified)
        if (alive + deceased) > 0:
            survival_rate = (alive / (alive + deceased)) * 100
        else:
            survival_rate = 0

        # Create statistics text
        stats_text = "Survival Analysis (Simplified):\n\n"
        stats_text += f"Total Patients: {total_patients}\n"
        stats_text += f"Alive: {alive} patients\n"
        stats_text += f"Deceased: {deceased} patients\n"
        stats_text += f"Lost to Follow-up: {lost} patients\n"
        stats_text += f"\nEstimated Survival Rate: {survival_rate:.1f}%\n"
        stats_text += "(Based on patients with known outcomes, excluding lost to follow-up)"

        # Text summary
        text_frame = ttk.Frame(parent, style='TFrame')
        text_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(text_frame, text=stats_text, 
                 font=('Helvetica', 12), justify='left').pack(anchor='w')

        # Create visualizations frame
        viz_frame = ttk.Frame(parent, style='TFrame')
        viz_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Create bar chart
        fig, ax = plt.subplots(figsize=(8, 6))
        
        categories = ["Alive", "Deceased", "Lost to Follow-up"]
        counts = [alive, deceased, lost]
        colors = ['#2ecc71', '#e74c3c', '#f39c12']
        
        ax.bar(categories, counts, color=colors)
        ax.set_xlabel('Outcome')
        ax.set_ylabel('Number of Patients')
        ax.set_title('Patient Outcomes')
        fig.tight_layout()

        # Display the plot in Tkinter
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        
        canvas = FigureCanvasTkAgg(fig, master=viz_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # Store figure for export/print
        self.current_figures = [fig]

    def print_statistics_report(self):
        """Create a Word document with the current statistics report"""
        if not hasattr(self, 'analysis_var'):
            messagebox.showerror("Error", "No analysis to print")
            return

        try:
            # Create a Word document
            doc = Document()
            
            # Add title
            doc.add_heading(f'OncoCare Statistics Report', 0)
            doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Analysis Type: {self.analysis_var.get()}")
            
            # Add filters info
            filters = []
            malignancy_filter = self.stats_malignancy_filter.get()
            if malignancy_filter != "All":
                filters.append(f"Malignancy: {malignancy_filter}")
            
            start_date = self.start_date_var.get()
            end_date = self.end_date_var.get()
            if start_date or end_date:
                date_range = f"Date Range: {start_date if start_date else 'Start'} to {end_date if end_date else 'End'}"
                filters.append(date_range)
            
            if filters:
                doc.add_paragraph("Filters Applied:")
                for f in filters:
                    doc.add_paragraph(f, style='List Bullet')
            
            # Add some basic statistics
            doc.add_heading('Summary Statistics', level=1)
            
            # Get filtered data
            filtered_data = self.apply_statistics_filters()
            doc.add_paragraph(f"Total Patients in Analysis: {len(filtered_data)}")
            
            # Add matplotlib figures to Word document
            if hasattr(self, 'current_figures'):
                for fig in self.current_figures:
                    # Save figure to a temporary image file
                    temp_img = "temp_fig.png"
                    fig.savefig(temp_img, dpi=300, bbox_inches='tight')
                    
                    # Add image to document
                    doc.add_picture(temp_img, width=Inches(6))
                    doc.add_paragraph()  # Add space after image
                    
                    # Remove temporary image
                    os.remove(temp_img)
            
            # Save to temporary file
            temp_file = "temp_statistics_report.docx"
            doc.save(temp_file)
            
            # Open the document for the user to review and save
            os.startfile(temp_file)
            messagebox.showinfo("Document Created", "Statistics report has been created. Please review and save it as needed.")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report: {str(e)}")

    def manage_users(self):
        """Manage user accounts"""
        if self.current_user != "mej.esam" and self.users[self.current_user]["role"] != "admin":
            messagebox.showerror("Access Denied", "Only admins can manage users")
            return
            
        self.clear_frame()
        
        # Main container with gradient background
        main_frame = tk.Frame(self.root, bg='white')
        main_frame.pack(expand=True, fill=tk.BOTH)
        
        # Left side with logo and app name
        left_frame = tk.Frame(main_frame, bg='#3498db')
        left_frame.pack(side=tk.LEFT, fill=tk.Y, expand=False)
        
        # App logo and name
        logo_frame = tk.Frame(left_frame, bg='#3498db')
        logo_frame.pack(expand=True, fill=tk.BOTH, padx=40, pady=40)
        
        # App name with modern font
        tk.Label(logo_frame, text="OncoCare", font=('Helvetica', 24, 'bold'), 
                bg='#3498db', fg='white').pack(pady=(0, 10))
        
        tk.Label(logo_frame, text="User Management", 
                font=('Helvetica', 14), bg='#3498db', fg='white').pack(pady=(0, 40))
        
        # Right side with content
        right_frame = tk.Frame(main_frame, bg='white')
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        # Main content
        main_frame = ttk.Frame(right_frame, padding="20 20 20 20", style='TFrame')
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # User list
        list_frame = ttk.Frame(main_frame, style='TFrame')
        list_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        scrollbar = ttk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.user_list = tk.Listbox(list_frame, yscrollcommand=scrollbar.set, 
                                  selectmode=tk.SINGLE, font=('Helvetica', 12),
                                  height=10, bg='white', bd=0, highlightthickness=0)
        self.user_list.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=self.user_list.yview)
        
        for username in self.users:
            self.user_list.insert(tk.END, f"{username} ({self.users[username]['role']})")
        
        # Add user form
        form_frame = ttk.Frame(main_frame, style='TFrame')
        form_frame.pack(fill=tk.X, pady=10)
        
        ttk.Label(form_frame, text="New Username:", 
                 font=('Helvetica', 11)).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        
        self.new_username = ttk.Entry(form_frame, font=('Helvetica', 11))
        self.new_username.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(form_frame, text="Password:", 
                 font=('Helvetica', 11)).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        
        self.new_password = ttk.Entry(form_frame, show="*", font=('Helvetica', 11))
        self.new_password.grid(row=1, column=1, padx=5, pady=5)
        
        ttk.Label(form_frame, text="Role:", 
                 font=('Helvetica', 11)).grid(row=2, column=0, padx=5, pady=5, sticky="e")
        
        self.new_role = ttk.Combobox(form_frame, values=["admin", "editor", "viewer", "pharmacist"], 
                                    font=('Helvetica', 11))
        self.new_role.grid(row=2, column=1, padx=5, pady=5)
        
        # Buttons
        btn_frame = ttk.Frame(main_frame, style='TFrame')
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="Add User", command=self.add_user, 
                  style='Blue.TButton').pack(side=tk.LEFT, padx=5)
        
        # Only mej.esam can delete users
        if self.current_user == "mej.esam":
            ttk.Button(btn_frame, text="Delete User", command=self.delete_user,
                      style='Red.TButton').pack(side=tk.LEFT, padx=5)
        
        ttk.Button(btn_frame, text="Change Password", command=self.change_user_password,
                  style='Green.TButton').pack(side=tk.LEFT, padx=5)
        
        ttk.Button(btn_frame, text="Back to Menu", command=self.main_menu,
                  style='Blue.TButton').pack(side=tk.RIGHT, padx=5)

    def add_user(self):
        """Add a new user"""
        username = self.new_username.get()
        password = self.new_password.get()
        role = self.new_role.get()
        
        if not all([username, password, role]):
            messagebox.showerror("Error", "All fields are required")
            return
        
        if username in self.users:
            messagebox.showerror("Error", "Username already exists")
            return
        
        hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        self.users[username] = {"password": hashed_password, "role": role}
        self.save_users_to_file()
        self.user_list.insert(tk.END, f"{username} ({role})")
        self.new_username.delete(0, tk.END)
        self.new_password.delete(0, tk.END)
        self.new_role.set("")
        
        messagebox.showinfo("Success", "User added successfully")
    
    def delete_user(self):
        """Delete a user"""
        selection = self.user_list.curselection()
        if not selection:
            messagebox.showerror("Error", "No user selected")
            return

        selected = self.user_list.get(selection[0])
        username = selected.split()[0]

        if username == self.current_user:
            messagebox.showerror("Error", "You cannot delete yourself.")
            return

        if username == "mej.esam":
            messagebox.showerror("Error", "Cannot delete 'mej.esam' user.")
            return

        confirm = messagebox.askyesno("Confirm Delete", f"Are you sure you want to delete the user '{username}'?")
        if not confirm:
            return

        del self.users[username]
        self.save_users_to_file()
        self.user_list.delete(selection[0])
        messagebox.showinfo("Success", f"User '{username}' deleted successfully.")

    def change_user_password(self):
        """Change a user's password"""
        selection = self.user_list.curselection()
        if not selection:
            messagebox.showerror("Error", "No user selected")
            return

        selected = self.user_list.get(selection[0])
        username = selected.split()[0]

        if username == "mej.esam" and self.current_user != "mej.esam":
            messagebox.showerror("Error", "Only 'mej.esam' can change their own password.")
            return

        self.clear_frame()

        # Header
        header_frame = ttk.Frame(self.root, style='TFrame')
        header_frame.pack(fill=tk.X, padx=20, pady=10)

        ttk.Label(header_frame, text=f"Change Password for {username}", 
                 font=('Helvetica', 18, 'bold'),
                 foreground=self.secondary_color).pack(side=tk.LEFT)

        # Form
        form_frame = ttk.Frame(self.root, padding="20 20 20 20", style='TFrame')
        form_frame.pack(fill=tk.BOTH, expand=True)

        ttk.Label(form_frame, text="New Password:", 
                 font=('Helvetica', 12)).pack(pady=5)
        
        new_password_entry = ttk.Entry(form_frame, show="*", font=('Helvetica', 12))
        new_password_entry.pack(pady=5)

        ttk.Label(form_frame, text="Confirm New Password:", 
                 font=('Helvetica', 12)).pack(pady=5)
        
        confirm_password_entry = ttk.Entry(form_frame, show="*", font=('Helvetica', 12))
        confirm_password_entry.pack(pady=5)

        # Buttons
        btn_frame = ttk.Frame(form_frame, style='TFrame')
        btn_frame.pack(pady=20)

        def save_new_password():
            new_password = new_password_entry.get()
            confirm_password = confirm_password_entry.get()

            if new_password != confirm_password:
                messagebox.showerror("Error", "Passwords do not match.")
                return

            hashed_password = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            self.users[username]["password"] = hashed_password
            self.save_users_to_file()
            messagebox.showinfo("Success", f"Password for {username} changed successfully!")
            self.manage_users()

        ttk.Button(btn_frame, text="Save", command=save_new_password, 
                  style='Blue.TButton').pack(side=tk.LEFT, padx=10)
        
        ttk.Button(btn_frame, text="Cancel", command=self.manage_users).pack(side=tk.LEFT, padx=10)

    def change_password(self):
        """Change the current user's password"""
        if self.users[self.current_user]["role"] == "viewer":
            messagebox.showerror("Access Denied", "Viewers cannot change passwords.")
            return

        self.clear_frame()

        # Header
        header_frame = ttk.Frame(self.root, style='TFrame')
        header_frame.pack(fill=tk.X, padx=20, pady=10)

        ttk.Label(header_frame, text="Change Password", 
                 font=('Helvetica', 18, 'bold'),
                 foreground=self.secondary_color).pack(side=tk.LEFT)

        # Form
        form_frame = ttk.Frame(self.root, padding="20 20 20 20", style='TFrame')
        form_frame.pack(fill=tk.BOTH, expand=True)

        ttk.Label(form_frame, text="Current Password:", 
                 font=('Helvetica', 12)).pack(pady=5)
        
        current_password_entry = ttk.Entry(form_frame, show="*", font=('Helvetica', 12))
        current_password_entry.pack(pady=5)

        ttk.Label(form_frame, text="New Password:", 
                 font=('Helvetica', 12)).pack(pady=5)
        
        new_password_entry = ttk.Entry(form_frame, show="*", font=('Helvetica', 12))
        new_password_entry.pack(pady=5)

        ttk.Label(form_frame, text="Confirm New Password:", 
                 font=('Helvetica', 12)).pack(pady=5)
        
        confirm_password_entry = ttk.Entry(form_frame, show="*", font=('Helvetica', 12))
        confirm_password_entry.pack(pady=5)

        # Buttons
        btn_frame = ttk.Frame(form_frame, style='TFrame')
        btn_frame.pack(pady=20)

        def save_new_password():
            current_password = current_password_entry.get()
            new_password = new_password_entry.get()
            confirm_password = confirm_password_entry.get()

            stored_hash = self.users[self.current_user]["password"].encode('utf-8')
            if not bcrypt.checkpw(current_password.encode('utf-8'), stored_hash):
                messagebox.showerror("Error", "Current password is incorrect.")
                return

            if new_password != confirm_password:
                messagebox.showerror("Error", "New passwords do not match.")
                return

            hashed_password = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            self.users[self.current_user]["password"] = hashed_password
            self.save_users_to_file()
            messagebox.showinfo("Success", "Password changed successfully!")
            self.main_menu()

        ttk.Button(btn_frame, text="Save", command=save_new_password, 
                  style='Blue.TButton').pack(side=tk.LEFT, padx=10)
        
        ttk.Button(btn_frame, text="Cancel", command=self.main_menu).pack(side=tk.LEFT, padx=10)
    
    def sync_data(self):
        """Full data synchronization implementation with proper progress handling"""
        if not self.internet_connected:
            messagebox.showerror("Error", "No internet connection available.")
            return
        
        if self.sync_in_progress:
            messagebox.showinfo("Info", "Sync already in progress.")
            return
        
        confirm = messagebox.askyesno("Confirm Sync", 
            "This will synchronize all data with Firebase and Google Drive.\n"
            "This may take some time depending on your internet speed and amount of data.\n\n"
            "Do you want to continue?")
        if not confirm:
            return
        
        self.sync_in_progress = True
        self.sync_status.config(text="Syncing... Please wait")
        self.sync_btn.config(state=tk.DISABLED)
        
        try:
            # Create progress window
            self.sync_progress_window = tk.Toplevel(self.root)
            self.sync_progress_window.title("Synchronization Progress")
            self.sync_progress_window.geometry("500x250")
            self.sync_progress_window.grab_set()
            
            # Progress UI components
            progress_frame = ttk.Frame(self.sync_progress_window)
            progress_frame.pack(pady=20, padx=20, fill=tk.BOTH, expand=True)
            
            tk.Label(progress_frame, text="Synchronization Progress", 
                    font=('Helvetica', 14, 'bold')).pack(pady=10)
            
            self.sync_status_label = tk.Label(progress_frame, 
                                            text="Initializing synchronization...",
                                            wraplength=450)
            self.sync_status_label.pack(pady=5)
            
            self.sync_progress_bar = ttk.Progressbar(
                progress_frame,
                orient='horizontal',
                length=400,
                mode='determinate'
            )
            self.sync_progress_bar.pack(pady=10)
            
            self.sync_details_label = tk.Label(progress_frame, 
                                             text="",
                                             fg='#666666')
            self.sync_details_label.pack(pady=5)
            
            close_btn = ttk.Button(progress_frame, 
                                 text="Close", 
                                 state=tk.DISABLED,
                                 command=self.sync_progress_window.destroy)
            close_btn.pack(pady=10)

            def update_progress(message, detail="", step=None, total_steps=4):
                """Nested function to update progress safely"""
                if step is not None:
                    progress = (step/total_steps)*100
                    self.sync_progress_bar['value'] = progress
                self.sync_status_label.config(text=message)
                self.sync_details_label.config(text=detail)
                self.sync_progress_window.update()

            def sync_task():
                """Main synchronization task"""
                try:
                    # Sync patients with Firebase
                    update_progress("Syncing patient data...", step=1)
                    fb_success, fb_msg = self.firebase.sync_patients(self.patient_data)
                    
                    if fb_success:
                        firebase_data = self.firebase.get_all_patients()
                        if firebase_data:
                            self.patient_data = sorted(firebase_data, 
                                                     key=lambda x: int(x.get("FILE NUMBER", 0)))
                            self.save_patient_data()
                    
                    # Sync users
                    update_progress("Syncing user accounts...", step=2)
                    users_list = [{'username': k, **v} for k, v in self.users.items()]
                    fb_user_success, fb_user_msg = self.firebase.sync_users(users_list)
                    
                    if fb_user_success:
                        firebase_users = self.firebase.get_all_users()
                        self.users = {u['username']: u for u in firebase_users}
                        self.save_users_to_file()
                    
                    # Sync dropdowns
                    update_progress("Syncing dropdown lists...", step=3)
                    try:
                        if self.internet_connected:  # Ensure you have this flag available
                            # Get fresh data from Firebase
                            firebase_dropdowns = self.firebase.get_all_dropdowns()
        
                            # Create temporary file
                            temp_path = os.path.join(tempfile.gettempdir(), f"temp_dropdowns_{os.getpid()}.json")
        
                            # Write new data to temp file
                            with open(temp_path, 'w') as f:
                                json.dump(firebase_dropdowns, f, indent=4)
        
                            # Atomic replacement
                            shutil.move(temp_path, DROPDOWN_FILE)
                            dd_success = True
                            dd_msg = "Dropdown lists updated from Firebase"
                        else:
                            dd_success = True
                            dd_msg = "Offline - Using local dropdown lists"
                    except Exception as e:
                        dd_success = False
                        dd_msg = f"Dropdown sync failed: {str(e)}"
                        # Clean up temp file if exists
                        if os.path.exists(temp_path):
                            os.remove(temp_path)

                    # Google Drive Sync
                    update_progress("Syncing with Google Drive...", 
                                   "This may take several minutes", 
                                   step=4)
                    if self.drive.initialized:
                        try:
                            for idx, patient in enumerate(self.patient_data):
                                update_progress(
                                    "Uploading patient data to Google Drive...",
                                    f"Patient {idx+1}/{len(self.patient_data)}",
                                    step=4
                                )
                                self.drive.upload_patient_data(patient)
                            
                            total_patients = len(self.patient_data)
                            for idx, patient in enumerate(self.patient_data):
                                file_no = patient["FILE NUMBER"]
                                update_progress(
                                    "Syncing patient files...",
                                    f"Patient {file_no} ({idx+1}/{total_patients})",
                                    step=4
                                )
                                local_folder = f"Patient_{file_no}"
                                os.makedirs(local_folder, exist_ok=True)
                                self.drive.sync_patient_files(file_no)
                        
                        except Exception as drive_error:
                            print(f"Google Drive sync error: {drive_error}")
                    
                    # Finalize
                    self.last_sync_time = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
                    update_progress("Synchronization complete!", 
                                  f"Last sync: {self.last_sync_time}", 
                                  step=4)
                    close_btn.config(state=tk.NORMAL)
                    self.sync_in_progress = False
                    
                except Exception as e:
                    update_progress("Synchronization failed!", 
                                  f"Error: {str(e)}", 
                                  step=4)
                    close_btn.config(state=tk.NORMAL)
                    self.sync_in_progress = False
                    raise

            # Start sync thread with proper reference to update_progress
            import threading
            sync_thread = threading.Thread(target=sync_task, daemon=True)
            sync_thread.start()
            
        except Exception as e:
            messagebox.showerror("Sync Error", f"Failed to start sync: {str(e)}")
            self.sync_in_progress = False
            self.sync_btn.config(state=tk.NORMAL)
            if hasattr(self, 'sync_progress_window'):
                self.sync_progress_window.destroy()

    def update_sync_progress(self, message):
        """Update the sync progress window with current status"""
        if hasattr(self, 'sync_progress_label') and self.sync_progress_label.winfo_exists():
            self.sync_progress_label.config(text=message)
            self.sync_progress_window.update()
    
    def sync_complete(self, success, message):
        """Handle sync completion with detailed feedback"""
        self.sync_in_progress = False
        self.sync_status.config(text="")
        self.sync_btn.config(state=tk.NORMAL)
        
        if hasattr(self, 'sync_progress_window') and self.sync_progress_window.winfo_exists():
            self.sync_progress_bar.stop()
            self.sync_progress_label.config(text="Sync completed!")
            self.sync_close_btn.config(state=tk.NORMAL)
        
        if success:
            messagebox.showinfo("Sync Complete", message)
        else:
            messagebox.showerror("Sync Failed", message)

if __name__ == "__main__":
    root = tk.Tk()
    app = OncologyApp(root)
    root.mainloop()
